import matplotlib.pyplot as plt
from shapely.geometry import Point, Polygon, MultiPoint
from pyproj import Transformer, CRS
import numpy as np
import contextily as cx
import sqlite3
import datetime
import calendar
import os
import sys
import locale 

# Importar para manejar documentos Word
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN DE LA BASE DE DATOS ---
DB_NAME = "inspecciones_anp.db"

def conectar_db():
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        application_path = os.path.dirname(sys.executable)
    else:
        application_path = os.path.dirname(os.path.abspath(__file__))
    db_path = os.path.join(application_path, DB_NAME)
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row 
    return conn

def inicializar_db():
    conn = conectar_db()
    cursor = conn.cursor()
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS observaciones_embarcaciones (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        matricula TEXT NOT NULL,
        nombre_embarcacion TEXT,
        timestamp TEXT NOT NULL,
        latitud_wgs84 REAL NOT NULL,
        longitud_wgs84 REAL NOT NULL,
        tipo_embarcacion_id TEXT,
        estatus_categoria_id TEXT,
        notas_adicionales TEXT,
        nombre_patron TEXT
    )
    """)
    conn.commit()
    conn.close()
    application_path = os.path.dirname(sys.executable if getattr(sys,'frozen',False) else os.path.abspath(__file__))
    db_path_display = os.path.join(application_path, DB_NAME)
    print(f"Base de datos '{DB_NAME}' inicializada/verificada en: {db_path_display}")

def agregar_observacion_db(matricula, nombre_embarcacion, avistamiento_timestamp, lat_wgs84, lon_wgs84, tipo_emb_id, estatus_cat_id, notas="", nombre_patron=""):
    conn = conectar_db()
    cursor = conn.cursor()
    try:
        cursor.execute("""
        INSERT INTO observaciones_embarcaciones
        (matricula, nombre_embarcacion, timestamp, latitud_wgs84, longitud_wgs84, tipo_embarcacion_id, estatus_categoria_id, notas_adicionales, nombre_patron)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (matricula.upper(), nombre_embarcacion, avistamiento_timestamp, lat_wgs84, lon_wgs84, tipo_emb_id, estatus_cat_id, notas, nombre_patron))
        conn.commit()
        print(f"Observación para '{matricula}' (Avistamiento: {avistamiento_timestamp}) guardada.")
    except sqlite3.Error as e:
        print(f"Error al guardar observación en la base de datos: {e}")
    finally:
        conn.close()

def buscar_historial_embarcacion(matricula):
    conn = conectar_db()
    cursor = conn.cursor()
    cursor.execute("""
    SELECT * FROM observaciones_embarcaciones 
    WHERE matricula = ? 
    ORDER BY timestamp DESC
    """, (matricula.upper(),))
    registros = cursor.fetchall()
    conn.close()
    return [dict(row) for row in registros]

def buscar_por_nombre_o_patron(nombre_embarcacion, nombre_patron):
    conn = conectar_db()
    cursor = conn.cursor()
    cursor.execute("""
    SELECT * FROM observaciones_embarcaciones 
    WHERE LOWER(nombre_embarcacion) = LOWER(?) OR LOWER(nombre_patron) = LOWER(?)
    ORDER BY timestamp DESC
    """, (nombre_embarcacion, nombre_patron))
    registros = cursor.fetchall()
    conn.close()
    return [dict(row) for row in registros]

def eliminar_observacion_db(id_observacion):
    conn = conectar_db()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM observaciones_embarcaciones WHERE id = ?", (id_observacion,))
        conn.commit()
        if cursor.rowcount > 0:
            print(f"Observación con ID {id_observacion} eliminada exitosamente.")
            return True
        else:
            print(f"No se encontró observación con ID {id_observacion} para eliminar.")
            return False
    except sqlite3.Error as e:
        print(f"Error al eliminar observación de la base de datos: {e}")
        return False
    finally:
        conn.close()

# --- TRANSFORMADORES DE COORDENADAS ---
crs_utm_anp = CRS("EPSG:32613")
crs_geo = CRS("EPSG:4326")
crs_mercator = CRS("EPSG:3857")
try:
    transformer_utm_to_geo = Transformer.from_crs(crs_utm_anp, crs_geo, always_xy=True)
    transformer_geo_to_utm = Transformer.from_crs(crs_geo, crs_utm_anp, always_xy=True)
    transformer_geo_to_mercator = Transformer.from_crs(crs_geo, crs_mercator, always_xy=True)
    transformer_mercator_to_geo = Transformer.from_crs(crs_mercator, crs_geo, always_xy=True)
except Exception as e:
    print(f"Error crítico al crear los transformadores de coordenadas: {e}"); exit()

# --- FUNCIONES AUXILIARES ---
def transform_coords_list(coords_list, transformer):
    if not coords_list: return []
    if not all(isinstance(coord, tuple) and len(coord) == 2 for coord in coords_list):
        if all(isinstance(coord, Point) for coord in coords_list):
             x_coords = np.array([coord.x for coord in coords_list])
             y_coords = np.array([coord.y for coord in coords_list])
        else: return []
    else:
        x_coords = np.array([coord[0] for coord in coords_list])
        y_coords = np.array([coord[1] for coord in coords_list])
    if x_coords.size == 0: return []
    trans_x, trans_y = transformer.transform(x_coords, y_coords)
    return list(zip(trans_x, trans_y))

def dd_to_gmm_str(dd_val, is_latitude):
    hemisphere = ('N' if dd_val >= 0 else 'S') if is_latitude else ('E' if dd_val >= 0 else 'W')
    abs_dd = abs(dd_val); degrees = int(abs_dd); minutes_decimal = (abs_dd - degrees) * 60
    return f"{degrees}°{minutes_decimal:.3f}' {hemisphere}"

def gms_to_dd(grados, minutos, segundos, hemisferio):
    try:
        dd = float(grados) + (float(minutos) / 60.0) + (float(segundos) / 3600.0)
        if hemisferio.upper() in ['S', 'O', 'W']: dd *= -1
        elif hemisferio.upper() not in ['N', 'E']: raise ValueError(f"Hemisferio GMS '{hemisferio}' no reconocido.")
        return dd
    except ValueError: raise ValueError("Valores de GMS inválidos.")

def gdm_to_dd(grados_str, minutos_decimales_str, hemisferio):
    try:
        grados = float(grados_str)
        minutos_decimales = float(minutos_decimales_str)
        dd = grados + (minutos_decimales / 60.0)
        if hemisferio.upper() in ['S', 'O', 'W']:
            dd *= -1
        elif hemisferio.upper() not in ['N', 'E']:
            raise ValueError(f"Hemisferio GDM '{hemisferio}' no reconocido. Usar N, S, E, W.")
        return dd
    except ValueError:
        raise ValueError("Valores de grados o minutos decimales (GDM) inválidos.")

# --- DEFINICIONES DE COLORES, MARCADORES Y ESTATUS ---
STATUS_CATEGORIES_INSIDE_ANP = {
    1: {"id": "paso_inocente", "desc": "Paso Inocente", "color_key": "blanco"},
    2: {"id": "turistico_autorizado", "desc": "Turístico Autorizado", "color_key": "verde"},
    3: {"id": "investigacion", "desc": "Investigación Autorizada", "color_key": "azul_marino"},
    4: {"id": "doc_nav_issue", "desc": "Inconsistencias Doc. / Nav.", "color_key": "amarillo"},
    5: {"id": "pesca_lgpas_issue", "desc": "Infracción LGPAS (Pesca/Acuacultura)", "color_key": "anaranjado"},
    6: {"id": "delito", "desc": "Delito Ambiental / Otro", "color_key": "rojo"},
}
STATUS_COLORS = {
    "blanco": "white",
    "verde": "green",
    "azul_marino": "darkblue",
    "amarillo": "yellow",
    "anaranjado": "orange",
    "rojo": "red",
    "outside_anp": "deepskyblue",
    "unknown_status": "lightgray"
}
VESSEL_TYPES = {
    1: {"id": "panga", "desc": "Panga / Emb. Menor", "marker_char": "^", "size_factor": 100},
    2: {"id": "yate", "desc": "Yate / Emb. Mayor", "marker_char": "s", "size_factor": 150},
    3: {"id": "otra", "desc": "Otra / No especificada", "marker_char": "o", "size_factor": 80}
}
DEFAULT_VESSEL_TYPE_INFO = {"id": "default", "desc": "Desconocido", "marker_char": "o", "size_factor": 80}

# --- COORDENADAS UTM ORIGINALES DEL ANP (Completas) ---
anp_maritime_boundary_coords_utm = [
    (327983.720703, 2441179.856690), (406635.043518, 2359341.216920),
    (368280.494690, 2319018.426330), (287871.587708, 2401061.958310),
    (327983.720703, 2441179.856690)
]
isla_maria_madre_coords_utm = [
    (340647.669678, 2396167.727910), (340579.711304, 2396063.339720), (340611.461304, 2395997.193730),
    (340587.648682, 2395913.849670), (340595.586304, 2395900.620480), (340645.857117, 2395890.037290),
    (340611.461304, 2395736.578670), (340628.659302, 2395618.838680), (340570.450684, 2395366.161070),
    (340525.471497, 2395249.744320), (340471.231689, 2395225.931700), (340487.543884, 2395182.234130),
    (340368.570129, 2395208.520080), (340324.400085, 2395007.740110), (340479.769714, 2394974.235470),
    (340393.179688, 2394903.139530), (340373.335876, 2394827.733090), (340450.065125, 2394766.878720),
    (340397.148315, 2394584.315920), (340256.918884, 2394516.846920), (340233.106323, 2394463.930300),
    (340152.408325, 2394445.409300), (340148.439514, 2394399.107120), (340182.835510, 2394362.065490),
    (340312.481506, 2394327.669490), (340373.335876, 2394281.367310), (340362.752502, 2394200.669310),
    (340446.096313, 2394192.731690), (340432.867126, 2394129.231690), (340327.033691, 2393958.575070),
    (340280.731506, 2393929.470890), (340268.825073, 2393908.304080), (340129.918701, 2393850.095700),
    (340141.824890, 2393802.470520), (340083.937500, 2393740.018130), (340033.490112, 2393757.320130),
    (339996.430115, 2393660.810120), (339967.940125, 2393661.430110), (339604.670105, 2393789.380130),
    (339527.150085, 2393658.210080), (339518.860107, 2393626.080080), (339516.220093, 2393591.930110),
    (339584.590088, 2393563.660100), (339650.840088, 2393463.400090), (339887.670105, 2393439.740110),
    (340161.430115, 2393477.410100), (340297.929321, 2393484.969910), (340489.607100, 2393022.572000),
    (340412.379000, 2392975.619700), (340361.409100, 2392954.032300), (340321.785500, 2392931.381300),
    (340295.962100, 2392918.257900), (340290.458700, 2392910.214500), (340269.927000, 2392865.976100),
    (340255.322000, 2392819.409400), (340243.892000, 2392791.892600), (340230.980300, 2392769.244300),
    (340215.105300, 2392759.719200), (340187.800200, 2392760.777600), (340168.593600, 2392758.504800),
    (340127.318500, 2392710.879700), (340112.203000, 2392694.661800), (340098.476700, 2392672.512400),
    (340074.983900, 2392646.056700), (340039.476900, 2392595.735600), (340036.037100, 2392587.001500),
    (340042.598800, 2392550.806500), (340065.458800, 2392523.713100), (340072.867200, 2392501.064700),
    (340077.698000, 2392477.208100), (340080.063900, 2392465.716300), (340076.677200, 2392463.388000),
    (340055.933800, 2392472.489700), (340049.372100, 2392473.336300), (340037.307100, 2392463.599600),
    (340057.819400, 2392442.841800), (340062.006000, 2392433.844500), (340061.159400, 2392424.107800),
    (340061.177000, 2392414.711600), (340047.829000, 2392415.398000), (340032.986600, 2392421.967400),
    (340029.475400, 2392418.726200), (340031.380400, 2392404.121200), (340038.365500, 2392383.589500),
    (340007.825700, 2392369.514200), (340002.805400, 2392359.882800), (340003.440400, 2392353.321100),
    (340018.680400, 2392324.322700), (340029.119900, 2392301.251200), (340036.102300, 2392288.919400),
    (340046.766000, 2392275.041300), (340037.942100, 2392262.092600), (340036.989600, 2392259.870100),
    (340040.799600, 2392257.647600), (340063.310200, 2392241.197300), (340060.715200, 2392232.911600),
    (340028.624400, 2392200.293200), (340015.895800, 2392181.650300), (339983.620400, 2392158.470700),
    (339965.552000, 2392150.173600), (339954.756900, 2392150.491100), (339942.374400, 2392157.476100),
    (339938.405700, 2392156.682400), (339936.818200, 2392137.632300), (339919.831900, 2392104.771000),
    (339939.020600, 2392084.476200), (339937.134600, 2392080.204300), (339911.418100, 2392080.799700),
    (339912.211900, 2392075.402200), (339925.388100, 2392064.607200), (339940.568600, 2392056.879300),
    (339941.277200, 2392048.862500), (339909.033400, 2392023.385200), (339908.651200, 2392005.784300),
    (339922.730700, 2391995.347100), (339944.120700, 2391991.740800), (339965.710700, 2391962.848200),
    (339996.368800, 2391935.821500), (340012.285400, 2391933.116600), (340015.928400, 2391935.073400),
    (340036.299500, 2391932.656400), (340036.593700, 2391910.177000), (340038.724700, 2391895.429900),
    (340044.230000, 2391884.233600), (340067.388800, 2391852.406400), (340078.899700, 2391843.468000),
    (340079.534700, 2391831.244200), (340072.769000, 2391832.081600), (340056.039700, 2391848.865500),
    (340034.535600, 2391848.404200), (340020.955800, 2391857.279300), (339987.300800, 2391874.900500),
    (339979.065600, 2391870.153000), (339987.955600, 2391845.705500), (339999.610500, 2391840.840700),
    (339990.879300, 2391836.342700), (340004.227500, 2391814.308200), (340023.198200, 2391823.797900),
    (340009.929300, 2391845.074000), (340007.010100, 2391843.477700), (339992.698300, 2391852.199200),
    (339993.015800, 2391857.279300), (340016.543900, 2391847.345900), (340033.900600, 2391843.747600),
    (340047.024000, 2391843.959200), (340057.131600, 2391841.214800), (340066.896700, 2391823.734500),
    (340075.493200, 2391821.945800), (340085.018200, 2391825.650000), (340089.102700, 2391828.828400),
    (340087.491400, 2391843.791800), (340084.344500, 2391845.818100), (340072.472200, 2391853.463100),
    (340061.777700, 2391871.174100), (340053.269200, 2391890.081400), (340041.236600, 2391908.580700),
    (340042.825000, 2391925.748300), (340041.017300, 2391938.588400), (340035.482100, 2391941.921100),
    (340029.687100, 2391943.956900), (340027.147100, 2391939.988200), (339995.147900, 2391941.743200),
    (339983.332000, 2391953.481900), (339968.885700, 2391965.229500), (339949.412100, 2391992.424300),
    (339943.803200, 2391996.027000), (339925.229400, 2391998.884500), (339912.211900, 2392005.869600),
    (339912.826100, 2392019.582900), (339943.485700, 2392040.953400), (339946.025700, 2392057.622200),
    (339944.279400, 2392060.638400), (339926.499400, 2392067.940900), (339919.196900, 2392074.132200),
    (339930.150700, 2392074.132200), (339938.246900, 2392077.783400), (339943.803200, 2392078.418400),
    (339945.708200, 2392084.609700), (339933.711500, 2392095.572300), (339925.388100, 2392107.311000),
    (339938.324900, 2392126.729100), (339939.389800, 2392129.708000), (339942.691900, 2392151.284800),
    (339960.630700, 2392144.299800), (339985.810800, 2392154.601400), (340018.796800, 2392178.369800),
    (340039.329100, 2392206.108200), (340055.922600, 2392219.451200), (340069.641900, 2392235.258700),
    (340068.263400, 2392242.883800), (340065.312800, 2392247.591800), (340043.339600, 2392260.187600),
    (340052.809200, 2392270.746200), (340052.547100, 2392275.745100), (340038.234200, 2392299.114700),
    (340021.643800, 2392326.439400), (340006.455700, 2392356.684200), (340013.690700, 2392363.109200),
    (340022.854200, 2392371.247000), (340043.183000, 2392375.030800), (340044.924900, 2392383.375200),
    (340036.672100, 2392404.544500), (340035.402100, 2392415.551200), (340047.890500, 2392410.259500),
    (340063.553800, 2392412.164500), (340066.305500, 2392431.637900), (340065.791800, 2392439.425500),
    (340042.598800, 2392463.388000), (340052.970500, 2392469.103000), (340073.076600, 2392460.422000),
    (340080.910500, 2392458.943000), (340085.910400, 2392465.840500), (340078.111300, 2392504.006200),
    (340069.903800, 2392526.464800), (340047.467100, 2392553.346500), (340042.138000, 2392575.950800),
    (340041.185500, 2392588.015800), (340041.961600, 2392592.078900), (340064.797200, 2392622.610600),
    (340077.698000, 2392642.625900), (340095.725000, 2392659.389000), (340115.006300, 2392686.786200),
    (340129.964400, 2392707.175500), (340164.375700, 2392746.607400), (340171.995700, 2392751.369900),
    (340180.830600, 2392754.932900), (340216.586900, 2392753.369200), (340233.943600, 2392764.375900),
    (340249.395300, 2392789.987600), (340259.784700, 2392818.362500), (340275.342200, 2392865.828800),
    (340296.456000, 2392909.485200), (340323.705900, 2392923.935900), (340362.813600, 2392949.014000),
    (340418.058700, 2392971.397800), (340445.046300, 2392986.637800), (340487.116700, 2393014.422800),
    (340491.723900, 2393017.465600), (340554.575684, 2392865.843690), (340735.815674, 2392689.895510),
    (340799.315918, 2392570.832700), (340914.409729, 2392618.457700), (340956.743286, 2392580.093080),
    (341315.254272, 2392087.967100), (341263.660522, 2392045.633730), (341321.868896, 2391935.831480),
    (341380.077271, 2392013.883730), (341389.337891, 2392037.696290), (341409.181702, 2392068.123290),
    (341415.997925, 2392074.462520), (341423.062317, 2392059.846680), (341452.431091, 2392036.431090),
    (341517.121887, 2392039.209110), (341533.790710, 2392066.196720), (341543.315674, 2392033.652890),
    (341538.156311, 2391976.502690), (341546.093689, 2391966.183900), (341556.015686, 2391974.121520),
    (341568.715698, 2392018.571470), (341602.450073, 2392075.324890), (341612.220276, 2392084.024110),
    (341634.077881, 2392050.925480), (341727.563110, 2392140.209110),
    (343552.698730, 2389885.991700), (343507.912109, 2389799.561280), (343518.495483, 2389691.081910),
    (343379.588928, 2389615.675480), (343288.307495, 2389620.967100), (343190.411499, 2389757.227720),
    (343072.671509, 2389714.894290), (343126.911316, 2389597.154480), (342863.650330, 2389430.466670),
    (342888.785889, 2389097.091130), (342875.556702, 2389093.122310), (342862.327271, 2388978.028320),
    (342916.567078, 2388898.653080), (342921.858704, 2388836.475890), (343022.400696, 2388661.850520),
    (343048.859131, 2388582.475520), (343100.876892, 2388586.024900), (343201.853088, 2388431.768920),
    (340521.691101, 2384841.964480), (341115.563904, 2386033.625120), (339699.978699, 2386665.757080),
    (339141.029114, 2385565.767700), (339139.558472, 2385562.847720), (333486.840088, 2388086.395690),
    (333487.039673, 2388086.762510), (333846.045471, 2388718.355900), (333589.775696, 2388845.270320),
    (333089.439514, 2388318.086910), (333088.560913, 2388317.270080), (331178.738525, 2391778.943910),
    (331506.668701, 2391960.168090), (331401.720276, 2392177.387330), (331079.552490, 2392206.675110),
    (330913.667908, 2392089.244320),
    (340647.669678, 2396167.727910)
]
puerto_balleto_coords_utm = [
    (340647.669623, 2396167.727920), (340871.403535, 2393096.078400), (340981.527284, 2393138.588180),
    (340974.280217, 2393158.040830), (340982.290133, 2393161.473650), (340989.537199, 2393142.402420),
    (341020.051165, 2393153.463730), (341042.173790, 2393168.720720), (341044.080913, 2393179.019180),
    (341005.938456, 2393236.232860), (341015.855495, 2393240.809960), (341066.966387, 2393165.669320),
    (341057.812197, 2393158.803680), (341046.369460, 2393161.855070), (341025.009684, 2393146.598090),
    (340992.588596, 2393134.392510), (340999.835663, 2393115.321280), (340992.970021, 2393113.795580),
    (340984.197256, 2393132.485380), (340873.875105, 2393088.869650), (341727.563090, 2392140.209000),
    (341634.077853, 2392050.925350), (341612.220175, 2392084.024120), (341602.450143, 2392075.324770),
    (341568.715700, 2392018.571530), (341556.015675, 2391974.121450), (341546.093780, 2391966.183930),
    (341538.156264, 2391976.502700), (341543.315650, 2392033.652810), (341533.790631, 2392066.196630),
    (341517.121847, 2392039.209080), (341452.431093, 2392036.430940), (341423.062284, 2392059.846620),
    (341415.997933, 2392074.462520), (341409.181570, 2392068.123300), (341389.337780, 2392037.696150),
    (341380.077345, 2392013.883610), (341321.868895, 2391935.831370), (341263.660445, 2392045.633670),
    (341315.254298, 2392087.967090), (340956.743165, 2392580.093070), (340914.409747, 2392618.457730),
    (340799.315766, 2392570.832640), (340735.815639, 2392689.895380), (340554.575694, 2392865.843640),
    (340297.929347, 2393484.969880), (340161.430002, 2393477.410000), (339887.670000, 2393439.740000),
    (339650.840002, 2393463.400000), (339584.590002, 2393563.660000), (339516.219999, 2393591.930000),
    (339518.860003, 2393626.080000), (339527.150000, 2393658.210000), (339604.670000, 2393789.380000),
    (339967.940001, 2393661.430000), (339996.430002, 2393660.810000), (340033.490000, 2393757.320000),
    (340083.937515, 2393740.018110), (340141.824868, 2393802.470520), (340129.918594, 2393850.095610),
    (340268.825122, 2393908.304060), (340280.731396, 2393929.470770), (340327.033572, 2393958.575000),
    (340432.867117, 2394129.231590), (340446.096310, 2394192.731710), (340362.752393, 2394200.669230),
    (340373.335748, 2394281.367310), (340312.481459, 2394327.669480), (340182.835367, 2394362.065390),
    (340148.439465, 2394399.107130), (340152.408223, 2394445.409300), (340233.106301, 2394463.930170),
    (340256.918848, 2394516.846950), (340397.148295, 2394584.315830), (340450.065068, 2394766.878700),
    (340373.335748, 2394827.732980), (340393.179538, 2394903.139390), (340479.769674, 2394974.235460),
    (340324.400000, 2395007.740000), (340368.570002, 2395208.520000), (340487.543742, 2395182.234030),
    (340471.231777, 2395225.931700), (340525.471469, 2395249.744250), (340570.450725, 2395366.161140),
    (340628.659175, 2395618.838730), (340611.461224, 2395736.578550), (340645.857126, 2395890.037190),
    (340595.586192, 2395900.620550), (340587.648676, 2395913.849740), (340611.461224, 2395997.193730),
    (340579.711161, 2396063.339620),
    (340647.669623, 2396167.727920)
]
islas_menores_data_utm = {
    "Isla San Juanito (V)": {"coords": [(328030.906100, 2404586.986800), (327633.883600, 2404310.100800)], "marker": "o", "color": "darkviolet"},
    "Islote El Morro (V)": {"coords": [(323879.578800, 2405116.471600), (323867.899400, 2405109.464000)], "marker": "s", "color": "firebrick"},
    "Isla María Magdalena (V)": {"coords": [(350862.351900, 2377854.999000), (351247.036700, 2377788.188800)], "marker": "^", "color": "olive"},
    "Isla María Cleofas (V)": {"coords": [
        (369429.510100, 2359596.764100), (369848.633600, 2359593.460300), (369855.558600, 2359590.690300),
        (369866.088000, 2359593.322700), (369908.061600, 2359592.991800), (372259.125600, 2359574.458900),
        (372486.810100, 2359572.664100)
    ], "marker": "P", "color": "teal"},
    "Islote La Mona 1 (V)": {"coords": [(366477.352800, 2358421.914700), (366325.153200, 2358358.393600)], "marker": "*", "color": "darkorange"},
    "Islote La Mona 2 (V)": {"coords": [(367454.190600, 2356053.133500), (367435.503600, 2356048.461700)], "marker": "X", "color": "darkmagenta"},
    "Islote La Mona 3 (V)": {"coords": [(368306.785600, 2355808.450400), (368277.587200, 2355808.450400)], "marker": "D", "color": "navy"},
}

# --- TRANSFORMACIONES GLOBALES INICIALES ---
print("Cargando y transformando geometrías base del ANP...")
try:
    anp_maritime_boundary_coords_geo = transform_coords_list(anp_maritime_boundary_coords_utm, transformer_utm_to_geo)
    anp_maritime_polygon_geo = Polygon(anp_maritime_boundary_coords_geo) if anp_maritime_boundary_coords_geo else Polygon()
    anp_maritime_boundary_coords_mercator = transform_coords_list(anp_maritime_boundary_coords_geo, transformer_geo_to_mercator)
    isla_maria_madre_coords_geo = transform_coords_list(isla_maria_madre_coords_utm, transformer_utm_to_geo)
    isla_maria_madre_polygon_geo = Polygon(isla_maria_madre_coords_geo) if isla_maria_madre_coords_geo else Polygon()
    isla_maria_madre_coords_mercator = transform_coords_list(isla_maria_madre_coords_geo, transformer_geo_to_mercator)
    puerto_balleto_coords_geo = transform_coords_list(puerto_balleto_coords_utm, transformer_utm_to_geo)
    puerto_balleto_polygon_geo = Polygon(puerto_balleto_coords_geo) if puerto_balleto_coords_geo else Polygon()
    puerto_balleto_coords_mercator = transform_coords_list(puerto_balleto_coords_geo, transformer_geo_to_mercator)
    islas_menores_data_geo = { name: {"coords": transform_coords_list(data["coords"], transformer_utm_to_geo), "marker": data["marker"], "color": data["color"]} for name, data in islas_menores_data_utm.items() }
    islas_menores_data_mercator = { name: {"coords": transform_coords_list(data_geo["coords"], transformer_geo_to_mercator), "marker": data_geo["marker"], "color": data_geo["color"]} for name, data_geo in islas_menores_data_geo.items() }
    print("Geometrías base del ANP cargadas y transformadas.")
except Exception as e:
    print(f"Error fatal durante la transformación de coordenadas base del ANP: {e}"); exit()

# --- FUNCIÓN PARA GENERAR REPORTE EN WORD (DOCX) ---
def generar_reporte_word(fig, observations_data, title, filename="reporte_inspeccion.docx"):
    print(f"DEBUG_WORD: Intentando generar reporte Word '{filename}' desde cero...")
    document = Document() 
    
    document.add_heading(title, level=1)
    document.add_paragraph() 

    temp_img_path = "temp_map_report.png"
    try:
        print(f"DEBUG_WORD: Guardando imagen temporal del mapa en '{temp_img_path}'...")
        fig.savefig(temp_img_path, dpi=300, bbox_inches='tight', pad_inches=0.1)
        print(f"DEBUG_WORD: Imagen temporal guardada: '{temp_img_path}' (Existe: {os.path.exists(temp_img_path)})")
    except Exception as e:
        print(f"ERROR_WORD: Falló al guardar la imagen temporal del mapa para el Word: {e}")
        return

    try:
        document.add_picture(temp_img_path, width=Inches(6.5)) 
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print(f"DEBUG_WORD: Imagen añadida al documento Word.")
    except Exception as e:
        print(f"ERROR_WORD: Falló al añadir la imagen al documento Word: {e}")
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)
        return

    document.add_page_break()

    document.add_heading('Resumen de Observaciones:', level=2)
    
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8') 
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Mexico.1252') 
        except locale.Error:
            print("ADVERTENCIA: No se pudo configurar el locale español. Los meses se mostrarán en inglés.")

    sorted_observations = sorted(observations_data, key=lambda x: x.get('timestamp_avistamiento', '') or x.get('timestamp', ''))

    for i, obs in enumerate(sorted_observations):
        timestamp_str = obs.get('timestamp_avistamiento') or obs.get('timestamp')
        try:
            if isinstance(timestamp_str, str):
                if 'T' in timestamp_str:
                    obs_dt_obj = datetime.datetime.fromisoformat(timestamp_str.replace('Z', '+00:00'))
                elif len(timestamp_str) == 19 and timestamp_str[10] == ' ':
                    obs_dt_obj = datetime.datetime.strptime(timestamp_str, '%Y-%m-%d %H:%M:%S')
                elif len(timestamp_str) == 16 and timestamp_str[10] == ' ':
                    obs_dt_obj = datetime.datetime.strptime(timestamp_str, '%Y-%m-%d %H:%M')
                else:
                    raise ValueError("Formato de timestamp desconocido.")
            else:
                raise ValueError("Timestamp no es una cadena.")
            
            date_str = obs_dt_obj.strftime('%d')
            month_str = obs_dt_obj.strftime('%B')
            year_str = obs_dt_obj.strftime('%Y')
            time_str = obs_dt_obj.strftime('%H:%M')
        except ValueError:
            date_str, month_str, year_str, time_str = "N/A", "N/A", "N/A", "N/A"
            print(f"ADVERTENCIA: No se pudo parsear el timestamp '{timestamp_str}'. Usando N/A.")

        lat_dd = obs.get('lat_dd') or obs.get('latitud_wgs84')
        lon_dd = obs.get('lon_dd') or obs.get('longitud_wgs84')
        lat_gmm = dd_to_gmm_str(lat_dd, True) if lat_dd is not None else "N/A"
        lon_gmm = dd_to_gmm_str(lon_dd, False) if lon_dd is not None else "N/A"

        vessel_name = obs.get('nombre_embarcacion', obs.get('name', 'N/A'))
        matricula = obs.get('matricula', 'N/A')
        nombre_patron = obs.get('nombre_patron', 'N/A')

        status_id = obs.get('estatus_categoria_id') or obs.get('status_category_id')
        status_desc = "Estatus Desconocido"
        if status_id == "outside_anp": status_desc = "Fuera del Polígono ANP"
        else:
            for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                if cat_info['id'] == status_id: status_desc = cat_info['desc']; break
        notes = obs.get('notas_adicionales', '')

        patron_report_text = ""
        if nombre_patron and nombre_patron.strip().lower() != 'n/a':
            patron_report_text = f" a cargo del C. \"{nombre_patron}\" quien se identificó como patron/capitan de citada embarcacion"

        report_paragraph = document.add_paragraph()
        runner = report_paragraph.add_run(
            f"{i+1}.- El día {date_str} de {month_str} de {year_str} a las {time_str} horas "
            f"en la situación geográfica {lat_gmm} y {lon_gmm} se detectó a la embarcación "
            f"de nombre \"{vessel_name}\" matrícula {matricula}{patron_report_text}, fue reportada la situación de "
            f"{status_desc} teniendo como nota: {notes}."
        )
        font = runner.font
        font.size = Pt(10)
    
    try:
        document.save(filename)
        print(f"DEBUG_WORD: Reporte Word '{filename}' generado exitosamente.")
    except Exception as e:
        print(f"ERROR_WORD: Falló al guardar el archivo Word final: {e}")
    finally:
        if os.path.exists(temp_img_path):
            print(f"DEBUG_WORD: Eliminando archivo temporal '{temp_img_path}'...")
            os.remove(temp_img_path)
            print(f"DEBUG_WORD: Archivo temporal eliminado.")

# --- FUNCIÓN PARA GRAFICAR HISTORIAL O MAPA DE SESIÓN (SOLO MUESTRA EL MAPA) ---
def graficar_mapa_observaciones(registros_o_puntos_sesion, titulo_mapa, es_historial_individual=False, ax_mapa_existente=None):
    if not registros_o_puntos_sesion and not es_historial_individual and ax_mapa_existente is None:
        return None, None
    elif not registros_o_puntos_sesion and es_historial_individual:
        print(f"No hay registros de historial para: {titulo_mapa}")
        return None, None

    if ax_mapa_existente is None:
        fig, ax = plt.subplots(figsize=(10 if es_historial_individual else 13, 8 if es_historial_individual else 11))
    else:
        fig = ax_mapa_existente.get_figure()
        ax = ax_mapa_existente

    ax.set_title(titulo_mapa)

    if ax_mapa_existente is None or not es_historial_individual:
        # CORRECCIÓN: Usar comillas dobles para toda la cadena 'label' para evitar el unterminated string literal
        if anp_maritime_boundary_coords_mercator: x_anp_m,y_anp_m = zip(*anp_maritime_boundary_coords_mercator); ax.plot(x_anp_m,y_anp_m,color="blue",linewidth=1.5 if es_historial_individual else 2.5,zorder=2,alpha=0.6,label="Límite ANP"); ax.fill(x_anp_m,y_anp_m,alpha=0.1,color="lightblue",zorder=1)
        if isla_maria_madre_coords_mercator and isla_maria_madre_polygon_geo.is_valid: x_imm_m,y_imm_m = zip(*isla_maria_madre_coords_mercator); ax.plot(x_imm_m,y_imm_m,color="darkgreen",linewidth=0.8,zorder=3,alpha=0.7,label="Isla María Madre"); ax.fill(x_imm_m,y_imm_m,color="lightgreen",alpha=0.5,zorder=2)
        if puerto_balleto_coords_mercator and puerto_balleto_polygon_geo.is_valid: x_pb_m,y_pb_m = zip(*puerto_balleto_coords_mercator); ax.plot(x_pb_m,y_pb_m,color="saddlebrown",linewidth=0.8,zorder=4,alpha=0.7,label="Puerto Balleto"); ax.fill(x_pb_m,y_pb_m,color="peru",alpha=0.6,zorder=3)
        for n_isla,d_isla_m in islas_menores_data_mercator.items():
            if d_isla_m["coords"] and all(isinstance(c,tuple)and len(c)==2 for c in d_isla_m["coords"]): p_coll_m_isla=MultiPoint(d_isla_m["coords"]); ax.plot([p.x for p in p_coll_m_isla.geoms],[p.y for p in p_coll_m_isla.geoms],marker=d_isla_m["marker"],color=d_isla_m["color"],linestyle='None',markersize=6,label=n_isla,zorder=5,alpha=0.8)

    legend_elements_types_used_on_this_map = {}
    
    if registros_o_puntos_sesion:
        for i, data_point in enumerate(registros_o_puntos_sesion):
            if es_historial_individual:
                lon_wgs84, lat_wgs84 = data_point['longitud_wgs84'], data_point['latitud_wgs84']
                x_m, y_m = transformer_geo_to_mercator.transform(lon_wgs84, lat_wgs84)
                v_type_id = data_point.get('tipo_embarcacion_id', DEFAULT_VESSEL_TYPE_INFO['id'])
                s_cat_id = data_point.get('estatus_categoria_id', "unknown_status")
                nombre_display = data_point.get('nombre_embarcacion', f"Emb. {data_point.get('matricula','N/A')}")
                matricula_display = data_point.get('matricula', 'N/A')
                timestamp_display = data_point.get('timestamp')
                notas_display = data_point.get('notas_adicionales', "")
                nombre_patron_display = data_point.get('nombre_patron', 'N/A')
                vessel_desc_display = VESSEL_TYPES.get(v_type_id, DEFAULT_VESSEL_TYPE_INFO)['desc']
                
                status_desc_display = "Estatus Desconocido"
                if s_cat_id == "outside_anp": status_desc_display = "Fuera del Polígono ANP"
                else:
                    for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                        if cat_info['id'] == s_cat_id: status_desc_display = cat_info['desc']; break
                lat_gmm_display, lon_gmm_display = dd_to_gmm_str(lat_wgs84,True), dd_to_gmm_str(lon_wgs84,False)
            else: # Es un punto de la sesión actual
                x_m,y_m=data_point["x_mercator"],data_point["y_mercator"]
                v_type_id=data_point["vessel_type_id"]
                s_cat_id=data_point["status_category_id"]
                nombre_display=data_point['name']
                matricula_display=data_point['matricula']
                timestamp_display=data_point['timestamp_avistamiento']
                notas_display = data_point.get('notas_adicionales', "")
                nombre_patron_display = data_point.get('nombre_patron', 'N/A')
                vessel_desc_display=data_point['vessel_desc']
                status_desc_display=data_point['status_desc']
                lat_gmm_display,lon_gmm_display=data_point['lat_gmm'],data_point['lon_gmm']

            # AHORA: DEFINIR marker_details y color para CADA data_point
            marker_details = VESSEL_TYPES.get(v_type_id, DEFAULT_VESSEL_TYPE_INFO)
            color = STATUS_COLORS.get("unknown_status") # Default color
            if s_cat_id == "outside_anp":
                color = STATUS_COLORS["outside_anp"]
            else:
                for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                    if cat_info['id'] == s_cat_id:
                        color = STATUS_COLORS[cat_info['color_key']]
                        break

            label_for_scatter = None
            if not es_historial_individual and marker_details['desc'] not in legend_elements_types_used_on_this_map:
                legend_elements_types_used_on_this_map[marker_details['desc']] = plt.Line2D([0],[0], marker=marker_details['marker_char'], color='w', label=marker_details['desc'], linestyle='None', markeredgecolor='black', markerfacecolor='dimgray', markersize=np.sqrt(marker_details['size_factor'])/1.5 )

            ax.scatter(x_m,y_m,marker=marker_details['marker_char'],color=color,
                       s=marker_details['size_factor']*(0.7 if es_historial_individual else 1.0),
                       edgecolors='black',linewidths=0.4,zorder=10 if es_historial_individual else 11,alpha=0.85,
                       label=label_for_scatter if label_for_scatter and not es_historial_individual else None)

            try: ts_obj=datetime.datetime.fromisoformat(timestamp_display); ts_fmt=ts_obj.strftime('%y-%m-%d %H:%M')
            except: ts_fmt=timestamp_display[:16] if timestamp_display else "N/A"

            # Incluir nombre del patrón en la anotación del mapa (si está presente)
            patron_map_text = f"C. {nombre_patron_display}" if nombre_patron_display and nombre_patron_display.strip().lower() != 'n/a' else "N/A"
            annot_text=(f"{nombre_display} ({matricula_display})\nPatrón: {patron_map_text}\n{ts_fmt} ({vessel_desc_display})\n"
                        f"{lat_gmm_display} | {lon_gmm_display}\nEstatus: {status_desc_display}\n"
                        f"Notas: {notas_display[:25]}{'...' if len(notas_display)>25 else ''}")
            ax.annotate(annot_text,(x_m,y_m),xytext=(0,(marker_details['size_factor']*(0.7 if es_historial_individual else 1.0))*0.05+7),textcoords='offset points',fontsize=5.5 if es_historial_individual else 6.5,ha='center',va='bottom',bbox=dict(boxstyle="round,pad=0.1",fc=color,alpha=0.7,ec='black',lw=0.4))

    ax.set_xlabel("X (Web Mercator)"); ax.set_ylabel("Y (Web Mercator)")
    try:
        print(f"DEBUG: Intentando cx.add_basemap para {'historial' if es_historial_individual else 'sesión'}...")
        cx.add_basemap(ax,crs=crs_mercator.to_string(),source=cx.providers.Esri.WorldImagery,zorder=0,alpha=0.9)
        print(f"DEBUG: cx.add_basemap ({'historial' if es_historial_individual else 'sesión'}) ejecutado.")
    except Exception as e_ctx:
        print(f"--- ERROR DETALLADO AL CARGAR MAPA BASE ({'HISTORIAL' if es_historial_individual else 'SESIÓN'}) ---"); import traceback; traceback.print_exc(); print(f"--- FIN TRACEBACK ({'HISTORIAL' if es_historial_individual else 'SESIÓN'}) ---"); print("Mapa base no cargado.")

    handles_fig, labels_fig = ax.get_legend_handles_labels();
    final_handles_display, final_labels_display, temp_legend_dict = [], [], {}

    for h,l in zip(handles_fig,labels_fig):
        if l not in temp_legend_dict and l is not None and l not in [info['desc'] for info in VESSEL_TYPES.values()]:
            temp_legend_dict[l]=h; final_handles_display.append(h); final_labels_display.append(l)

    final_handles_display.extend(legend_elements_types_used_on_this_map.values())
    final_labels_display.extend(legend_elements_types_used_on_this_map.keys())

    final_unique_dict = {}; clean_handles, clean_labels = [], []
    for h, l in zip(final_handles_display, final_labels_display):
        if l not in final_unique_dict: final_unique_dict[l] = h; clean_handles.append(h); clean_labels.append(l)
    final_handles_display, final_labels_display = clean_handles, final_labels_display

    if final_handles_display:
        leg1_title = "Leyenda Mapa" if es_historial_individual else "Leyenda Principal"
        leg1=ax.legend(final_handles_display,final_labels_display,fontsize='xx-small',loc='upper left',bbox_to_anchor=(1.02,1),borderaxespad=0.,title=leg1_title); ax.add_artist(leg1)

    status_legend_handles=[];
    all_s_disp_ordered = [
        ("blanco", "Paso Inocente"),
        ("verde", "Turístico Autorizado"),
        ("azul_marino", "Investigación Autorizada"),
        ("amarillo", "Inconsistencias Doc. / Nav."),
        ("anaranjado", "Infracción LGPAS (Pesca/Acuacultura)"),
        ("rojo", "Delito Ambiental / Otro"),
        ("outside_anp", "Fuera del Polígono ANP"),
        ("unknown_status", "Estatus Desconocido")
    ]
    for color_key, description in all_s_disp_ordered:
        status_legend_handles.append(plt.Line2D([0],[0],marker='s',color='w',label=description,markerfacecolor=STATUS_COLORS[color_key],markersize=7))

    if status_legend_handles:
        leg2_title="Semáforo Estatus";
        leg2=ax.legend(handles=status_legend_handles,fontsize='xx-small',loc='center left',bbox_to_anchor=(1.02,0.55 if not es_historial_individual else 0.65),borderaxespad=0.,title=leg2_title); ax.add_artist(leg2)

    if es_historial_individual:
        vessel_type_legend_hist = []
        for vt_info in VESSEL_TYPES.values():
            vessel_type_legend_hist.append(plt.Line2D([0],[0], marker=vt_info['marker_char'], color='w', label=vt_info['desc'], linestyle='None', markeredgecolor='black', markerfacecolor='dimgray', markersize=np.sqrt(vt_info['size_factor'])/1.5 ))
        if vessel_type_legend_hist:
            leg_hist_v_types=ax.legend(handles=vessel_type_legend_hist,fontsize='xx-small',loc='lower left',bbox_to_anchor=(1.02,0.25),borderaxespad=0.,title="Tipos Emb. (Definidos)"); ax.add_artist(leg_hist_v_types)

    if ax_mapa_existente is None:
        plt.subplots_adjust(left=0.06,right=0.70,bottom=0.05,top=0.92)

    return fig, ax

# --- INTERACCIÓN DE CLIC EN MAPA PRINCIPAL ---
current_click_annotation_main = None
current_click_marker_main = None
ax_main_map_ref = None

current_hover_annotation = None

def on_click_get_geo_coords_main_map(event):
    global current_click_annotation_main, current_click_marker_main, ax_main_map_ref
    if ax_main_map_ref is None or event.inaxes != ax_main_map_ref or event.xdata is None or event.ydata is None: return
    click_x_m,click_y_m=event.xdata,event.ydata; click_lon,click_lat=transformer_mercator_to_geo.transform(click_x_m,click_y_m)
    poly_bound_geo=anp_maritime_polygon_geo.exterior
    if not poly_bound_geo.is_empty:
        dist_on_line=poly_bound_geo.project(Point(click_lon,click_lat)); closest_pt_geo=poly_bound_geo.interpolate(dist_on_line)
        cp_lon,cp_lat=closest_pt_geo.x,closest_pt_geo.y; lat_gmm,lon_gmm=dd_to_gmm_str(cp_lat,True),dd_to_gmm_str(cp_lon,False)
        print(f"- Clic: M({click_x_m:.2f},{click_y_m:.2f}) G({click_lon:.6f},{click_lat:.6f}). Borde cercano: G({cp_lon:.6f},{cp_lat:.6f}) {lon_gmm},{lat_gmm}")
        # CORRECCIÓN: Usar comillas dobles en el f-string para evitar conflicto con las comillas simples internas
        current_click_annotation_main=ax_main_map_ref.annotate(f"{lat_gmm}\n{lon_gmm}",(cp_x_m,cp_y_m),xytext=(10,-25),textcoords='offset points',fontsize=8,bbox=dict(boxstyle="round,pad=0.3",fc="lightcyan",alpha=0.9,ec='k',lw=0.5),arrowprops=dict(arrowstyle="->",connectionstyle="arc3,rad=.2",color='k'))
        if current_click_annotation_main and current_click_annotation_main.get_visible():
             current_click_annotation_main.remove()
        if current_click_marker_main and current_click_marker_main in ax_main_map_ref.lines:
             current_click_marker_main.remove()

        cp_x_m,cp_y_m=transformer_geo_to_mercator.transform(cp_lon,cp_lat)
        current_click_marker_main,=ax_main_map_ref.plot(cp_x_m,cp_y_m,'mo',ms=7,mec='k',zorder=12)
        current_click_annotation_main=ax_main_map_ref.annotate(f"{lat_gmm}\n{lon_gmm}",(cp_x_m,cp_y_m),xytext=(10,-25),textcoords='offset points',fontsize=8,bbox=dict(boxstyle="round,pad=0.3",fc="lightcyan",alpha=0.9,ec='k',lw=0.5),arrowprops=dict(arrowstyle="->",connectionstyle="arc3,rad=.2",color='k'))
        ax_main_map_ref.figure.canvas.draw_idle()

def on_mouse_motion(event):
    global current_hover_annotation
    if ax_main_map_ref is None or event.inaxes != ax_main_map_ref or event.xdata is None or event.ydata is None:
        if current_hover_annotation:
            current_hover_annotation.remove()
            current_hover_annotation = None
            event.canvas.draw_idle()
        return

    x_mercator, y_mercator = event.xdata, event.ydata
    lon_dd, lat_dd = transformer_mercator_to_geo.transform(x_mercator, y_mercator)

    lat_gdm_str = dd_to_gmm_str(lat_dd, True)
    lon_gdm_str = dd_to_gmm_str(lon_dd, False)

    coords_text = f"Lat: {lat_gdm_str}\nLon: {lon_gdm_str}"

    if current_hover_annotation:
        current_hover_annotation.set_text(coords_text)
    else:
        current_hover_annotation = ax_main_map_ref.annotate(
            coords_text,
            xy=(1, 0),
            xycoords='axes fraction',
            xytext=(-5, 5),
            textcoords='offset points',
            fontsize=8,
            bbox=dict(boxstyle="round,pad=0.2", fc="white", alpha=0.8, ec="gray", lw=0.5),
            ha='right',
            va='bottom',
            zorder=20
        )
    event.canvas.draw_idle()


# --- BLOQUE PRINCIPAL (verificador11.py) ---
if __name__ == "__main__":
    # Asegurarse de que la columna 'nombre_patron' exista en la DB
    conn_check = conectar_db()
    cursor_check = conn_check.cursor()
    try:
        cursor_check.execute("ALTER TABLE observaciones_embarcaciones ADD COLUMN nombre_patron TEXT")
        conn_check.commit()
        print("Columna 'nombre_patron' añadida a la base de datos.")
    except sqlite3.OperationalError as e:
        if "duplicate column name: nombre_patron" in str(e):
            print("La columna 'nombre_patron' ya existe. No es necesaria la migración.")
        else:
            print(f"Error al verificar/añadir columna 'nombre_patron': {e}")
    finally:
        conn_check.close()
    
    inicializar_db(); user_points_data_session = []
    print("\n--- Verificación de Embarcaciones/Puntos ANP ---")
    while True:
        verificar_puntos_usuario = input("¿Deseas ingresar datos? (s/n): ").lower()
        if verificar_puntos_usuario in ['s','n']: break
        print("Respuesta inválida.")

    if verificar_puntos_usuario == 's':
        num_points_to_enter = 0
        while True:
            try:
                num_points_to_enter = int(input("¿Cuántas embarcaciones/puntos deseas ingresar (1-50)?: "))
                if 1 <= num_points_to_enter <= 50:
                    break
                else:
                    print("Número debe estar entre 1 y 50.")
            except ValueError:
                print("Entrada numérica inválida. Intente de nuevo.")

        for i in range(num_points_to_enter):
            print(f"\n--- Datos para Embarcación/Punto {i+1} de {num_points_to_enter} ---")
            matricula_input = "";
            while not matricula_input: matricula_input = input("Matrícula (ID único, obligatorio): ").strip().upper();
            nombre_embarcacion_input = input(f"Nombre embarcación para '{matricula_input}' (o etiqueta): ") or f"Emb. {matricula_input}"
            nombre_patron_input = input(f"Nombre del Patrón/Capitán para '{matricula_input}': ") or "N/A"

            # ALERTA: Buscar si el nombre de la embarcación o el patrón ya existen (excluyendo la matrícula actual si ya existía)
            conn_alerta = conectar_db()
            cursor_alerta = conn_alerta.cursor()
            cursor_alerta.execute("""
            SELECT * FROM observaciones_embarcaciones 
            WHERE (LOWER(nombre_embarcacion) = LOWER(?) OR LOWER(nombre_patron) = LOWER(?))
            AND matricula != ? 
            ORDER BY timestamp DESC
            """, (nombre_embarcacion_input, nombre_patron_input, matricula_input))
            coincidencias_generales = [dict(row) for row in cursor_alerta.fetchall()] 
            conn_alerta.close()

            if coincidencias_generales:
                print("\n¡ALERTA! Se encontraron registros previos con el mismo nombre de embarcación o patrón (en OTRA matrícula):")
                for reg_idx, reg_c in enumerate(coincidencias_generales[:min(3, len(coincidencias_generales))]):
                    try: ts_obj_c = datetime.datetime.fromisoformat(reg_c['timestamp']); ts_str_c = ts_obj_c.strftime('%Y-%m-%d %H:%M')
                    except: ts_str_c = reg_c['timestamp']
                    status_desc_c_short = reg_c['estatus_categoria_id'].replace('_',' ').title()
                    for cat_info_c_val in STATUS_CATEGORIES_INSIDE_ANP.values():
                        if cat_info_c_val['id'] == reg_c['estatus_categoria_id']: status_desc_c_short = cat_info_c_val['desc']; break
                    if reg_c['estatus_categoria_id'] == "outside_anp": status_desc_c_short = "Fuera del Polígono ANP"
                    print(f"  - {ts_str_c}: Matrícula '{reg_c['matricula']}', Nombre Emb. '{reg_c['nombre_embarcacion']}', Patrón '{reg_c.get('nombre_patron', 'N/A')}', Estatus '{status_desc_c_short}'")
                input("Presiona Enter para continuar con la observación actual...") 

            # AHORA SÍ: DESPUÉS DE LA ALERTA DE COINCIDENCIA, VERIFICAR EL HISTORIAL DE LA MATRÍCULA ACTUAL
            historial = buscar_historial_embarcacion(matricula_input) 
            if historial:
                print(f"\n¡ALERTA! '{matricula_input} ({nombre_embarcacion_input})' tiene {len(historial)} registro(s) previos.")
                for reg_idx, reg_h in enumerate(historial[:min(3, len(historial))]):
                    try: ts_obj_hist = datetime.datetime.fromisoformat(reg_h['timestamp']); ts_str_hist = ts_obj_hist.strftime('%Y-%m-%d %H:%M')
                    except: ts_str_hist = reg_h['timestamp']
                    status_desc_hist_short = reg_h['estatus_categoria_id'].replace('_',' ').title()
                    for cat_info_val_hist in STATUS_CATEGORIES_INSIDE_ANP.values():
                        if cat_info_val_hist['id'] == reg_h['estatus_categoria_id']: status_desc_hist_short = cat_info_val_hist['desc']; break
                    if reg_h['estatus_categoria_id'] == "outside_anp": status_desc_hist_short = "Fuera del Polígono ANP"

                    print(f"  - {ts_str_hist}: Estatus '{status_desc_hist_short}', en Lat {reg_h['latitud_wgs84']:.4f}, Lon {reg_h['longitud_wgs84']:.4f}, Patrón: {reg_h.get('nombre_patron', 'N/A')}")

                while True:
                    ver_hist = input("¿Ver mapa de historial completo? (s/n): ").lower()
                    if ver_hist in ['s','n']: break
                    print("Opción inválida.")
                if ver_hist == 's':
                    hist_fig, hist_ax = graficar_mapa_observaciones(historial, f"Historial: {matricula_input} ({nombre_embarcacion_input})", es_historial_individual=True)
                    if hist_fig:
                        observations_for_report_word_hist = historial 
                        
                        if observations_for_report_word_hist:
                            base_name = f"Historial_{matricula_input}_{nombre_embarcacion_input}".replace(' ', '_').replace(':', '').replace('/', '_').replace('(', '').replace(')', '')
                            report_word_filename = f"{base_name}_Reporte.docx"
                            
                            while True:
                                generar_word_choice = input(f"¿Deseas guardar un reporte en Word (DOCX) para el historial de '{matricula_input}'? (s/n): ").lower()
                                if generar_word_choice in ['s', 'n']:
                                    break
                                else:
                                    print("Respuesta inválida. Por favor, ingresa 's' o 'n'.")

                            if generar_word_choice == 's':
                                print(f"DEBUG: El usuario eligió generar el reporte Word para el historial. Llamando a generar_reporte_word...")
                                generar_reporte_word(hist_fig, observations_for_report_word_hist, f"Historial: {matricula_input} ({nombre_embarcacion_input})", report_word_filename)
                            else:
                                print(f"Generación de reporte Word para el historial de '{matricula_input}' omitida por el usuario.")
                        
                        hist_fig.canvas.mpl_connect('motion_notify_event', on_mouse_motion)
                        plt.show(block=True)


                if historial:
                    while True:
                        del_choice = input("¿Deseas eliminar algún registro del historial de esta embarcación? (s/n): ").lower()
                        if del_choice in ['s', 'n']: break
                        print("Opción inválida.")
                    if del_choice == 's':
                        print("\nRegistros en el historial (ID de base de datos):")
                        temp_historial_list_for_deletion = list(historial)
                        for idx_h, reg_h_del in enumerate(temp_historial_list_for_deletion):
                            try: ts_obj_del = datetime.datetime.fromisoformat(reg_h_del['timestamp']); ts_str_del = ts_obj_del.strftime('%Y-%m-%d %H:%M')
                            except: ts_str_del = reg_h_del['timestamp']
                            s_desc_del_short = reg_h_del['estatus_categoria_id'].replace('_',' ').title()
                            for cat_val_del in STATUS_CATEGORIES_INSIDE_ANP.values():
                                if cat_val_del['id'] == reg_h_del['estatus_categoria_id']: s_desc_del_short = cat_val_del['desc']; break
                            if reg_h_del['estatus_categoria_id'] == "outside_anp": s_desc_del_short = "Fuera del Polígono ANP"
                            print(f"  {idx_h + 1}. ID_DB: {reg_h_del['id']}, Fecha: {ts_str_del}, Estatus: {s_desc_del_short}, Patrón: {reg_h_del.get('nombre_patron', 'N/A')}")
                        try:
                            num_to_del_input = input("Ingresa el NÚMERO DE LISTA del registro a eliminar (0 para cancelar): ")
                            num_to_del = int(num_to_del_input)
                            if 0 < num_to_del <= len(temp_historial_list_for_deletion):
                                id_a_eliminar = temp_historial_list_for_deletion[num_to_del - 1]['id']
                                if eliminar_observacion_db(id_a_eliminar):
                                    print("Actualizando historial...")
                                    historial = buscar_historial_embarcacion(matricula_input)
                                    if input("¿Ver mapa de historial actualizado? (s/n): ").lower() == 's':
                                        hist_fig_upd, hist_ax_upd = graficar_mapa_observaciones(historial, f"Historial Actualizado: {matricula_input}", True)
                                        if hist_fig_upd:
                                            hist_fig_upd.canvas.mpl_connect('motion_notify_event', on_mouse_motion)
                                            plt.show(block=True)
                            elif num_to_del != 0: print("Número de registro inválido.")
                        except ValueError: print("Entrada inválida. Debe ser un número.")

            print(f"\n--- Ingresando OBSERVACIÓN ACTUAL para '{matricula_input}' ---")
            avist_ts_str=""; avist_dt_obj=None
            while not avist_dt_obj:
                avist_in_str=input(f"Fecha y Hora del AVISTAMIENTO (AAAA-MM-DD HH:MM, Enter para actual): ").strip()
                if not avist_in_str: avist_dt_obj=datetime.datetime.now(); avist_ts_str=avist_dt_obj.isoformat(timespec='seconds'); print(f"  Usando actual: {avist_dt_obj.strftime('%Y-%m-%d %H:%M')}"); break
                try: avist_dt_obj=datetime.datetime.strptime(avist_in_str,"%Y-%m-%d %H:%M"); avist_ts_str=avist_dt_obj.isoformat(timespec='seconds')
                except ValueError: print("  Formato incorrecto.")

            print("\n--- Tipo de Embarcación (Observación Actual) ---")
            for key, value in VESSEL_TYPES.items(): print(f"{key}. {value['desc']}")
            v_choice=0
            while v_choice not in VESSEL_TYPES:
                try: v_choice=int(input("Seleccione tipo: "));
                except ValueError: print("Inválido."); continue
                if v_choice not in VESSEL_TYPES: print("Opción inválida.")
            selected_vessel_details = VESSEL_TYPES[v_choice]

            print("\n--- Formato de Coordenadas Actuales ---");
            print("1. GMS (Grados, Minutos, Segundos)")
            print("2. DD (Grados Decimales)")
            print("3. UTM (Zona 13N)")
            print("4. GDM (Grados, Minutos Decimales)")
            cf_choice='';lon_dd,lat_dd=None,None
            while cf_choice not in['1','2','3','4']:
                cf_choice=input(f"Opción formato para '{matricula_input}' (1-4): ")
                if cf_choice not in ['1','2','3','4']: print("Inválido.")
            try:
                if cf_choice=='1':
                    print("\n--- Formato GMS ---")
                    lat_g=input("LatG:");lat_m=input("LatM:");lat_s=input("LatS:");lat_h=input("LatH(N/S):").upper()
                    lon_g=input("LonG:");lon_m=input("LonM:");lon_s=input("LonS:");lon_h=input("LonH(E/O/W):").upper()
                    lat_dd=gms_to_dd(lat_g,lat_m,lat_s,lat_h);lon_dd=gms_to_dd(lon_g,lon_m,lon_s,lon_h)
                elif cf_choice=='2':
                    print("\n--- Formato Grados Decimales ---")
                    lat_dd=float(input("Latitud DD (ej. 21.633989): "))
                    lon_dd=float(input("Longitud DD (ej. -106.537233): "))
                elif cf_choice=='3':
                    print("\n--- Formato UTM ---")
                    utm_x=float(input("UTMX:"));utm_y=float(input("UTMY:"));lon_dd,lat_dd=transformer_utm_to_geo.transform(utm_x,utm_y); print(f"UTM->Geo: Lon={lon_dd:.6f},Lat={lat_dd:.6f}")
                elif cf_choice=='4':
                    print("\n--- Formato Grados, Minutos Decimales (GDM) ---")
                    print("Ej. Lat: Grados=21, MinutosDecimales=18.739, Hemisferio=N")
                    lat_g_gdm = input("Latitud - Grados (ej. 21): ")
                    lat_m_gdm = input("Latitud - Minutos.Decimales (ej. 18.739): ")
                    lat_h_gdm = input("Latitud - Hemisferio (N/S): ").upper()
                    lat_dd = gdm_to_dd(lat_g_gdm, lat_m_gdm, lat_h_gdm)
                    print("Ej. Lon: Grados=106, MinutosDecimales=13.259, Hemisferio=W")
                    lon_g_gdm = input("Longitud - Grados (ej. 106): ")
                    lon_m_gdm = input("Longitud - Minutos.Decimales (ej. 13.259): ")
                    lon_h_gdm = input("Longitud - Hemisferio (E/W/O): ").upper()
                    lon_dd = gdm_to_dd(lon_g_gdm, lon_m_gdm, lon_h_gdm)
                    print(f"GDM->Geo: Lon={lon_dd:.6f},Lat={lat_dd:.6f}")

                if lon_dd is None or lat_dd is None: raise ValueError("Coordenadas no determinadas.")

                curr_pt_geo=Point(lon_dd,lat_dd); is_in=anp_maritime_polygon_geo.intersects(curr_pt_geo) if not anp_maritime_polygon_geo.is_empty else False
                # Aquí se determina el estatus de la observación actual
                curr_s_id_actual="outside_anp"; 
                curr_s_details={"id":"outside_anp","desc":"Fuera del Polígono ANP","color_key":"outside_anp"}
                if is_in:
                    print(f"\n'{matricula_input}' DENTRO ANP. Estatus:");
                    for k,v in STATUS_CATEGORIES_INSIDE_ANP.items(): print(f"{k}. {v['desc']}")
                    s_choice=0
                    while s_choice not in STATUS_CATEGORIES_INSIDE_ANP:
                        try:s_choice=int(input("Seleccione estatus: "))
                        except ValueError:print("Inválido.");continue
                        if s_choice not in STATUS_CATEGORIES_INSIDE_ANP:print("Opción inválida.")
                    curr_s_details=STATUS_CATEGORIES_INSIDE_ANP[s_choice];
                    curr_s_id_actual=curr_s_details['id']
                    if curr_s_id_actual == "paso_inocente": curr_s_details["color_key"] = "blanco"
                    elif curr_s_id_actual == "turistico_autorizado": curr_s_details["color_key"] = "verde"
                    elif curr_s_id_actual == "investigacion": curr_s_details["color_key"] = "azul_marino"
                    elif curr_s_id_actual == "doc_nav_issue": curr_s_details["color_key"] = "amarillo"
                    elif curr_s_id_actual == "pesca_lgpas_issue": curr_s_details["color_key"] = "anaranjado"
                    elif curr_s_id_actual == "delito": curr_s_details["color_key"] = "rojo"

                else:print(f"'{matricula_input}' FUERA ANP.")
                notas=input("Notas adicionales (opcional): ")

                # Guardar nombre del patrón en la base de datos
                agregar_observacion_db(matricula_input,nombre_embarcacion_input,avist_ts_str,lat_dd,lon_dd,selected_vessel_details['id'],curr_s_id_actual,notas, nombre_patron_input)

                user_x_m,user_y_m=transformer_geo_to_mercator.transform(lon_dd,lat_dd)

                user_points_data_session.append({
                    "name":nombre_embarcacion_input,"matricula":matricula_input,"timestamp_avistamiento":avist_ts_str,
                    "x_mercator":user_x_m,"y_mercator":user_y_m,
                    "lat_gmm":dd_to_gmm_str(lat_dd,True),"lon_gmm":dd_to_gmm_str(lon_dd,False),
                    "lat_dd":lat_dd, 
                    "lon_dd":lon_dd, 
                    "vessel_type_id":selected_vessel_details['id'],
                    "vessel_desc":selected_vessel_details['desc'],
                    "status_category_id":curr_s_id_actual,
                    "status_desc":curr_s_details['desc'],
                    "display_color":STATUS_COLORS.get(curr_s_details['color_key'],"lightgray"),
                    "display_marker_char":selected_vessel_details['marker_char'],
                    "display_marker_size_factor":selected_vessel_details['size_factor'],
                    "notas_adicionales": notas,
                    "nombre_patron": nombre_patron_input 
                })
                print(f"Observación actual para '{matricula_input}' registrada para mapa sesión.")
            except ValueError as ve: print(f"Error entrada para '{matricula_input}': {ve}. Punto omitido.")
            except Exception as e:
                print(f"Error procesando '{matricula_input}': {type(e).__name__} - {e}. Punto omitido.")

    # --- MAPA DE SESIÓN ACTUAL ---
    fig_main_session, ax_main_map_ref = plt.subplots(figsize=(13, 11))

    fig_main_session, ax_main_map_ref = graficar_mapa_observaciones(user_points_data_session,
                                "ANP Islas Marías - Observaciones de Sesión Actual",
                                es_historial_individual=False,
                                ax_mapa_existente=ax_main_map_ref)

    if fig_main_session:
        # Aquí se maneja la lógica de generar el reporte de Word para la sesión actual
        if user_points_data_session:
            observations_for_report_word_session = user_points_data_session 

            base_name = f"ANP_Islas_Marías_-_Observaciones_de_Sesión_Actual".replace(' ', '_').replace(':', '').replace('/', '_').replace('(', '').replace(')', '')
            report_word_filename = f"{base_name}_Reporte.docx"
            
            while True:
                generar_word_choice = input(f"¿Deseas guardar un reporte en Word (DOCX) para la sesión actual? (s/n): ").lower()
                if generar_word_choice in ['s', 'n']:
                    break
                else:
                    print("Respuesta inválida. Por favor, ingresa 's' o 'n'.")

            if generar_word_choice == 's':
                print(f"DEBUG: El usuario eligió generar el reporte Word para la sesión actual. Llamando a generar_reporte_word...")
                generar_reporte_word(fig_main_session, observations_for_report_word_session, "ANP Islas Marías - Observaciones de Sesión Actual", report_word_filename)
            else:
                print(f"Generación de reporte Word para la sesión actual omitida por el usuario.")

        # Conectar eventos de interacción al mapa principal
        fig_main_session.canvas.mpl_connect('button_press_event', on_click_get_geo_coords_main_map)
        fig_main_session.canvas.mpl_connect('motion_notify_event', on_mouse_motion)
        plt.show(block=False)
    else:
        print("No se ingresaron puntos válidos en esta sesión o el usuario no optó por ingresar. No se mostrará el mapa de sesión principal.")
        if fig_main_session: plt.close(fig_main_session)

    print("\nPrograma terminado.")
    input("\nPresiona Enter para cerrar esta consola...")