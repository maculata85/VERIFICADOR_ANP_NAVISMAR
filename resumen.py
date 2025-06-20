import matplotlib.pyplot as plt
from shapely.geometry import Point, Polygon, MultiPoint
from pyproj import Transformer, CRS
import numpy as np
import contextily as cx
import sqlite3
import datetime
import calendar 
import os
import sys
import locale 

# Importar para manejar documentos Word
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN DE LA BASE DE DATOS ---
DB_NAME = "inspecciones_anp.db"

def conectar_db():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    db_path = os.path.join(script_dir, DB_NAME)
    if not os.path.exists(db_path):
        print(f"ERROR: La base de datos '{DB_NAME}' no se encuentra en la ruta: {db_path}")
        print("Asegúrate de que el archivo exista y esté en la misma carpeta que el script,")
        print("o ejecuta primero el script de ingreso de datos para crearlo.")
        return None
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn

def inicializar_db(): # Aunque este script es de reportes, tenerla no daña.
    conn = conectar_db()
    if not conn: return # No se pudo conectar, ya se imprimió error.
    cursor = conn.cursor()
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS observaciones_embarcaciones (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        matricula TEXT NOT NULL,
        nombre_embarcacion TEXT,
        timestamp TEXT NOT NULL, 
        latitud_wgs84 REAL NOT NULL,
        longitud_wgs84 REAL NOT NULL,
        tipo_embarcacion_id TEXT, 
        estatus_categoria_id TEXT,
        notas_adicionales TEXT,
        nombre_patron TEXT
    )
    """)
    conn.commit()
    conn.close()
    # No imprimimos nada aquí para un script de solo lectura, a menos que se cree la tabla.

# --- TRANSFORMADORES DE COORDENADAS ---
crs_utm_anp = CRS("EPSG:32613")
crs_geo = CRS("EPSG:4326")
crs_mercator = CRS("EPSG:3857")
try:
    transformer_utm_to_geo = Transformer.from_crs(crs_utm_anp, crs_geo, always_xy=True)
    transformer_geo_to_mercator = Transformer.from_crs(crs_geo, crs_mercator, always_xy=True)
    transformer_mercator_to_geo = Transformer.from_crs(crs_mercator, crs_geo, always_xy=True) 
except Exception as e:
    print(f"Error crítico al crear los transformadores de coordenadas: {e}"); exit()

# --- FUNCIONES AUXILIARES ---
def transform_coords_list(coords_list, transformer):
    if not coords_list: return []
    if not all(isinstance(coord, tuple) and len(coord) == 2 for coord in coords_list):
        if all(isinstance(coord, Point) for coord in coords_list):
             x_coords = np.array([coord.x for coord in coords_list])
             y_coords = np.array([coord.y for coord in coords_list])
        else: return []
    else:
        x_coords = np.array([coord[0] for coord in coords_list])
        y_coords = np.array([coord[1] for coord in coords_list])
    if x_coords.size == 0: return []
    trans_x, trans_y = transformer.transform(x_coords, y_coords) 
    return list(zip(trans_x, trans_y))

def dd_to_gmm_str(dd_val, is_latitude):
    hemisphere = ('N' if dd_val >= 0 else 'S') if is_latitude else ('E' if dd_val >= 0 else 'W')
    abs_dd = abs(dd_val); degrees = int(abs_dd); minutes_decimal = (abs_dd - degrees) * 60
    return f"{degrees}°{minutes_decimal:.3f}' {hemisphere}"

# --- DEFINICIONES DE VISUALIZACIÓN ---
STATUS_CATEGORIES_INSIDE_ANP = { 
    1: {"id": "paso_inocente", "desc": "Paso Inocente", "color_key": "blanco"},
    2: {"id": "turistico_autorizado", "desc": "Turístico Autorizado", "color_key": "verde"},
    3: {"id": "investigacion", "desc": "Investigación Autorizada", "color_key": "azul_marino"},
    4: {"id": "doc_nav_issue", "desc": "Inconsistencias Doc. / Nav.", "color_key": "amarillo"},
    5: {"id": "pesca_lgpas_issue", "desc": "Infracción LGPAS (Pesca/Acuacultura)", "color_key": "anaranjado"},
    6: {"id": "delito", "desc": "Delito Ambiental / Otro", "color_key": "rojo"},
}
STATUS_COLORS = {
    "blanco": "white",
    "verde": "green",
    "azul_marino": "darkblue",
    "amarillo": "yellow",
    "anaranjado": "orange",
    "rojo": "red",
    "outside_anp": "deepskyblue",
    "unknown_status": "lightgray" 
}
VESSEL_TYPES = { 
    "panga": {"id": "panga", "desc": "Panga / Emb. Menor", "marker_char": "^", "size_factor": 100}, 
    "yate": {"id": "yate", "desc": "Yate / Emb. Mayor", "marker_char": "s", "size_factor": 150},   
    "otra": {"id": "otra", "desc": "Otra / No especificada", "marker_char": "o", "size_factor": 80} 
}

DEFAULT_VESSEL_TYPE_INFO = {"id": "default", "desc": "Desconocido", "marker_char": "o", "size_factor": 80}


# --- COORDENADAS UTM ORIGINALES DEL ANP (Completas) ---
anp_maritime_boundary_coords_utm = [
    (327983.720703, 2441179.856690), (406635.043518, 2359341.216920),
    (368280.494690, 2319018.426330), (287871.587708, 2401061.958310),
    (327983.720703, 2441179.856690)
]
isla_maria_madre_coords_utm = [
    (340647.669678, 2396167.727910), (340579.711304, 2396063.339720), (340611.461304, 2395997.193730),
    (340587.648682, 2395913.849670), (340595.586304, 2395900.620480), (340645.857117, 2395890.037290),
    (340611.461304, 2395736.578670), (340628.659302, 2395618.838680), (340570.450684, 2395366.161070),
    (340525.471497, 2395249.744320), (340471.231689, 2395225.931700), (340487.543884, 2395182.234130),
    (340368.570129, 2395208.520080), (340324.400085, 2395007.740110), (340479.769714, 2394974.235470),
    (340393.179688, 2394903.139530), (340373.335876, 2394827.733090), (340450.065125, 2394766.878720),
    (340397.148315, 2394584.315920), (340256.918884, 2394516.846920), (340233.106323, 2394463.930300),
    (340152.408325, 2394445.409300), (340148.439514, 2394399.107120), (340182.835510, 2394362.065490),
    (340312.481506, 2394327.669490), (340373.335876, 2394281.367310), (340362.752502, 2394200.669310),
    (340446.096313, 2394192.731690), (340432.867126, 2394129.231690), (340327.033691, 2393958.575070),
    (340280.731506, 2393929.470890), (340268.825073, 2393908.304080), (340129.918701, 2393850.095700),
    (340141.824890, 2393802.470520), (340083.937500, 2393740.018130), (340033.490112, 2393757.320130),
    (339996.430115, 2393660.810120), (339967.940125, 2393661.430110), (339604.670105, 2393789.380130),
    (339527.150085, 2393658.210080), (339518.860107, 2393626.080080), (339516.220093, 2393591.930110),
    (339584.590088, 2393563.660100), (339650.840088, 2393463.400090), (339887.670105, 2393439.740110),
    (340161.430115, 2393477.410100), (340297.929321, 2393484.969910), (340489.607100, 2393022.572000),
    (340412.379000, 2392975.619700), (340361.409100, 2392954.032300), (340321.785500, 2392931.381300),
    (340295.962100, 2392918.257900), (340290.458700, 2392910.214500), (340269.927000, 2392865.976100),
    (340255.322000, 2392819.409400), (340243.892000, 2392791.892600), (340230.980300, 2392769.244300),
    (340215.105300, 2392759.719200), (340187.800200, 2392760.777600), (340168.593600, 2392758.504800),
    (340127.318500, 2392710.879700), (340112.203000, 2392694.661800), (340098.476700, 2392672.512400),
    (340074.983900, 2392646.056700), (340039.476900, 2392595.735600), (340036.037100, 2392587.001500),
    (340042.598800, 2392550.806500), (340065.458800, 2392523.713100), (340072.867200, 2392501.064700),
    (340077.698000, 2392477.208100), (340080.063900, 2392465.716300), (340076.677200, 2392463.388000),
    (340055.933800, 2392472.489700), (340049.372100, 2392473.336300), (340037.307100, 2392463.599600),
    (340057.819400, 2392442.841800), (340062.006000, 2392433.844500), (340061.159400, 2392424.107800),
    (340061.177000, 2392414.711600), (340047.829000, 2392415.398000), (340032.986600, 2392421.967400),
    (340029.475400, 2392418.726200), (340031.380400, 2392404.121200), (340038.365500, 2392383.589500),
    (340007.825700, 2392369.514200), (340002.805400, 2392359.882800), (340003.440400, 2392353.321100),
    (340018.680400, 2392324.322700), (340029.119900, 2392301.251200), (340036.102300, 2392288.919400),
    (340046.766000, 2392275.041300), (340037.942100, 2392262.092600), (340036.989600, 2392259.870100),
    (340040.799600, 2392257.647600), (340063.310200, 2392241.197300), (340060.715200, 2392232.911600),
    (340028.624400, 2392200.293200), (340015.895800, 2392181.650300), (339983.620400, 2392158.470700),
    (339965.552000, 2392150.173600), (339954.756900, 2392150.491100), (339942.374400, 2392157.476100),
    (339938.405700, 2392156.682400), (339936.818200, 2392137.632300), (339919.831900, 2392104.771000),
    (339939.020600, 2392084.476200), (339937.134600, 2392080.204300), (339911.418100, 2392080.799700),
    (339912.211900, 2392075.402200), (339925.388100, 2392064.607200), (339940.568600, 2392056.879300),
    (339941.277200, 2392048.862500), (339909.033400, 2392023.385200), (339908.651200, 2392005.784300),
    (339922.730700, 2391995.347100), (339944.120700, 2391991.740800), (339965.710700, 2391962.848200),
    (339996.368800, 2391935.821500), (340012.285400, 2391933.116600), (340015.928400, 2391935.073400),
    (340036.299500, 2391932.656400), (340036.593700, 2391910.177000), (340038.724700, 2391895.429900),
    (340044.230000, 2391884.233600), (340067.388800, 2391852.406400), (340078.899700, 2391843.468000),
    (340079.534700, 2391831.244200), (340072.769000, 2391832.081600), (340056.039700, 2391848.865500),
    (340034.535600, 2391848.404200), (340020.955800, 2391857.279300), (339987.300800, 2391874.900500),
    (339979.065600, 2391870.153000), (339987.955600, 2391845.705500), (339999.610500, 2391840.840700),
    (339990.879300, 2391836.342700), (340004.227500, 2391814.308200), (340023.198200, 2391823.797900),
    (340009.929300, 2391845.074000), (340007.010100, 2391843.477700), (339992.698300, 2391852.199200),
    (339993.015800, 2391857.279300), (340016.543900, 2391847.345900), (340033.900600, 2391843.747600),
    (340047.024000, 2391843.959200), (340057.131600, 2391841.214800), (340066.896700, 2391823.734500),
    (340075.493200, 2391821.945800), (340085.018200, 2391825.650000), (340089.102700, 2391828.828400),
    (340087.491400, 2391843.791800), (340084.344500, 2391845.818100), (340072.472200, 2391853.463100),
    (340061.777700, 2391871.174100), (340053.269200, 2391890.081400), (340041.236600, 2391908.580700),
    (340042.825000, 2391925.748300), (340041.017300, 2391938.588400), (340035.482100, 2391941.921100),
    (340029.687100, 2391943.956900), (340027.147100, 2391939.988200), (339995.147900, 2391941.743200),
    (339983.332000, 2391953.481900), (339968.885700, 2391965.229500), (339949.412100, 2391992.424300),
    (339943.803200, 2391996.027000), (339925.229400, 2391998.884500), (339912.211900, 2392005.869600),
    (339912.826100, 2392019.582900), (339943.485700, 2392040.953400), (339946.025700, 2392057.622200),
    (339944.279400, 2392060.638400), (339926.499400, 2392067.940900), (339919.196900, 2392074.132200),
    (339930.150700, 2392074.132200), (339938.246900, 2392077.783400), (339943.803200, 2392078.418400),
    (339945.708200, 2392084.609700), (339933.711500, 2392095.572300), (339925.388100, 2392107.311000),
    (339938.324900, 2392126.729100), (339939.389800, 2392129.708000), (339942.691900, 2392151.284800),
    (339960.630700, 2392144.299800), (339985.810800, 2392154.601400), (340018.796800, 2392178.369800),
    (340039.329100, 2392206.108200), (340055.922600, 2392219.451200), (340069.641900, 2392235.258700),
    (340068.263400, 2392242.883800), (340065.312800, 2392247.591800), (340043.339600, 2392260.187600),
    (340052.809200, 2392270.746200), (340052.547100, 2392275.745100), (340038.234200, 2392299.114700),
    (340021.643800, 2392326.439400), (340006.455700, 2392356.684200), (340013.690700, 2392363.109200),
    (340022.854200, 2392371.247000), (340043.183000, 2392375.030800), (340044.924900, 2392383.375200),
    (340036.672100, 2392404.544500), (340035.402100, 2392415.551200), (340047.890500, 2392410.259500),
    (340063.553800, 2392412.164500), (340066.305500, 2392431.637900), (340065.791800, 2392439.425500),
    (340042.598800, 2392463.388000), (340052.970500, 2392469.103000), (340073.076600, 2392460.422000),
    (340080.910500, 2392458.943000), (340085.910400, 2392465.840500), (340078.111300, 2392504.006200),
    (340069.903800, 2392526.464800), (340047.467100, 2392553.346500), (340042.138000, 2392575.950800),
    (340041.185500, 2392588.015800), (340041.961600, 2392592.078900), (340064.797200, 2392622.610600),
    (340077.698000, 2392642.625900), (340095.725000, 2392659.389000), (340115.006300, 2392686.786200),
    (340129.964400, 2392707.175500), (340164.375700, 2392746.607400), (340171.995700, 2392751.369900),
    (340180.830600, 2392754.932900), (340216.586900, 2392753.369200), (340233.943600, 2392764.375900),
    (340249.395300, 2392789.987600), (340259.784700, 2392818.362500), (340275.342200, 2392865.828800),
    (340296.456000, 2392909.485200), (340323.705900, 2392923.935900), (340362.813600, 2392949.014000),
    (340418.058700, 2392971.397800), (340445.046300, 2392986.637800), (340487.116700, 2393014.422800),
    (340491.723900, 2393017.465600), (340554.575684, 2392865.843690), (340735.815674, 2392689.895510),
    (340799.315918, 2392570.832700), (340914.409729, 2392618.457700), (340956.743286, 2392580.093080),
    (341315.254272, 2392087.967100), (341263.660522, 2392045.633730), (341321.868896, 2391935.831480),
    (341380.077271, 2392013.883730), (341389.337891, 2392037.696290), (341409.181702, 2392068.123290),
    (341415.997925, 2392074.462520), (341423.062317, 2392059.846680), (341452.431091, 2392036.431090),
    (341517.121887, 2392039.209110), (341533.790710, 2392066.196720), (341543.315674, 2392033.652890),
    (341538.156311, 2391976.502690), (341546.093689, 2391966.183900), (341556.015686, 2391974.121520),
    (341568.715698, 2392018.571470), (341602.450073, 2392075.324890), (341612.220276, 2392084.024110),
    (341634.077881, 2392050.925480), (341727.563110, 2392140.209110), 
    (343552.698730, 2389885.991700), (343507.912109, 2389799.561280), (343518.495483, 2389691.081910), 
    (343379.588928, 2389615.675480), (343288.307495, 2389620.967100), (343190.411499, 2389757.227720), 
    (343072.671509, 2389714.894290), (343126.911316, 2389597.154480), (342863.650330, 2389430.466670), 
    (342888.785889, 2389097.091130), (342875.556702, 2389093.122310), (342862.327271, 2388978.028320), 
    (342916.567078, 2388898.653080), (342921.858704, 2388836.475890), (343022.400696, 2388661.850520), 
    (343048.859131, 2388582.475520), (343100.876892, 2388586.024900), (343201.853088, 2388431.768920), 
    (340521.691101, 2384841.964480), (341115.563904, 2386033.625120), (339699.978699, 2386665.757080), 
    (339141.029114, 2385565.767700), (339139.558472, 2385562.847720), (333486.840088, 2388086.395690), 
    (333487.039673, 2388086.762510), (333846.045471, 2388718.355900), (333589.775696, 2388845.270320), 
    (333089.439514, 2388318.086910), (333088.560913, 2388317.270080), (331178.738525, 2391778.943910), 
    (331506.668701, 2391960.168090), (331401.720276, 2392177.387330), (331079.552490, 2392206.675110), 
    (330913.667908, 2392089.244320),
    (340647.669678, 2396167.727910) 
]
puerto_balleto_coords_utm = [
    (340647.669623, 2396167.727920), (340871.403535, 2393096.078400), (340981.527284, 2393138.588180),
    (340974.280217, 2393158.040830), (340982.290133, 2393161.473650), (340989.537199, 2393142.402420),
    (341020.051165, 2393153.463730), (341042.173790, 2393168.720720), (341044.080913, 2393179.019180),
    (341005.938456, 2393236.232860), (341015.855495, 2393240.809960), (341066.966387, 2393165.669320),
    (341057.812197, 2393158.803680), (341046.369460, 2393161.855070), (341025.009684, 2393146.598090),
    (340992.588596, 2393134.392510), (340999.835663, 2393115.321280), (340992.970021, 2393113.795580),
    (340984.197256, 2393132.485380), (340873.875105, 2393088.869650), (341727.563090, 2392140.209000),
    (341634.077853, 2392050.925350), (341612.220175, 2392084.024120), (341602.450143, 2392075.324770),
    (341568.715700, 2392018.571530), (341556.015675, 2391974.121450), (341546.093780, 2391966.183930),
    (341538.156264, 2391976.502700), (341543.315650, 2392033.652810), (341533.790631, 2392066.196630),
    (341517.121847, 2392039.209080), (341452.431093, 2392036.430940), (341423.062284, 2392059.846620),
    (341415.997933, 2392074.462520), (341409.181570, 2392068.123300), (341389.337780, 2392037.696150),
    (341380.077345, 2392013.883610), (341321.868895, 2391935.831370), (341263.660445, 2392045.633670),
    (341315.254298, 2392087.967090), (340956.743165, 2392580.093070), (340914.409747, 2392618.457730),
    (340799.315766, 2392570.832640), (340735.815639, 2392689.895380), (340554.575694, 2392865.843640),
    (340297.929347, 2393484.969880), (340161.430002, 2393477.410000), (339887.670000, 2393439.740000),
    (339650.840002, 2393463.400000), (339584.590002, 2393563.660000), (339516.219999, 2393591.930000),
    (339518.860003, 2393626.080000), (339527.150000, 2393658.210000), (339604.670000, 2393789.380000),
    (339967.940001, 2393661.430000), (339996.430002, 2393660.810000), (340033.490000, 2393757.320000),
    (340083.937515, 2393740.018110), (340141.824868, 2393802.470520), (340129.918594, 2393850.095610),
    (340268.825122, 2393908.304060), (340280.731396, 2393929.470770), (340327.033572, 2393958.575000),
    (340432.867117, 2394129.231590), (340446.096310, 2394192.731710), (340362.752393, 2394200.669230),
    (340373.335748, 2394281.367310), (340312.481459, 2394327.669480), (340182.835367, 2394362.065390),
    (340148.439465, 2394399.107130), (340152.408223, 2394445.409300), (340233.106301, 2394463.930170),
    (340256.918848, 2394516.846950), (340397.148295, 2394584.315830), (340450.065068, 2394766.878700),
    (340373.335748, 2394827.732980), (340393.179538, 2394903.139390), (340479.769674, 2394974.235460),
    (340324.400000, 2395007.740000), (340368.570002, 2395208.520000), (340487.543742, 2395182.234030),
    (340471.231777, 2395225.931700), (340525.471469, 2395249.744250), (340570.450725, 2395366.161140),
    (340628.659175, 2395618.838730), (340611.461224, 2395736.578550), (340645.857126, 2395890.037190),
    (340595.586192, 2395900.620550), (340587.648676, 2395913.849740), (340611.461224, 2395997.193730),
    (340579.711161, 2396063.339620),
    (340647.669623, 2396167.727920)
]
islas_menores_data_utm = {
    "Isla San Juanito (V)": {"coords": [(328030.906100, 2404586.986800), (327633.883600, 2404310.100800)], "marker": "o", "color": "darkviolet"},
    "Islote El Morro (V)": {"coords": [(323879.578800, 2405116.471600), (323867.899400, 2405109.464000)], "marker": "s", "color": "firebrick"},
    "Isla María Magdalena (V)": {"coords": [(350862.351900, 2377854.999000), (351247.036700, 2377788.188800)], "marker": "^", "color": "olive"},
    "Isla María Cleofas (V)": {"coords": [
        (369429.510100, 2359596.764100), (369848.633600, 2359593.460300), (369855.558600, 2359590.690300),
        (369866.088000, 2359593.322700), (369908.061600, 2359592.991800), (372259.125600, 2359574.458900),
        (372486.810100, 2359572.664100)
    ], "marker": "P", "color": "teal"},
    "Islote La Mona 1 (V)": {"coords": [(366477.352800, 2358421.914700), (366325.153200, 2358358.393600)], "marker": "*", "color": "darkorange"},
    "Islote La Mona 2 (V)": {"coords": [(367454.190600, 2356053.133500), (367435.503600, 2356048.461700)], "marker": "X", "color": "darkmagenta"},
    "Islote La Mona 3 (V)": {"coords": [(368306.785600, 2355808.450400), (368277.587200, 2355808.450400)], "marker": "D", "color": "navy"},
}

# --- TRANSFORMACIONES GLOBALES INICIALES ---
print("Cargando y transformando geometrías base del ANP...")
try:
    anp_maritime_boundary_coords_geo = transform_coords_list(anp_maritime_boundary_coords_utm, transformer_utm_to_geo)
    anp_maritime_polygon_geo = Polygon(anp_maritime_boundary_coords_geo) if anp_maritime_boundary_coords_geo else Polygon()
    anp_maritime_boundary_coords_mercator = transform_coords_list(anp_maritime_boundary_coords_geo, transformer_geo_to_mercator)

    isla_maria_madre_coords_geo = transform_coords_list(isla_maria_madre_coords_utm, transformer_utm_to_geo)
    isla_maria_madre_polygon_geo = Polygon(isla_maria_madre_coords_geo) if isla_maria_madre_coords_geo else Polygon()
    isla_maria_madre_coords_mercator = transform_coords_list(isla_maria_madre_coords_geo, transformer_geo_to_mercator)

    puerto_balleto_coords_geo = transform_coords_list(puerto_balleto_coords_utm, transformer_utm_to_geo)
    puerto_balleto_polygon_geo = Polygon(puerto_balleto_coords_geo) if puerto_balleto_coords_geo else Polygon()
    puerto_balleto_coords_mercator = transform_coords_list(puerto_balleto_coords_geo, transformer_geo_to_mercator)
    
    islas_menores_data_geo = { name: {"coords": transform_coords_list(data["coords"], transformer_utm_to_geo), "marker": data["marker"], "color": data["color"]} for name, data in islas_menores_data_utm.items() }
    islas_menores_data_mercator = { name: {"coords": transform_coords_list(data_geo["coords"], transformer_geo_to_mercator), "marker": data_geo["marker"], "color": data_geo["color"]} for name, data_geo in islas_menores_data_geo.items() }
    print("Geometrías base del ANP cargadas y transformadas.")
except Exception as e:
    print(f"Error fatal durante la transformación de coordenadas base del ANP: {e}"); exit()


# --- FUNCIÓN PARA GENERAR REPORTE EN WORD (DOCX) ---
def generar_reporte_word(fig, observations_data, title, filename="reporte_inspeccion.docx"):
    print(f"DEBUG_WORD: Intentando generar reporte Word '{filename}' desde cero...")
    document = Document() 
    
    document.add_heading(title, level=1)
    document.add_paragraph() 

    temp_img_path = "temp_map_report.png"
    try:
        print(f"DEBUG_WORD: Guardando imagen temporal del mapa en '{temp_img_path}'...")
        fig.savefig(temp_img_path, dpi=300, bbox_inches='tight', pad_inches=0.1)
        print(f"DEBUG_WORD: Imagen temporal guardada: '{temp_img_path}' (Existe: {os.path.exists(temp_img_path)})")
    except Exception as e:
        print(f"ERROR_WORD: Falló al guardar la imagen temporal del mapa para el Word: {e}")
        return

    try:
        document.add_picture(temp_img_path, width=Inches(6.5)) 
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print(f"DEBUG_WORD: Imagen añadida al documento Word.")
    except Exception as e:
        print(f"ERROR_WORD: Falló al añadir la imagen al documento Word: {e}")
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)
        return

    document.add_page_break()

    document.add_heading('Resumen de Observaciones:', level=2)
    
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8') 
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Mexico.1252') 
        except locale.Error:
            print("ADVERTENCIA: No se pudo configurar el locale español. Los meses se mostrarán en inglés.")

    sorted_observations = sorted(observations_data, key=lambda x: x.get('timestamp_avistamiento', '') or x.get('timestamp', ''))

    for i, obs in enumerate(sorted_observations):
        timestamp_str = obs.get('timestamp_avistamiento') or obs.get('timestamp')
        try:
            if isinstance(timestamp_str, str):
                if 'T' in timestamp_str:
                    obs_dt_obj = datetime.datetime.fromisoformat(timestamp_str.replace('Z', '+00:00'))
                elif len(timestamp_str) == 19 and timestamp_str[10] == ' ':
                    obs_dt_obj = datetime.datetime.strptime(timestamp_str, '%Y-%m-%d %H:%M:%S')
                elif len(timestamp_str) == 16 and timestamp_str[10] == ' ':
                    obs_dt_obj = datetime.datetime.strptime(timestamp_str, '%Y-%m-%d %H:%M')
                else:
                    raise ValueError("Formato de timestamp desconocido.")
            else:
                raise ValueError("Timestamp no es una cadena.")
            
            date_str = obs_dt_obj.strftime('%d')
            month_str = obs_dt_obj.strftime('%B')
            year_str = obs_dt_obj.strftime('%Y')
            time_str = obs_dt_obj.strftime('%H:%M')
        except ValueError:
            date_str, month_str, year_str, time_str = "N/A", "N/A", "N/A", "N/A"
            print(f"ADVERTENCIA: No se pudo parsear el timestamp '{timestamp_str}'. Usando N/A.")

        lat_dd = obs.get('lat_dd') or obs.get('latitud_wgs84')
        lon_dd = obs.get('lon_dd') or obs.get('longitud_wgs84')
        lat_gmm = dd_to_gmm_str(lat_dd, True) if lat_dd is not None else "N/A"
        lon_gmm = dd_to_gmm_str(lon_dd, False) if lon_dd is not None else "N/A"

        vessel_name = obs.get('nombre_embarcacion', obs.get('name', 'N/A'))
        matricula = obs.get('matricula', 'N/A')
        nombre_patron = obs.get('nombre_patron', 'N/A')

        status_id = obs.get('estatus_categoria_id') or obs.get('status_category_id')
        status_desc = "Estatus Desconocido"
        if status_id == "outside_anp": status_desc = "Fuera del Polígono ANP"
        else:
            for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                if cat_info['id'] == status_id: status_desc = cat_info['desc']; break
        notes = obs.get('notas_adicionales', '')

        patron_report_text = ""
        if nombre_patron and nombre_patron.strip().lower() != 'n/a':
            patron_report_text = f" a cargo del C. \"{nombre_patron}\" quien se identificó como patron/capitan de citada embarcacion"

        report_paragraph = document.add_paragraph()
        runner = report_paragraph.add_run(
            f"{i+1}.- El día {date_str} de {month_str} de {year_str} a las {time_str} horas "
            f"en la situación geográfica {lat_gmm} y {lon_gmm} se detectó a la embarcación "
            f"de nombre \"{vessel_name}\" matrícula {matricula}{patron_report_text}, fue reportada la situación de "
            f"{status_desc} teniendo como nota: {notes}."
        )
        font = runner.font
        font.size = Pt(10)
    
    try:
        document.save(filename)
        print(f"DEBUG_WORD: Reporte Word '{filename}' generado exitosamente.")
    except Exception as e:
        print(f"ERROR_WORD: Falló al guardar el archivo Word final: {e}")
    finally:
        if os.path.exists(temp_img_path):
            print(f"DEBUG_WORD: Eliminando archivo temporal '{temp_img_path}'...")
            os.remove(temp_img_path)
            print(f"DEBUG_WORD: Archivo temporal eliminado.")

# --- FUNCIÓN PARA GRAFICAR MAPA DE INSPECCIONES (RESUMEN O HISTORIAL) ---
def graficar_mapa_inspecciones(registros, titulo_mapa, es_historial_individual=False):
    if not registros:
        print(f"No hay registros para graficar para: {titulo_mapa}")
        return None, None # Devuelve None si no hay registros para que el flujo principal sepa

    fig, ax = plt.subplots(figsize=(15, 12))
    ax.set_title(titulo_mapa)

    # Dibujar polígonos base del ANP
    if anp_maritime_boundary_coords_mercator:
        x_anp_m, y_anp_m = zip(*anp_maritime_boundary_coords_mercator)
        ax.plot(x_anp_m, y_anp_m, color="blue", linewidth=1.5, zorder=2, alpha=0.6, label="Límite ANP")
        ax.fill(x_anp_m, y_anp_m, alpha=0.10, color="lightblue", zorder=1)
    
    if isla_maria_madre_coords_mercator and isla_maria_madre_polygon_geo.is_valid:
        x_imm_m, y_imm_m = zip(*isla_maria_madre_coords_mercator)
        ax.plot(x_imm_m, y_imm_m, color="darkgreen", linewidth=0.8, zorder=3, alpha=0.7, label="Isla María Madre")
        ax.fill(x_imm_m, y_imm_m, color="lightgreen", alpha=0.5, zorder=2)

    if puerto_balleto_coords_mercator and puerto_balleto_polygon_geo.is_valid:
        x_pb_m, y_pb_m = zip(*puerto_balleto_coords_mercator)
        ax.plot(x_pb_m, y_pb_m, color="saddlebrown", linewidth=0.8, zorder=4, alpha=0.7, label="Puerto Balleto")
        ax.fill(x_pb_m, y_pb_m, color="peru", alpha=0.6, zorder=3)

    for nombre_isla, data_mercator in islas_menores_data_mercator.items():
        if data_mercator["coords"]:
            points_collection_m = MultiPoint(data_mercator["coords"])
            ax.plot([p.x for p in points_collection_m.geoms], [p.y for p in points_collection_m.geoms],
                    marker=data_mercator["marker"], color=data_mercator["color"], linestyle='None',
                    markersize=6, label=nombre_isla, zorder=5, alpha=0.8)
    
    lats_wgs84 = [r['latitud_wgs84'] for r in registros]
    lons_wgs84 = [r['longitud_wgs84'] for r in registros]
    
    coords_mercator = []
    if lats_wgs84 and lons_wgs84:
        coords_transformed = transformer_geo_to_mercator.transform(lons_wgs84, lats_wgs84)
        if coords_transformed and len(coords_transformed) == 2:
             x_mercator_all, y_mercator_all = coords_transformed
             coords_mercator = list(zip(x_mercator_all, y_mercator_all))

    legend_elements_types_used = {} # Para leyenda de tipos usados en este mapa
    
    for i, registro in enumerate(registros):
        if i < len(coords_mercator): # Asegurarse de que el punto existe en coords_mercator
            x_m, y_m = coords_mercator[i]
            
            v_type_id = registro['tipo_embarcacion_id']
            s_cat_id = registro['estatus_categoria_id']
            
            marker_details = VESSEL_TYPES.get(v_type_id, DEFAULT_VESSEL_TYPE_INFO)
            
            color = STATUS_COLORS.get("unknown_status") # Default color si no se encuentra
            if s_cat_id == "outside_anp":
                color = STATUS_COLORS["outside_anp"]
            else:
                for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                    if cat_info['id'] == s_cat_id:
                        color = STATUS_COLORS[cat_info['color_key']]
                        break
            
            # Añadir a la leyenda de tipos si es la primera vez que se ve este tipo
            if marker_details['desc'] not in legend_elements_types_used:
                legend_elements_types_used[marker_details['desc']] = plt.Line2D([0],[0], marker=marker_details['marker_char'], color='w', label=marker_details['desc'], linestyle='None', markeredgecolor='black', markerfacecolor='dimgray', markersize=7)

            ax.scatter(x_m, y_m, 
                       marker=marker_details['marker_char'], color=color, 
                       s=marker_details['size_factor'] * (0.7 if es_historial_individual else 0.5), 
                       edgecolors='black', linewidths=0.4, zorder=10, alpha=0.75)
            
            # Formatear la anotación para el mapa
            try: ts_obj = datetime.datetime.fromisoformat(registro['timestamp']); ts_fmt = ts_obj.strftime('%y-%m-%d %H:%M')
            except: ts_fmt = registro['timestamp'][:16] if registro['timestamp'] else "N/A" # Fallback si hay problemas con formato completo
            
            status_desc_display = "Estatus Desconocido"
            found_status_desc = False
            for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                if cat_info['id'] == s_cat_id: status_desc_display = cat_info['desc']; found_status_desc = True; break
            if not found_status_desc and s_cat_id == "outside_anp": status_desc_display = "Fuera del Polígono ANP"
            
            # No anotar si hay demasiados puntos en resumen general para evitar saturación
            if es_historial_individual or len(registros) < 15: # Mostrar anotaciones solo en historial o pocos puntos
                patron_map_text = f"C. {registro.get('nombre_patron', 'N/A')}" if registro.get('nombre_patron') and registro['nombre_patron'].strip().lower() != 'n/a' else "N/A"
                annot_text = (f"{registro['matricula']}\nPatrón: {patron_map_text}\n{ts_fmt}\n{status_desc_display}")
                ax.annotate(annot_text, (x_m, y_m),
                            xytext=(0, marker_details['size_factor'] * (0.7 if es_historial_individual else 0.5) * 0.05 + 7), 
                            textcoords='offset points', fontsize=4 if not es_historial_individual else 5.5, 
                            ha='center', va='bottom',
                            bbox=dict(boxstyle="round,pad=0.1", fc=color, alpha=0.6, ec='none'))
    
    ax.set_xlabel("X (Web Mercator)"); ax.set_ylabel("Y (Web Mercator)")
    try: cx.add_basemap(ax, crs=crs_mercator.to_string(), source=cx.providers.Esri.WorldImagery, zorder=0, alpha=0.9)
    except Exception as e_ctx: print(f"No se pudo cargar mapa base para resumen: {e_ctx}")

    # --- LEYENDAS ---
    handles_fig, labels_fig = ax.get_legend_handles_labels(); 
    
    # Leyenda para elementos del mapa (ANP, islas)
    map_legend_handles_display = []
    map_legend_labels_display = []
    seen_labels = set()

    for h,l in zip(handles_fig, labels_fig):
        if l not in seen_labels:
            map_legend_handles_display.append(h)
            map_legend_labels_display.append(l)
            seen_labels.add(l)

    if map_legend_handles_display:
        leg1 = ax.legend(handles=map_legend_handles_display, labels=map_legend_labels_display, 
                         fontsize='xx-small', loc='upper left', bbox_to_anchor=(1.02, 1), 
                         borderaxespad=0., title="Elementos del Mapa")
        ax.add_artist(leg1)

    # Leyenda para tipos de embarcación (solo los que aparecen en los datos filtrados)
    if legend_elements_types_used:
        leg_vessel_types = ax.legend(handles=list(legend_elements_types_used.values()), 
                                     fontsize='xx-small', loc='center left', 
                                     bbox_to_anchor=(1.02, 0.65), 
                                     borderaxespad=0., title="Tipos Embarcación")
        ax.add_artist(leg_vessel_types)

    # Leyenda para el Semáforo de Estatus
    status_legend_handles = []
    all_status_display_ordered = [
        ("blanco", "Paso Inocente"),
        ("verde", "Turístico Autorizado"),
        ("azul_marino", "Investigación Autorizada"),
        ("amarillo", "Inconsistencias Doc. / Nav."),
        ("anaranjado", "Infracción LGPAS (Pesca/Acuacultura)"),
        ("rojo", "Delito Ambiental / Otro"),
        ("outside_anp", "Fuera del Polígono ANP"),
        ("unknown_status", "Estatus Desconocido")
    ]
    for color_key, description in all_status_display_ordered:
        status_legend_handles.append(plt.Line2D([0],[0], marker='s', color='w', label=description, 
                                                 markerfacecolor=STATUS_COLORS[color_key], markersize=7))
    if status_legend_handles:
        leg_status = ax.legend(handles=status_legend_handles, fontsize='xx-small', loc='lower left', 
                           bbox_to_anchor=(1.02, 0.05), borderaxespad=0., title="Semaforo Estatus")
        ax.add_artist(leg_status)
    
    plt.subplots_adjust(left=0.06, right=0.70, bottom=0.05, top=0.92) # Ajustar para leyendas externas
    return fig, ax # Retornar la figura y el eje para que el flujo principal pueda cerrarla/manejarla.

# --- BLOQUE PRINCIPAL DEL SCRIPT DE REPORTES ---
if __name__ == "__main__":
    # Configurar el locale al inicio del script para que todos los format_time lo usen
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Mexico.1252')
        except locale.Error:
            print("ADVERTENCIA: No se pudo configurar el locale español del sistema. Los nombres de meses se mostrarán en inglés.")

    inicializar_db() 
    
    print("--- Generador de Mapas Resumen de Inspecciones ANP ---")
    print("Seleccione el periodo para el resumen:")
    print("1. Semanal (por mes y selección de semana)")
    print("2. Mensual (por mes y año)")
    print("3. Anual (por año)")
    print("4. Neto (todas las inspecciones registradas)")
    print("5. Salir")

    choice = ""
    while choice not in ["1", "2", "3", "4", "5"]:
        choice = input("Ingrese su opción (1-5): ")
        if choice not in ["1", "2", "3", "4", "5"]: print("Opción no válida.")

    if choice == "5": print("Saliendo del programa de reportes."); exit()

    start_date_iso, end_date_iso, map_title_suffix = None, None, ""
    current_year = datetime.datetime.now().year

    try:
        if choice == "1": # Semanal por mes
            year = int(input(f"Ingrese el año (AAAA) [default: {current_year}]: ") or current_year)
            month = int(input("Ingrese el mes (1-12): "))
            if not (1 <= month <= 12): raise ValueError("Mes inválido.")

            print(f"\nSemanas para {calendar.month_name[month]} {year} (Domingo a Sábado):")
            weeks_in_month_options = []
            
            first_day_of_month = datetime.date(year, month, 1)
            day_of_week_first_day = first_day_of_month.weekday() # 0=Lunes, 6=Domingo
            days_until_prev_sunday = (day_of_week_first_day + 1) % 7 

            current_sunday = first_day_of_month - datetime.timedelta(days=days_until_prev_sunday)
            
            option_count = 1
            while True:
                end_of_week_date = current_sunday + datetime.timedelta(days=6)
                
                week_overlaps_month = False
                if (current_sunday.month == month and current_sunday.year == year) or \
                   (end_of_week_date.month == month and end_of_week_date.year == year) or \
                   (current_sunday < first_day_of_month and end_of_week_date >= first_day_of_month and end_of_week_date.month == month):
                    week_overlaps_month = True
                
                if week_overlaps_month:
                    display_text = (f"{option_count}. Del {current_sunday.day} {current_sunday.strftime('%B')} "
                                    f"al {end_of_week_date.day} {end_of_week_date.strftime('%B')} {end_of_week_date.year}")
                    weeks_in_month_options.append({
                        "number": option_count, 
                        "start_dt_obj": current_sunday, 
                        "end_dt_obj": end_of_week_date, 
                        "display": display_text
                    })
                    print(display_text)
                    option_count += 1
                
                current_sunday += datetime.timedelta(days=7)
                
                if current_sunday.year > year or (current_sunday.year == year and current_sunday.month > month):
                    last_day_of_month = datetime.date(year, month, calendar.monthrange(year, month)[1])
                    if current_sunday > last_day_of_month + datetime.timedelta(days=7): 
                        break
                    if not (current_sunday.month == month or end_of_week_date.month == month):
                        break


            if not weeks_in_month_options:
                print("No se generaron opciones de semana para el mes y año seleccionados. Verifique las fechas.")
                exit()

            selected_week_num = 0
            while selected_week_num not in [opt['number'] for opt in weeks_in_month_options]:
                try:
                    selected_week_num = int(input("Seleccione el número de semana: "))
                    if selected_week_num not in [opt['number'] for opt in weeks_in_month_options]:
                        print("Número de semana inválido.")
                except ValueError: print("Entrada inválida.")
            
            selected_week_data = next(opt for opt in weeks_in_month_options if opt['number'] == selected_week_num)
            start_dt = datetime.datetime.combine(selected_week_data['start_dt_obj'], datetime.time.min)
            end_dt = datetime.datetime.combine(selected_week_data['end_dt_obj'], datetime.time.max)
            start_date_iso, end_date_iso = start_dt.isoformat(timespec='seconds'), end_dt.isoformat(timespec='seconds')
            map_title_suffix = f"Semana del {selected_week_data['start_dt_obj'].strftime('%d %B')} al {selected_week_data['end_dt_obj'].strftime('%d %B %Y')}"

        elif choice == "2": # Mensual
            year = int(input(f"Ingrese el año (AAAA) [default: {current_year}]: ") or current_year)
            month = int(input("Ingrese el mes (1-12): "))
            if not (1 <= month <= 12): raise ValueError("Mes inválido.")
            _, num_days = calendar.monthrange(year, month)
            start_dt = datetime.datetime(year, month, 1, 0, 0, 0)
            end_dt = datetime.datetime(year, month, num_days, 23, 59, 59)
            start_date_iso, end_date_iso = start_dt.isoformat(timespec='seconds'), end_dt.isoformat(timespec='seconds')
            map_title_suffix = f"{datetime.date(year, month, 1).strftime('%B')} {year}"
        elif choice == "3": # Anual
            year = int(input(f"Ingrese el año (AAAA) [default: {current_year}]: ") or current_year)
            start_dt = datetime.datetime(year, 1, 1, 0, 0, 0)
            end_dt = datetime.datetime(year, 12, 31, 23, 59, 59)
            start_date_iso, end_date_iso = start_dt.isoformat(timespec='seconds'), end_dt.isoformat(timespec='seconds')
            map_title_suffix = f"Año {year}"
        elif choice == "4": # Neto
            map_title_suffix = "Todas las Inspecciones (Neto)"
    
    except ValueError as ve: print(f"Error en la entrada de fecha: {ve}"); exit()
    except Exception as e_date: print(f"Error al procesar fechas: {e_date}"); exit()


    conn = conectar_db()
    if conn:
        cursor = conn.cursor()
        query = "SELECT * FROM observaciones_embarcaciones"
        params = []
        if start_date_iso and end_date_iso:
            query += " WHERE timestamp BETWEEN ? AND ?"
            params.extend([start_date_iso, end_date_iso])
        query += " ORDER BY timestamp ASC"
        
        try:
            cursor.execute(query, params)
            registros_filtrados = cursor.fetchall()
        except sqlite3.Error as e_sql:
            print(f"Error al consultar la base de datos: {e_sql}")
            registros_filtrados = []
        conn.close()

        if registros_filtrados:
            registros_filtrados_dict = [dict(row) for row in registros_filtrados] 
            
            print(f"\nSe encontraron {len(registros_filtrados_dict)} registros para el periodo: {map_title_suffix}.")
            print("Generando mapa...")
            fig_resumen, ax_resumen = graficar_mapa_inspecciones(registros_filtrados_dict, f"Resumen Inspecciones: {map_title_suffix}")
            
            if fig_resumen: 
                plt.show(block=True) 
                
                while True:
                    generar_word_choice = input(f"¿Deseas guardar un reporte en Word (DOCX) para el resumen '{map_title_suffix}'? (s/n): ").lower()
                    if generar_word_choice in ['s', 'n']:
                        break
                    else:
                        print("Respuesta inválida. Por favor, ingresa 's' o 'n'.")

                if generar_word_choice == 's':
                    base_name = f"Resumen_{map_title_suffix}".replace(' ', '_').replace(':', '').replace('/', '_').replace('(', '').replace(')', '')
                    report_word_filename = f"{base_name}_Reporte.docx"
                    print(f"DEBUG: El usuario eligió generar el reporte Word para el resumen. Llamando a generar_reporte_word...")
                    generar_reporte_word(fig_resumen, registros_filtrados_dict, f"Resumen: {map_title_suffix}", report_word_filename)
                else:
                    print(f"Generación de reporte Word para el resumen omitida por el usuario.")
            else:
                print("No se pudo generar el mapa de resumen (quizás no hay datos válidos después de filtrar).")
        else:
            print(f"No se encontraron registros para el periodo: {map_title_suffix}.")
    else:
        print("No se pudo acceder a la base de datos para generar el reporte.")

    print("\nPrograma de reportes terminado.")
    input("Presiona Enter para cerrar esta consola...")