# app.py

# 1. --- TODAS TUS IMPORTACIONES ---
import os
import io
import base64
import csv
from flask import Flask, request, render_template, send_file, redirect, url_for, flash, session, jsonify
import matplotlib.pyplot as plt
import contextily as cx
from shapely.geometry import Point, Polygon, MultiPoint
from pyproj import Transformer, CRS
import numpy as np
import datetime
import calendar
import locale
from functools import wraps # Importar wraps para el decorador

# IMPORTACIONES CLAVE PARA POSTGRESQL
import psycopg2
from urllib.parse import urlparse

# Importar para manejar documentos Word
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Nuevas importaciones para Flask-Login y seguridad
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash


# 2. --- CONFIGURACIÓN DE LA BASE DE DATOS Y FUNCIONES DE CONEXIÓN/INICIALIZACIÓN ---
DATABASE_URL = os.environ.get('DATABASE_URL')

def conectar_db():
    """
    Establece una conexión a la base de datos PostgreSQL usando la URL de entorno.
    """
    if not DATABASE_URL:
        print("ERROR: DATABASE_URL no está configurada en el entorno.")
        return None
    
    try:
        url = urlparse(DATABASE_URL)
        conn = psycopg2.connect(
            database=url.path[1:],
            user=url.username,
            password=url.password,
            host=url.hostname,
            port=url.port,
            sslmode='require' # Supabase/Neon requieren SSL. Esto es importante.
        )
        return conn
    except Exception as e:
        print(f"ERROR: No se pudo conectar a la base de datos PostgreSQL: {e}")
        return None

def inicializar_db():
    """
    Inicializa las tablas 'observaciones_embarcaciones' y 'users' en PostgreSQL si no existen.
    Asegura que la tabla 'users' tiene las columnas 'is_approved' y 'role'.
    Inserta un usuario administrador por defecto si no existe ninguno.
    """
    conn = conectar_db()
    if not conn: return
    cursor = conn.cursor()
    try:
        # --- Creación de tabla 'observaciones_embarcaciones' ---
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS observaciones_embarcaciones (
            id SERIAL PRIMARY KEY,
            matricula TEXT NOT NULL,
            nombre_embarcacion TEXT,
            timestamp TIMESTAMP NOT NULL,
            latitud_wgs84 REAL NOT NULL,
            longitud_wgs84 REAL NOT NULL,
            tipo_embarcacion_id TEXT,
            estatus_categoria_id TEXT,
            notas_adicionales TEXT,
            nombre_patron TEXT
        )
        """)
        conn.commit()
        print("Tabla 'observaciones_embarcaciones' inicializada/verificada en PostgreSQL con TIMESTAMP.")

        # --- Creación y actualización de tabla 'users' ---
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id SERIAL PRIMARY KEY,
            username VARCHAR(80) UNIQUE NOT NULL,
            password_hash VARCHAR(255) NOT NULL
        );
        """)
        conn.commit()
        print("Tabla 'users' inicializada/verificada en PostgreSQL.")
        
        # Añadir columna 'is_approved' si no existe
        cursor.execute("""
        SELECT column_name FROM information_schema.columns 
        WHERE table_name='users' AND column_name='is_approved';
        """)
        if cursor.fetchone() is None:
            cursor.execute("ALTER TABLE users ADD COLUMN is_approved BOOLEAN DEFAULT FALSE;")
            conn.commit()
            print("Columna 'is_approved' añadida a la tabla 'users'.")
        else:
            print("La columna 'is_approved' ya existe en la tabla 'users'.")

        # Añadir columna 'role' si no existe
        # Los roles serán 'viewer', 'editor', 'admin'
        cursor.execute("""
        SELECT column_name FROM information_schema.columns 
        WHERE table_name='users' AND column_name='role';
        """)
        if cursor.fetchone() is None:
            cursor.execute("ALTER TABLE users ADD COLUMN role TEXT DEFAULT 'viewer';") # Nuevo rol por defecto
            conn.commit()
            print("Columna 'role' añadida a la tabla 'users'.")
        else:
            print("La columna 'role' ya existe en la tabla 'users'.")

        # Verificar y añadir la columna 'nombre_patron' si no existe
        cursor.execute("""
        SELECT column_name FROM information_schema.columns 
        WHERE table_name='observaciones_embarcaciones' AND column_name='nombre_patron';
        """)
        if cursor.fetchone() is None:
            cursor.execute("ALTER TABLE observaciones_embarcaciones ADD COLUMN nombre_patron TEXT")
            conn.commit()
            print("Columna 'nombre_patron' añadida a la base de datos PostgreSQL.")
        else:
            print("La columna 'nombre_patron' ya existe en PostgreSQL.")

        # Añadir índice único para evitar duplicados en importaciones (ya estaba, la mantengo)
        cursor.execute("""
        DO $$ BEGIN
            IF NOT EXISTS (SELECT 1 FROM pg_constraint WHERE conname = 'unique_matricula_timestamp') THEN
                ALTER TABLE observaciones_embarcaciones
                ADD CONSTRAINT unique_matricula_timestamp UNIQUE (matricula, timestamp);
            END IF;
        END $$;
        """)
        conn.commit()
        print("Índice único 'unique_matricula_timestamp' verificado/creado en PostgreSQL.")

        # --- Creación de usuario administrador por defecto si no existe ninguno ---
        cursor.execute("SELECT COUNT(*) FROM users WHERE role = 'admin';")
        admin_count = cursor.fetchone()[0]
        if admin_count == 0:
            default_admin_username = os.environ.get('DEFAULT_ADMIN_USERNAME', 'admin')
            default_admin_password = os.environ.get('DEFAULT_ADMIN_PASSWORD', 'adminpass') # ¡CAMBIAR EN PRODUCCIÓN!
            
            # Verificar si ya existe un usuario con el nombre de usuario de administrador por defecto
            cursor.execute("SELECT id FROM users WHERE username = %s;", (default_admin_username,))
            existing_admin = cursor.fetchone()

            if not existing_admin:
                admin_password_hash = generate_password_hash(default_admin_password)
                cursor.execute("""
                    INSERT INTO users (username, password_hash, is_approved, role)
                    VALUES (%s, %s, TRUE, 'admin') RETURNING id;
                """, (default_admin_username, admin_password_hash))
                conn.commit()
                print(f"Usuario administrador por defecto '{default_admin_username}' creado con ID: {cursor.fetchone()[0]}.")
            else:
                # Si ya existe pero no tiene rol de admin o no está aprobado, se actualiza
                cursor.execute("""
                    UPDATE users SET is_approved = TRUE, role = 'admin'
                    WHERE username = %s AND (is_approved = FALSE OR role != 'admin');
                """, (default_admin_username,))
                conn.commit()
                if cursor.rowcount > 0:
                    print(f"Usuario '{default_admin_username}' actualizado a administrador y aprobado.")
                else:
                    print(f"Usuario '{default_admin_username}' ya existe y es administrador aprobado.")

    except psycopg2.Error as e:
        print(f"Error al inicializar la base de datos PostgreSQL: {e}")
        conn.rollback()
    finally:
        cursor.close()
        conn.close()


# 3. --- DEFINICIÓN DE LA CLASE USER Y FUNCIONES DE INTERACCIÓN CON LA BASE DE DATOS PARA USUARIOS ---

# Clase User (Modelo de Usuario)
class User(UserMixin):
    def __init__(self, id, username, password_hash, is_approved=False, role='viewer'): # Rol por defecto 'viewer'
        self.id = id
        self.username = username
        self.password_hash = password_hash
        self.is_approved = is_approved # Nuevo atributo
        self.role = role             # Nuevo atributo

    def get_id(self):
        return str(self.id)

    # Propiedad requerida por Flask-Login para verificar si la cuenta está activa
    @property
    def is_active(self):
        return self.is_approved # Un usuario está "activo" si ha sido aprobado

    # Propiedad para verificar si el usuario tiene un rol específico
    def has_role(self, required_role):
        # Define una jerarquía de roles si es necesario para el acceso:
        # admin > editor > viewer
        roles_hierarchy = {
            'viewer': 0,
            'editor': 1,
            'admin': 2
        }
        user_role_level = roles_hierarchy.get(self.role, -1)
        required_role_level = roles_hierarchy.get(required_role, -1)
        
        return user_role_level >= required_role_level

# Función para buscar un usuario por su ID
def get_user_by_id(user_id):
    conn = None
    cur = None
    try:
        conn = conectar_db()
        if not conn: return None
        cur = conn.cursor()
        # Seleccionar también 'is_approved' y 'role'
        cur.execute("SELECT id, username, password_hash, is_approved, role FROM users WHERE id = %s", (user_id,))
        user_data = cur.fetchone()
        if user_data:
            return User(user_data[0], user_data[1], user_data[2], user_data[3], user_data[4])
        return None
    except Exception as e:
        print(f"Error al buscar usuario por ID: {e}")
        return None
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

# Función para buscar un usuario por su nombre de usuario
def get_user_by_username(username):
    conn = None
    cur = None
    try:
        print(f"DEBUG DB: Intentando buscar usuario '{username}' por nombre.")
        conn = conectar_db()
        if not conn:
            print("DEBUG DB: Falló la conexión a la base de datos en get_user_by_username.")
            return None
        
        cur = conn.cursor()
        # Seleccionar también 'is_approved' y 'role'
        cur.execute("SELECT id, username, password_hash, is_approved, role FROM users WHERE username = %s", (username,))
        user_data = cur.fetchone()

        if user_data:
            print(f"DEBUG DB: Usuario '{username}' ENCONTRADO. Datos: {user_data}")
            return User(user_data[0], user_data[1], user_data[2], user_data[3], user_data[4])
        else:
            print(f"DEBUG DB: Usuario '{username}' NO ENCONTRADO en la base de datos (get_user_by_username).")
            return None
    except Exception as e:
        print(f"DEBUG DB: Error en get_user_by_username para '{username}': {e}")
        return None
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

# Función para añadir un nuevo usuario (registrarse)
def add_user(username, password):
    conn = None
    cur = None
    try:
        conn = conectar_db()
        if not conn: return None
        cur = conn.cursor()
        password_hash = generate_password_hash(password)
        # Añade la columna is_approved con valor FALSE y role 'viewer' al insertar
        cur.execute("INSERT INTO users (username, password_hash, is_approved, role) VALUES (%s, %s, FALSE, 'viewer') RETURNING id", (username, password_hash))
        new_user_id = cur.fetchone()[0]
        conn.commit()
        print(f"Usuario {username} añadido con ID: {new_user_id} (No aprobado).")
        # Retorna la instancia de User con is_approved=False y role='viewer'
        return User(new_user_id, username, password_hash, is_approved=False, role='viewer')
    except psycopg2.IntegrityError as e:
        if "duplicate key value violates unique constraint" in str(e):
            print(f"Error: El nombre de usuario '{username}' ya existe.")
            conn.rollback() 
            return None
        else:
            print(f"Error de integridad al añadir usuario: {e}")
            conn.rollback() 
            return None
    except Exception as e:
        print(f"Error al añadir usuario: {e}")
        if conn: conn.rollback() 
        return None
    finally:
        if cur:
            cur.close()
        if conn:
            conn.close()

def get_all_users():
    """Obtiene todos los usuarios registrados."""
    conn = None
    cur = None
    try:
        conn = conectar_db()
        if not conn: return []
        cur = conn.cursor()
        cur.execute("SELECT id, username, is_approved, role FROM users ORDER BY username ASC;")
        # No recuperar password_hash por seguridad
        users_data = cur.fetchall()
        # Convertir a lista de diccionarios para facilitar el manejo en templates
        users_list = []
        for user_id, username, is_approved, role in users_data:
            users_list.append({
                'id': user_id,
                'username': username,
                'is_approved': is_approved,
                'role': role
            })
        return users_list
    except Exception as e:
        print(f"Error al obtener todos los usuarios: {e}")
        return []
    finally:
        if cur: cur.close()
        if conn: conn.close()

def update_user_status_and_role(user_id, is_approved, role):
    """Actualiza el estado de aprobación y el rol de un usuario."""
    conn = None
    cur = None
    try:
        conn = conectar_db()
        if not conn: return False
        cur = conn.cursor()
        cur.execute("""
            UPDATE users SET is_approved = %s, role = %s WHERE id = %s;
        """, (is_approved, role, user_id))
        conn.commit()
        return cur.rowcount > 0
    except Exception as e:
        print(f"Error al actualizar el usuario {user_id}: {e}")
        if conn: conn.rollback()
        return False
    finally:
        if cur: cur.close()
        if conn: conn.close()

# NUEVA FUNCIÓN: Actualizar la contraseña de un usuario
def update_user_password(user_id, new_password_hash):
    conn = None
    cur = None
    try:
        conn = conectar_db()
        if not conn: return False
        cur = conn.cursor()
        cur.execute("""
            UPDATE users SET password_hash = %s WHERE id = %s;
        """, (new_password_hash, user_id))
        conn.commit()
        return cur.rowcount > 0
    except Exception as e:
        print(f"Error al actualizar la contraseña del usuario {user_id}: {e}")
        if conn: conn.rollback()
        return False
    finally:
        if cur: cur.close()
        if conn: conn.close()

# NUEVA FUNCIÓN: Eliminar un usuario de la base de datos
def delete_user_db(user_id):
    """Elimina un usuario de la base de datos por su ID."""
    conn = None
    cur = None
    try:
        conn = conectar_db()
        if not conn: return False
        cur = conn.cursor()
        cur.execute("DELETE FROM users WHERE id = %s;", (user_id,))
        conn.commit()
        return cur.rowcount > 0 # Retorna True si se eliminó al menos una fila
    except Exception as e:
        print(f"Error al eliminar usuario {user_id}: {e}")
        if conn: conn.rollback()
        return False
    finally:
        if cur: cur.close()
        if conn: conn.close()


# 4. --- INICIALIZACIÓN DE LA APLICACIÓN FLASK ---
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'una_clave_secreta_muy_larga_y_aleatoria_para_desarrollo_cambiala_en_produccion')

# Inicialización de Flask-Login
login_manager = LoginManager()
login_manager.init_app(app) 
login_manager.login_view = 'login' 

@login_manager.user_loader
def load_user(user_id):
    return get_user_by_id(user_id)

# Decorador para requerir rol de administrador
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.has_role('admin'):
            flash('Acceso denegado: Necesitas ser un administrador.', 'error')
            return redirect(url_for('index')) # Redirige al index o a una página de error
        return f(*args, **kwargs)
    return decorated_function

# Decorador para requerir rol de editor (o superior)
def editor_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.has_role('editor'):
            flash('Acceso denegado: Necesitas ser un editor o administrador.', 'error')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

# Decorador para requerir cualquier usuario autenticado y aprobado (viewer o superior)
def viewer_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_active: # is_active ya verifica is_approved
            flash('Acceso denegado: Necesitas iniciar sesión y tu cuenta debe estar aprobada.', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


# Asegurarse de que DB se inicializa y la columna 'nombre_patron' existe al inicio
with app.app_context():
    inicializar_db()


# 5. --- TRANSFORMADORES DE COORDENADAS y DEFINICIONES GLOBALES ---
crs_utm_anp = CRS("EPSG:32613") 
crs_geo = CRS("EPSG:4326")     
crs_mercator = CRS("EPSG:3857") 
try:
    transformer_utm_to_geo = Transformer.from_crs(crs_utm_anp, crs_geo, always_xy=True)
    transformer_geo_to_utm = Transformer.from_crs(crs_geo, crs_utm_anp, always_xy=True)
    transformer_geo_to_mercator = Transformer.from_crs(crs_geo, crs_mercator, always_xy=True)
    transformer_mercator_to_geo = Transformer.from_crs(crs_mercator, crs_geo, always_xy=True)
except Exception as e:
    print(f"Error crítico al crear los transformadores de coordenadas: {e}"); exit()

# FUNCIONES AUXILIARES DE GEOMETRÍA Y FORMATO
def transform_coords_list(coords_list, transformer):
    if not coords_list: return []
    if not all(isinstance(coord, tuple) and len(coord) == 2 for coord in coords_list):
        if all(isinstance(coord, Point) for coord in coords_list):
             x_coords = np.array([coord.x for coord in coords_list])
             y_coords = np.array([coord.y for coord in coords_list])
        else: return []
    else:
        x_coords = np.array([coord[0] for coord in coords_list])
        y_coords = np.array([coord[1] for coord in coords_list])
    if x_coords.size == 0: return []
    trans_x, trans_y = transformer.transform(x_coords, y_coords)
    return list(zip(trans_x, trans_y))

def dd_to_gmm_str(dd_val, is_latitude):
    hemisphere = ('N' if dd_val >= 0 else 'S') if is_latitude else ('E' if dd_val >= 0 else 'W')
    abs_dd = abs(dd_val); degrees = int(abs_dd); minutes_decimal = (abs_dd - degrees) * 60
    return f"{degrees}°{minutes_decimal:.3f}' {hemisphere}"

def gms_to_dd(grados, minutos, segundos, hemisferio):
    try:
        dd = float(grados) + (float(minutos) / 60.0) + (float(segundos) / 3600.0)
        if hemisferio.upper() in ['S', 'O', 'W']: dd *= -1
        elif hemisferio.upper() not in ['N', 'E']: raise ValueError(f"Hemisferio GMS '{hemisferio}' no reconocido.")
        return dd
    except ValueError: raise ValueError("Valores de GMS inválidos.")

def gdm_to_dd(grados_str, minutos_decimales_str, hemisferio):
    try:
        grados = float(grados_str)
        minutos_decimales = float(minutos_decimales_str)
        dd = grados + (minutos_decimales / 60.0)
        if hemisferio.upper() in ['S', 'O', 'W']:
            dd *= -1
        elif hemisferio.upper() not in ['N', 'E']:
            raise ValueError(f"Hemisferio GDM '{hemisferio}' no reconocido. Usar N, S, E, W.")
        return dd
    except ValueError:
        raise ValueError("Valores de grados o minutos decimales (GDM) inválidos.")

# DEFINICIONES DE COLORES, MARCADORES Y ESTATUS
STATUS_CATEGORIES_INSIDE_ANP = {
    1: {"id": "paso_inocente", "desc": "Paso Inocente", "color_key": "blanco"},
    2: {"id": "turistico_autorizado", "desc": "Turístico Autorizado", "color_key": "verde"},
    3: {"id": "investigacion", "desc": "Investigación Autorizada", "color_key": "azul_marino"},
    4: {"id": "doc_nav_issue", "desc": "Inconsistencias Doc. / Nav.", "color_key": "amarillo"},
    5: {"id": "pesca_lgpas_issue", "desc": "Infracción LGPAS (Pesca/Acuacultura)", "color_key": "anaranjado"},
    6: {"id": "delito", "desc": "Delito Ambiental / Otro", "color_key": "rojo"},
}
STATUS_COLORS = {
    "blanco": "white",
    "verde": "green",
    "azul_marino": "darkblue",
    "amarillo": "yellow",
    "anaranjado": "orange",
    "rojo": "red",
    "outside_anp": "deepskyblue",
    "unknown_status": "lightgray"
}
VESSEL_TYPES = {
    "panga": {"id": "panga", "desc": "Panga / Emb. Menor", "marker_char": "^", "size_factor": 100},
    "yate": {"id": "yate", "desc": "Yate / Emb. Mayor", "marker_char": "s", "size_factor": 150},
    "otra": {"id": "otra", "desc": "Otra / No especificada", "marker_char": "o", "size_factor": 80}
}
DEFAULT_VESSEL_TYPE_INFO = {"id": "default", "desc": "Desconocido", "marker_char": "o", "size_factor": 80}

# COORDENADAS UTM ORIGINALES DEL ANP (Completas)
anp_maritime_boundary_coords_utm = [
    (327983.720703, 2441179.856690), (406635.043518, 2359341.216920),
    (368280.494690, 2319018.426330), (287871.587708, 2401061.958310),
    (327983.720703, 2441179.856690)
]
isla_maria_madre_coords_utm = [
    (340647.669678, 2396167.727910), (340579.711304, 2396063.339720), (340611.461304, 2395997.193730),
    (340587.648682, 2395913.849670), (340595.586304, 2395900.620480), (340645.857117, 2395890.037290),
    (340611.461304, 2395736.578670), (340628.659302, 2395618.838680), (340570.450684, 2395366.161070),
    (340525.471497, 2395249.744320), (340471.231689, 2395225.931700), (340487.543884, 2395182.234130),
    (340368.570129, 2395208.520080), (340324.400085, 2395007.740110), (340479.769714, 2394974.235470),
    (340393.179688, 2394903.139530), (340373.335876, 2394827.733090), (340450.065125, 2394766.878720),
    (340397.148315, 2394584.315920), (340256.918884, 2394516.846920), (340233.106323, 2394463.930300),
    (340152.408325, 2394445.409300), (340148.439514, 2394399.107120), (340182.835510, 2394362.065490),
    (340312.481506, 2394327.669490), (340373.335876, 2394281.367310), (340362.752502, 2394200.669310),
    (340446.096313, 2394192.731690), (340432.867126, 2394129.231690), (340327.033691, 2393958.575070),
    (340280.731506, 2393929.470890), (340268.825073, 2393908.304080), (340129.918701, 2393850.095700),
    (340141.824890, 2393802.470520), (340083.937500, 2393740.018130), (339996.430115, 2393660.810120),
    (339967.940125, 2393661.430110), (339604.670105, 2393789.380130),
    (339527.150085, 2393658.210080), (339518.860107, 2393626.080080), (339516.220093, 2393591.930110),
    (339584.590088, 2393563.660100), (339650.840088, 2393463.400090), (339887.670105, 2393439.740110),
    (340161.430115, 2393477.410100), (340297.929321, 2393484.969910), (340489.607100, 2393022.572000),
    (340412.379000, 2392975.619700), (340361.409100, 2392954.032300), (340321.785500, 2392931.381300),
    (340295.962100, 2392918.257900), (340290.458700, 2392910.214500), (340269.927000, 2392865.976100),
    (340255.322000, 2392819.409400), (340243.892000, 2392791.892600), (340230.980300, 2392769.244300),
    (340215.105300, 2392759.719200), (340187.800200, 2392760.777600), (340168.593600, 2392758.504800),
    (340127.318500, 2392710.879700), (340112.203000, 2392694.661800), (340098.476700, 2392672.512400),
    (340074.983900, 2392646.056700), (340039.476900, 2392595.735600), (340036.037100, 2392587.001500),
    (340042.598800, 2392550.806500), (340065.458800, 2392523.713100), (340072.867200, 2392501.064700),
    (340077.698000, 2392477.208100), (340080.063900, 2392465.716300), (340076.677200, 2392463.388000),
    (340055.933800, 2392472.489700), (340049.372100, 2392473.336300), (340037.307100, 2392463.599600),
    (340057.819400, 2392442.841800), (340062.006000, 2392433.844500), (340061.159400, 2392424.107800),
    (340061.177000, 2392414.711600), (340047.829000, 2392415.398000), (340032.986600, 2392421.967400),
    (340029.475400, 2392418.726200), (340031.380400, 2392404.121200), (340038.365500, 2392383.589500),
    (340007.825700, 2392369.514200), (340002.805400, 2392359.882800), (340003.440400, 2392353.321100),
    (340018.680400, 2392324.322700), (340029.119900, 2392301.251200), (340036.102300, 2392288.919400),
    (340046.766000, 2392275.041300), (340037.942100, 2392262.092600), (340036.989600, 2392259.870100),
    (340040.799600, 2392257.647600), (340063.310200, 2392241.197300), (340060.715200, 2392232.911600),
    (340028.624400, 2392200.293200), (340015.895800, 2392181.650300), (339983.620400, 2392158.470700),
    (339965.552000, 2392150.173600), (339954.756900, 2392150.491100), (339942.374400, 2392157.476100),
    (339938.405700, 2392156.682400), (339936.818200, 2392137.632300), (339919.831900, 2392104.771000),
    (339939.020600, 2392084.476200), (339937.134600, 2392080.204300), (339911.418100, 2392080.799700),
    (339912.211900, 2392075.402200), (339925.388100, 2392064.607200), (339940.568600, 2392056.879300),
    (339941.277200, 2392048.862500), (339909.033400, 2392023.385200), (339908.651200, 2392005.784300),
    (339922.730700, 2391995.347100), (339944.120700, 2391991.740800), (339965.710700, 2391962.848200),
    (339996.368800, 2391935.821500), (340012.285400, 2391933.116600), (340015.928400, 2391935.073400),
    (340036.299500, 2391932.656400), (340036.593700, 2391910.177000), (340038.724700, 2391895.429900),
    (340044.230000, 2391884.233600), (340067.388800, 2391852.406400), (340078.899700, 2391843.468000),
    (340079.534700, 2391831.244200), (340072.769000, 2391832.081600), (340056.039700, 2391848.865500),
    (340034.535600, 2391848.404200), (340020.955800, 2391857.279300), (339987.300800, 2391874.900500),
    (339979.065600, 2391870.153000), (339987.955600, 2391845.705500), (339999.610500, 2391840.840700),
    (339990.879300, 2391836.342700), (340004.227500, 2391814.308200), (340023.198200, 2391823.797900),
    (340009.929300, 2391845.074000), (340007.010100, 2391843.477700), (339992.698300, 2391852.199200),
    (339993.015800, 2391857.279300), (340016.543900, 2391847.345900), (340033.900600, 2391843.747600),
    (340047.024000, 2391843.959200), (340057.131600, 2391841.214800), (340066.896700, 2391823.734500),
    (340075.493200, 2391821.945800), (340085.018200, 2391825.650000), (340089.102700, 2391828.828400),
    (340087.491400, 2391843.791800), (340084.344500, 2391845.818100), (340072.472200, 2391853.463100),
    (340061.777700, 2391871.174100), (340053.269200, 2391890.081400), (340041.236600, 2391908.580700),
    (340042.825000, 2391925.748300), (340041.017300, 2391938.588400), (340035.482100, 2391941.921100),
    (340029.687100, 2391943.956900), (340027.147100, 2391939.988200), (339995.147900, 2391941.743200),
    (339983.332000, 2391953.481900), (339968.885700, 2391965.229500), (339949.412100, 2391992.424300),
    (339943.803200, 2391996.027000), (339925.229400, 2391998.884500), (339912.211900, 2392005.869600),
    (339912.826100, 2392019.582900), (339943.485700, 2392040.953400), (339946.025700, 2392057.622200),
    (339944.279400, 2392060.638400), (339926.499400, 2392067.940900), (339919.196900, 2392074.132200),
    (339930.150700, 2392074.132200), (339938.246900, 2392077.783400), (339943.803200, 2392078.418400),
    (339945.708200, 2392084.609700), (339933.711500, 2392095.572300), (339925.388100, 2392107.311000),
    (339938.324900, 2392126.729100), (339939.389800, 2392129.708000), (339942.691900, 2392151.284800),
    (339960.630700, 2392144.299800), (339985.810800, 2392154.601400), (340018.796800, 2392178.369800),
    (340039.329100, 2392206.108200), (340055.922600, 2392219.451200), (340069.641900, 2392235.258700),
    (340068.263400, 2392242.883800), (340065.312800, 2392247.591800), (340043.339600, 2392260.187600),
    (340052.809200, 2392270.746200), (340052.547100, 2392275.745100), (340038.234200, 2392299.114700),
    (340021.643800, 2392326.439400), (340006.455700, 2392356.684200), (340013.690700, 2392363.109200),
    (340022.854200, 2392371.247000), (340043.183000, 2392375.030800), (340044.924900, 2392383.375200),
    (340036.672100, 2392404.544500), (340035.402100, 2392415.551200), (340047.890500, 2392410.259500),
    (340063.553800, 2392412.164500), (340066.305500, 2392431.637900), (340065.791800, 2392439.425500),
    (340042.598800, 2392463.388000), (340052.970500, 2392469.103000), (340073.076600, 2392460.422000),
    (340080.910500, 2392458.943000), (340085.910400, 2392465.840500), (340078.111300, 2392504.006200),
    (340069.903800, 2392526.464800), (340047.467100, 2392553.346500), (340042.138000, 2392575.950800),
    (340041.185500, 2392588.015800), (340041.961600, 2392592.078900), (340064.797200, 2392622.610600),
    (340077.698000, 2392642.625900), (340095.725000, 2392659.389000), (340115.006300, 2392686.786200),
    (340129.964400, 2392707.175500), (340164.375700, 2392746.607400), (340171.995700, 2392751.369900),
    (340180.830600, 2392754.932900), (340216.586900, 2392753.369200), (340233.943600, 2392764.375900),
    (340249.395300, 2392789.987600), (340259.784700, 2392818.362500), (340275.342200, 2392865.828800),
    (340296.456000, 2392909.485200), (340323.705900, 2392923.935900), (340362.813600, 2392949.014000),
    (340418.058700, 2392971.397800), (340445.046300, 2392986.637800), (340487.116700, 2393014.422800),
    (340491.723900, 2393017.465600), (340554.575684, 2392865.843690), (340735.815674, 2392689.895510),
    (340799.315918, 2392570.832700), (340914.409729, 2392618.457700), (340956.743286, 2392580.093080),
    (341315.254272, 2392087.967100), (341263.660522, 2392045.633730), (341321.868896, 2391935.831480),
    (341380.077271, 2392013.883730), (341389.337891, 2392037.696290), (341409.181702, 2392068.123290),
    (341415.997925, 2392074.462520), (341423.062317, 2392059.846680), (341452.431091, 2392036.431090),
    (341517.121887, 2392039.209110), (341533.790710, 2392066.196720), (341543.315674, 2392033.652890),
    (341538.156311, 2391976.502690), (341546.093689, 2391966.183900), (341556.015686, 2391974.121520),
    (341568.715698, 2392018.571470), (341602.450073, 2392075.324890), (341612.220276, 2392084.024110),
    (341634.077881, 2392050.925480), (341727.563110, 2392140.209110),
    (343552.698730, 2389885.991700), (343507.912109, 2389799.561280), (343518.495483, 2389691.081910),
    (343379.588928, 2389615.675480), (343288.307495, 2389620.967100), (343190.411499, 2389757.227720),
    (343072.671509, 2389714.894290), (343126.911316, 2389597.154480), (342863.650330, 2389430.466670),
    (342888.785889, 2389097.091130), (342875.556702, 2389093.122310), (342862.327271, 2388978.028320),
    (342916.567078, 2388898.653080), (342921.858704, 2388836.475890), (343022.400696, 2388661.850520),
    (343048.859131, 2388582.475520), (343100.876892, 2388586.024900), (343201.853088, 2388431.768920),
    (340521.691101, 2384841.964480), (341115.563904, 2386033.625120), (339699.978699, 2386665.757080),
    (339141.029114, 2385565.767700), (339139.558472, 2385562.847720), (333486.840088, 2388086.395690),
    (333487.039673, 2388086.762510), (333846.045471, 2388718.355900), (333589.775696, 2388845.270320),
    (333089.439514, 2388318.086910), (333088.560913, 2388317.270080), (331178.738525, 2391778.943910),
    (331506.668701, 2391960.168090), (331401.720276, 2392177.387330), (331079.552490, 2392206.675110),
    (330913.667908, 2392089.244320),
    (340647.669678, 2396167.727910)
]
puerto_balleto_coords_utm = [
    (340647.669623, 2396167.727920), (340871.403535, 2393096.078400), (340981.527284, 2393138.588180),
    (340974.280217, 2393158.040830), (340982.290133, 2393161.473650), (340989.537199, 2393142.402420),
    (341020.051165, 2393153.463730), (341042.173790, 2393168.720720), (341044.080913, 2393179.019180),
    (341005.938456, 2393236.232860), (341015.855495, 2393240.809960), (341066.966387, 2393165.669320),
    (341057.812197, 2393158.803680), (341046.369460, 2393161.855070), (341025.009684, 2393146.598090),
    (340992.588596, 2393134.392510), (340999.835663, 2393115.321280), (340992.970021, 2393113.795580),
    (340984.197256, 2393132.485380), (340873.875105, 2393088.869650), (341727.563090, 2392140.209000),
    (341634.077853, 2392050.925350), (341612.220175, 2392084.024120), (341602.450143, 2392075.324770),
    (341568.715700, 2392018.571530), (341556.015675, 2391974.121450), (341546.093780, 2391966.183930),
    (341538.156264, 2391976.502700), (341543.315650, 2392033.652810), (341533.790631, 2392066.196630),
    (341517.121847, 2392039.209080), (341452.431093, 2392036.430940), (341423.062284, 2392059.846620),
    (341415.997933, 2392074.462520), (341409.181570, 2392068.123300), (341389.337780, 2392037.696150),
    (341380.077345, 2392013.883610), (341321.868895, 2391935.831370), (341263.660445, 2392045.633670),
    (341315.254298, 2392087.967090), (340956.743165, 2392580.093070), (340914.409747, 2392618.457730),
    (340799.315766, 2392570.832640), (340735.815639, 2392689.895380), (340554.575694, 2392865.843640),
    (340297.929347, 2393484.969880), (340161.430002, 2393477.410000), (339887.670000, 2393439.740000),
    (339650.840002, 2393463.400000), (339584.590002, 2393563.660000), (339516.219999, 2393591.930000),
    (339518.860003, 2393626.080000), (339527.150000, 2393658.210000), (339604.670000, 2393789.380000),
    (339967.940001, 2393661.430000), (339996.430002, 2393660.810000), (340033.490000, 2393757.320000),
    (340083.937515, 2393740.018110), (340141.824868, 2393802.470520), (340129.918594, 2393850.095610),
    (340268.825122, 2393908.304060), (340280.731396, 2393929.470770), (340327.033572, 2393958.575000),
    (340432.867117, 2394129.231590), (340446.096310, 2394192.731710), (340362.752393, 2394200.669230),
    (340373.335748, 2394281.367310), (340312.481459, 2394327.669480), (340182.835367, 2394362.065390),
    (340148.439465, 2394399.107130), (340152.408223, 2394445.409300), (340233.106301, 2394463.930170),
    (340256.918848, 2394516.846950), (340397.148295, 2394584.315830), (340450.065068, 2394766.878700),
    (340373.335748, 2394827.732980), (340393.179538, 2394903.139390), (340479.769674, 2394974.235460),
    (340324.400000, 2395007.740000), (340368.570002, 2395208.520000), (340487.543742, 2395182.234030),
    (340471.231777, 2395225.931700), (340525.471469, 2395249.744250), (340570.450725, 2395366.161140),
    (340628.659175, 2395618.838730), (340611.461224, 2395736.578550), (340645.857126, 2395890.037190),
    (340595.586192, 2395900.620550), (340587.648676, 2395913.849740), (340611.461224, 2395997.193730),
    (340579.711161, 2396063.339620),
    (340647.669623, 2396167.727920)
]
islas_menores_data_utm = {
    "Isla San Juanito (V)": {"coords": [(328030.906100, 2404586.986800), (327633.883600, 2404310.100800)], "marker": "o", "color": "darkviolet"},
    "Islote El Morro (V)": {"coords": [(323879.578800, 2405116.471600), (323867.899400, 2405109.464000)], "marker": "s", "color": "firebrick"},
    "Isla María Magdalena (V)": {"coords": [(350862.351900, 2377854.999000), (351247.036700, 2377788.188800)], "marker": "^", "color": "olive"},
    "Isla María Cleofas (V)": {"coords": [
        (369429.510100, 2359596.764100), (369848.633600, 2359593.460300), (369866.088000, 2359593.322700),
        (369908.061600, 2359592.991800), (372259.125600, 2359574.458900), (372486.810100, 2359572.664100)
    ], "marker": "P", "color": "teal"},
    "Islote La Mona 1 (V)": {"coords": [(366477.352800, 2358421.914700), (366325.153200, 2358358.393600)], "marker": "*", "color": "darkorange"},
    "Islote La Mona 2 (V)": {"coords": [(367454.190600, 2356053.133500), (367435.503600, 2356048.461700)], "marker": "X", "color": "darkmagenta"},
    "Islote La Mona 3 (V)": {"coords": [(368306.785600, 2355808.450400), (368277.587200, 2355808.450400)], "marker": "D", "color": "navy"},
}

print("Cargando y transformando geometrías base del ANP...")
try:
    anp_maritime_boundary_coords_geo = transform_coords_list(anp_maritime_boundary_coords_utm, transformer_utm_to_geo)
    anp_maritime_polygon_geo = Polygon(anp_maritime_boundary_coords_geo) if anp_maritime_boundary_coords_geo else Polygon()
    anp_maritime_boundary_coords_mercator = transform_coords_list(anp_maritime_boundary_coords_geo, transformer_geo_to_mercator)

    isla_maria_madre_coords_geo = transform_coords_list(isla_maria_madre_coords_utm, transformer_utm_to_geo)
    isla_maria_madre_polygon_geo = Polygon(isla_maria_madre_coords_geo) if isla_maria_madre_coords_geo else Polygon()
    isla_maria_madre_coords_mercator = transform_coords_list(isla_maria_madre_coords_geo, transformer_geo_to_mercator)

    puerto_balleto_coords_geo = transform_coords_list(puerto_balleto_coords_utm, transformer_utm_to_geo)
    puerto_balleto_polygon_geo = Polygon(puerto_balleto_coords_geo) if puerto_balleto_coords_geo else Polygon()
    puerto_balleto_coords_mercator = transform_coords_list(puerto_balleto_coords_geo, transformer_geo_to_mercator)
    
    islas_menores_data_geo = { name: {"coords": transform_coords_list(data["coords"], transformer_utm_to_geo), "marker": data["marker"], "color": data["color"]} for name, data in islas_menores_data_utm.items() }
    islas_menores_data_mercator = { name: {"coords": transform_coords_list(data_geo["coords"], transformer_geo_to_mercator), "marker": data_geo["marker"], "color": data_geo["color"]} for name, data_geo in islas_menores_data_geo.items() }
    print("Geometrías base del ANP cargadas y transformadas.")
except Exception as e:
    print(f"Error fatal durante la transformación de coordenadas base del ANP: {e}"); exit()


# 6. --- FUNCIONES AUXILIARES DE DB Y REPORTE (generar_reporte_word, graficar_mapa_general, etc.) ---

def _fetch_as_dict(cursor):
    """
    Ayudante para obtener resultados de la consulta como lista de diccionarios.
    """
    columns = [col[0] for col in cursor.description]
    return [dict(zip(columns, row)) for row in cursor.fetchall()]

def agregar_observacion_db(matricula, nombre_embarcacion, avistamiento_timestamp, lat_wgs84, lon_wgs84, tipo_emb_id, estatus_cat_id, notas="", nombre_patron=""):
    """
    Inserta una nueva observación de embarcación en la base de datos.
    'avistamiento_timestamp' ahora debe ser un objeto datetime de Python.
    """
    conn = conectar_db()
    if not conn: return
    cursor = conn.cursor()
    try:
        cursor.execute("""
        INSERT INTO observaciones_embarcaciones
        (matricula, nombre_embarcacion, timestamp, latitud_wgs84, longitud_wgs84, tipo_embarcacion_id, estatus_categoria_id, notas_adicionales, nombre_patron)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (matricula.upper(), nombre_embarcacion, avistamiento_timestamp, lat_wgs84, lon_wgs84, tipo_emb_id, estatus_cat_id, notas, nombre_patron))
        conn.commit()
        print(f"Observación para '{matricula}' (Avistamiento: {avistamiento_timestamp}) guardada en PostgreSQL.")
    except psycopg2.Error as e:
        print(f"Error al guardar observación en la base de datos PostgreSQL: {e}")
        conn.rollback()
    finally:
        cursor.close()
        conn.close()

# NUEVA FUNCIÓN: Actualizar una observación existente
def update_observacion_db(obs_id, matricula, nombre_embarcacion, avistamiento_timestamp, lat_wgs84, lon_wgs84, tipo_emb_id, estatus_cat_id, notas="", nombre_patron=""):
    """
    Actualiza una observación de embarcación existente en la base de datos por su ID.
    """
    conn = conectar_db()
    if not conn: return False
    cursor = conn.cursor()
    try:
        cursor.execute("""
        UPDATE observaciones_embarcaciones
        SET matricula = %s, nombre_embarcacion = %s, timestamp = %s, 
            latitud_wgs84 = %s, longitud_wgs84 = %s, tipo_embarcacion_id = %s, 
            estatus_categoria_id = %s, notas_adicionales = %s, nombre_patron = %s
        WHERE id = %s
        """, (matricula.upper(), nombre_embarcacion, avistamiento_timestamp, lat_wgs84, lon_wgs84, 
              tipo_emb_id, estatus_cat_id, notas, nombre_patron, obs_id))
        conn.commit()
        return cursor.rowcount > 0 # Retorna True si se actualizó una fila
    except psycopg2.Error as e:
        print(f"Error al actualizar observación ID {obs_id} en la base de datos PostgreSQL: {e}")
        conn.rollback()
        return False
    finally:
        cursor.close()
        conn.close()

# NUEVA FUNCIÓN: Obtener una observación por su ID
def get_observacion_by_id(obs_id):
    """
    Obtiene una observación de embarcación por su ID.
    """
    conn = conectar_db()
    if not conn: return None
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT * FROM observaciones_embarcaciones WHERE id = %s", (obs_id,))
        registro = _fetch_as_dict(cursor)
        return registro[0] if registro else None
    except psycopg2.Error as e:
        print(f"Error al obtener observación por ID {obs_id} en PostgreSQL: {e}")
        return None
    finally:
        cursor.close()
        conn.close()


def buscar_historial_embarcacion(matricula):
    """
    Busca todas las observaciones para una matrícula específica.
    psycopg2 devuelve 'timestamp' como objeto datetime.
    """
    conn = conectar_db()
    if not conn: return []
    cursor = conn.cursor()
    try:
        cursor.execute("""
        SELECT * FROM observaciones_embarcaciones
        WHERE matricula = %s
        ORDER BY timestamp DESC
        """, (matricula.upper(),))
        registros = _fetch_as_dict(cursor)
        return registros
    except psycopg2.Error as e:
        print(f"Error al buscar historial en PostgreSQL: {e}")
        return []
    finally:
        cursor.close()
        conn.close()

def buscar_por_nombre_o_patron(nombre_embarcacion, nombre_patron):
    """
    Busca observaciones por nombre de embarcación o nombre de patrón (parcial o completo).
    psycopg2 devuelve 'timestamp' como objeto datetime.
    """
    conn = conectar_db()
    if not conn: return []
    cursor = conn.cursor()
    try:
        query = "SELECT * FROM observaciones_embarcaciones WHERE 1=1"
        params = []
        if nombre_embarcacion:
            query += " AND LOWER(nombre_embarcacion) LIKE LOWER(%s)"
            params.append(f'%{nombre_embarcacion}%')
        if nombre_patron:
            query += " AND LOWER(nombre_patron) LIKE LOWER(%s)"
            params.append(f'%{nombre_patron}%')
        
        query += " ORDER BY timestamp DESC"
        
        cursor.execute(query, tuple(params))
        registros = _fetch_as_dict(cursor)
        return registros
    except psycopg2.Error as e:
        print(f"Error al buscar por nombre/patrón en PostgreSQL: {e}")
        return []
    finally:
        cursor.close()
        conn.close()

def obtener_observaciones_filtradas(start_date_obj=None, end_date_obj=None, status_category_filter=None):
    """
    Obtiene observaciones dentro de un rango de fechas y/o por estatus de categoría.
    'start_date_obj' y 'end_date_obj' deben ser objetos datetime de Python.
    'status_category_filter' es el ID de estatus por el que se desea filtrar.
    """
    conn = conectar_db()
    if not conn: return []
    cursor = conn.cursor()
    try:
        query = "SELECT * FROM observaciones_embarcaciones WHERE 1=1"
        params = []

        if start_date_obj and end_date_obj:
            query += " AND timestamp BETWEEN %s AND %s"
            params.extend([start_date_obj, end_date_obj])
        
        if status_category_filter:
            # Si el filtro es "outside_anp", se filtra por eso
            if status_category_filter == "outside_anp":
                query += " AND estatus_categoria_id = %s"
                params.append(status_category_filter)
            else:
                # Si el filtro es un estatus dentro del ANP, se filtra por su ID
                # Primero, encontrar el ID numérico si el filtro viene como nombre legible
                found_status_id_for_query = None
                for k_int, v_dict in STATUS_CATEGORIES_INSIDE_ANP.items():
                    if v_dict['id'] == status_category_filter:
                        found_status_id_for_query = status_category_filter # Usar el ID de texto como se almacena en DB
                        break
                
                if found_status_id_for_query:
                    query += " AND estatus_categoria_id = %s"
                    params.append(found_status_id_for_query)
                elif status_category_filter != "": # Evitar errores si el filtro está vacío (Todos los estatus)
                    print(f"ADVERTENCIA: Estatus de categoría '{status_category_filter}' no reconocido para el filtro.")
                    # Si no se encuentra, no se añade el filtro de estatus o se podría añadir un filtro que no devuelva nada
        
        query += " ORDER BY timestamp ASC"
        
        cursor.execute(query, tuple(params))
        registros = _fetch_as_dict(cursor)
        return registros
    except psycopg2.Error as e:
        print(f"Error al consultar la base de datos para resumen con filtro: {e}")
        return []
    finally:
        cursor.close()
        conn.close()

# NUEVA FUNCIÓN: Obtener conteo de observaciones por mes/año
def get_observation_counts_by_month_year():
    conn = conectar_db()
    if not conn: return []
    cursor = conn.cursor()
    try:
        cursor.execute("""
            SELECT 
                EXTRACT(YEAR FROM timestamp) AS year,
                EXTRACT(MONTH FROM timestamp) AS month,
                COUNT(*) AS count
            FROM observaciones_embarcaciones
            GROUP BY 1, 2
            ORDER BY 1 ASC, 2 ASC;
        """)
        return _fetch_as_dict(cursor)
    except Exception as e:
        print(f"Error al obtener conteos de observaciones por mes/año: {e}")
        return []
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# NUEVA FUNCIÓN: Obtener distribución de estatus
def get_status_distribution():
    conn = conectar_db()
    if not conn: return []
    cursor = conn.cursor()
    try:
        cursor.execute("""
            SELECT estatus_categoria_id, COUNT(*) AS count
            FROM observaciones_embarcaciones
            GROUP BY estatus_categoria_id
            ORDER BY count DESC;
        """)
        return _fetch_as_dict(cursor)
    except Exception as e:
        print(f"Error al obtener distribución de estatus: {e}")
        return []
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# NUEVA FUNCIÓN: Obtener embarcaciones recurrentes (ej. top 10 por matrícula)
def get_top_recurrent_vessels(limit=10):
    conn = conectar_db()
    if not conn: return []
    cursor = conn.cursor()
    try:
        cursor.execute("""
            SELECT matricula, COUNT(*) AS count
            FROM observaciones_embarcaciones
            GROUP BY matricula
            ORDER BY count DESC
            LIMIT %s;
        """, (limit,))
        return _fetch_as_dict(cursor)
    except Exception as e:
        print(f"Error al obtener embarcaciones recurrentes: {e}")
        return []
    finally:
        if cursor: cursor.close()
        if conn: conn.close()

# NUEVA FUNCIÓN: Obtener embarcaciones con estatus de infracción/delito repetido
def get_repeated_infraction_vessels(min_infractions=2):
    conn = conectar_db()
    if not conn: return []
    cursor = conn.cursor()
    try:
        # IDs de estatus que se consideran "infracción" o "delito"
        infraction_status_ids = [
            STATUS_CATEGORIES_INSIDE_ANP[5]['id'], # pesca_lgpas_issue
            STATUS_CATEGORIES_INSIDE_ANP[6]['id'], # delito
        ]
        
        # Filtramos por los estatus de infracción/delito y contamos las ocurrencias por matrícula
        cursor.execute(f"""
            SELECT matricula, COUNT(*) AS infraction_count, 
                   array_agg(estatus_categoria_id) AS all_status_ids,
                   array_agg(timestamp ORDER BY timestamp DESC) AS last_timestamps
            FROM observaciones_embarcaciones
            WHERE estatus_categoria_id IN %s
            GROUP BY matricula
            HAVING COUNT(*) >= %s
            ORDER BY infraction_count DESC;
        """, (tuple(infraction_status_ids), min_infractions))
        
        results = _fetch_as_dict(cursor)
        
        # Opcional: para cada infracción, obtener el último timestamp
        for res in results:
            if res['last_timestamps']:
                res['last_infraction_date'] = max(res['last_timestamps']).strftime('%Y-%m-%d %H:%M')
            else:
                res['last_infraction_date'] = 'N/A'
            # Convertir la lista de IDs de estatus a descripciones legibles
            res['all_status_descriptions'] = []
            for status_id in res['all_status_ids']:
                if status_id == 'outside_anp':
                    res['all_status_descriptions'].append("Fuera del Polígono ANP")
                else:
                    for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                        if cat_info['id'] == status_id:
                            res['all_status_descriptions'].append(cat_info['desc'])
                            break
        
        return results

    except Exception as e:
        print(f"Error al obtener embarcaciones con infracciones repetidas: {e}")
        return []
    finally:
        if cursor: cursor.close()
        if conn: conn.close()


def eliminar_observacion_db(id_observacion):
    """
    Elimina una observación por su ID.
    """
    conn = conectar_db()
    if not conn: return False
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM observaciones_embarcaciones WHERE id = %s", (id_observacion,))
        conn.commit()
        if cursor.rowcount > 0:
            print(f"Observación con ID {id_observacion} eliminada exitosamente de PostgreSQL.")
            return True
        else:
            print(f"No se encontró observación con ID {id_observacion} en PostgreSQL para eliminar.")
            return False
    except psycopg2.Error as e:
        print(f"Error al eliminar observación de la base de datos PostgreSQL: {e}")
        conn.rollback()
        return False
    finally:
        cursor.close()
        conn.close()

# FUNCIÓN PARA GENERAR REPORTE EN WORD (DOCX)
def generar_reporte_word(fig, observations_data, title, filename_or_buffer="reporte_inspeccion.docx"):
    """
    Genera un documento Word (.docx) con el mapa de Matplotlib y un resumen de observaciones.
    Acepta un buffer en memoria o un nombre de archivo para guardar.
    Las 'observations_data' deben contener objetos datetime para el timestamp.
    """
    print(f"DEBUG_WORD: Intentando generar reporte Word '{filename_or_buffer}' desde cero...")
    document = Document() 
    
    document.add_heading(title, level=1)
    document.add_paragraph() 

    temp_img_buffer = io.BytesIO() 
    try:
        print(f"DEBUG_WORD: Guardando imagen temporal del mapa en buffer...")
        fig.savefig(temp_img_buffer, format='png', dpi=300, bbox_inches='tight', pad_inches=0.1)
        plt.close(fig) 
        temp_img_buffer.seek(0) 
        print(f"DEBUG_WORD: Imagen temporal guardada en buffer.")
    except Exception as e:
        print(f"ERROR_WORD: Falló al guardar la imagen temporal del mapa para el Word: {e}")
        return
    
    # Añadir un salto de sección antes de la imagen para centrarla sin afectar el encabezado
    # Esta funcionalidad es más compleja en python-docx. Para simplificar,
    # se intentará añadir directamente y alinear.
    try:
        document.add_picture(temp_img_buffer, width=Inches(6.5)) 
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print(f"DEBUG_WORD: Imagen añadida al documento Word.")
    except Exception as e:
        print(f"ERROR_WORD: Falló al añadir la imagen al documento Word: {e}")
        return

    document.add_page_break()

    document.add_heading('Resumen de Observaciones:', level=2)
    
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8') 
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Mexico.1252') 
        except locale.Error:
            print("ADVERTENCIA: No se pudo configurar el locale español. Los meses se mostrarán en inglés.")

    sorted_observations = sorted(observations_data, key=lambda x: x.get('timestamp'))

    for i, obs in enumerate(sorted_observations):
        timestamp_dt_obj = obs.get('timestamp')
        
        if isinstance(timestamp_dt_obj, datetime.datetime):
            date_str = timestamp_dt_obj.strftime('%d')
            month_str = timestamp_dt_obj.strftime('%B')
            year_str = timestamp_dt_obj.strftime('%Y')
            time_str = timestamp_dt_obj.strftime('%H:%M')
        else:
            date_str, month_str, year_str, time_str = "N/A", "N/A", "N/A", "N/A"
            print(f"ADVERTENCIA: Timestamp no es objeto datetime para reporte Word. Tipo: {type(timestamp_dt_obj)}. Usando N/A.")

        lat_dd = obs.get('lat_dd') or obs.get('latitud_wgs84')
        lon_dd = obs.get('lon_dd') or obs.get('longitud_wgs84')
        lat_gmm = dd_to_gmm_str(lat_dd, True) if lat_dd is not None else "N/A"
        lon_gmm = dd_to_gmm_str(lon_dd, False) if lon_dd is not None else "N/A"

        vessel_name = obs.get('nombre_embarcacion', obs.get('name', 'N/A'))
        matricula = obs.get('matricula', 'N/A')
        nombre_patron = obs.get('nombre_patron', 'N/A')

        status_id = obs.get('estatus_categoria_id') or obs.get('status_category_id')
        status_desc = "Estatus Desconocido"
        if status_id == "outside_anp": status_desc = "Fuera del Polígono ANP"
        else:
            # Buscar el estatus en STATUS_CATEGORIES_INSIDE_ANP por su 'id' de texto
            for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                if cat_info['id'] == status_id: status_desc = cat_info['desc']; break
        notes = obs.get('notas_adicionales', '')

        patron_report_text = ""
        if nombre_patron and nombre_patron.strip().lower() != 'n/a':
            patron_report_text = f" a cargo del C. \"{nombre_patron}\" quien se identificó como patron/capitan de citada embarcacion"

        report_paragraph = document.add_paragraph()
        runner = report_paragraph.add_run(
            f"{i+1}.- El día {date_str} de {month_str} de {year_str} a las {time_str} horas "
            f"en la situación geográfica {lat_gmm} y {lon_gmm} se detectó a la embarcación "
            f"de nombre \"{vessel_name}\" matrícula {matricula}{patron_report_text}, fue reportada la situación de "
            f"{status_desc} teniendo como nota: {notes}."
        )
        font = runner.font
        font.size = Pt(10)
    
    if isinstance(filename_or_buffer, io.BytesIO):
        document.save(filename_or_buffer)
        print(f"DEBUG_WORD: Reporte Word generado exitosamente en buffer.")
    else:
        document.save(filename_or_buffer)
        print(f"DEBUG_WORD: Reporte Word '{filename_or_buffer}' generado exitosamente.")


# FUNCIÓN PARA GRAFICAR HISTORIAL O MAPA DE SESIÓN (SOLO MUESTRA EL MAPA)
def graficar_mapa_general(registros_data, titulo_mapa, es_historial_individual=False):
    """
    Genera un mapa con las observaciones de embarcaciones, límites del ANP y leyendas.
    'registros_data' debe contener objetos datetime para el timestamp.
    """
    if not registros_data:
        print(f"No hay registros para graficar para: {titulo_mapa}")
        return None, None

    fig, ax = plt.subplots(figsize=(15, 12))
    ax.set_title(titulo_mapa)

    # Dibujar polígonos base del ANP
    if anp_maritime_boundary_coords_mercator:
        x_anp_m, y_anp_m = zip(*anp_maritime_boundary_coords_mercator)
        ax.plot(x_anp_m, y_anp_m, color="blue", linewidth=1.5, zorder=2, alpha=0.6, label="Límite ANP")
        ax.fill(x_anp_m, y_anp_m, alpha=0.10, color="lightblue", zorder=1)
    
    if isla_maria_madre_coords_mercator and isla_maria_madre_polygon_geo.is_valid:
        x_imm_m, y_imm_m = zip(*isla_maria_madre_coords_mercator)
        ax.plot(x_imm_m, y_imm_m, color="darkgreen", linewidth=0.8, zorder=3, alpha=0.7, label="Isla María Madre")
        ax.fill(x_imm_m, y_imm_m, color="lightgreen", alpha=0.5, zorder=2)

    if puerto_balleto_coords_mercator and puerto_balleto_polygon_geo.is_valid:
        x_pb_m, y_pb_m = zip(*puerto_balleto_coords_mercator)
        ax.plot(x_pb_m, y_pb_m, color="saddlebrown", linewidth=0.8, zorder=4, alpha=0.7, label="Puerto Balleto")
        ax.fill(x_pb_m, y_pb_m, color="peru", alpha=0.6, zorder=3)

    for nombre_isla, data_mercator in islas_menores_data_mercator.items():
        if data_mercator["coords"]:
            points_collection_m = MultiPoint(data_mercator["coords"])
            ax.plot([p.x for p in points_collection_m.geoms], [p.y for p in points_collection_m.geoms],
                    marker=data_mercator["marker"], color=data_mercator["color"], linestyle='None',
                    markersize=6, label=nombre_isla, zorder=5, alpha=0.8)
    
    legend_elements_types_used_on_this_map = {}
    
    for i, data_point in enumerate(registros_data):
        lon_wgs84 = data_point['longitud_wgs84']
        lat_wgs84 = data_point['latitud_wgs84']
        
        x_m, y_m = transformer_geo_to_mercator.transform(lon_wgs84, lat_wgs84)
        
        v_type_id = data_point.get('tipo_embarcacion_id')
        s_cat_id = data_point.get('estatus_categoria_id')
        
        if isinstance(v_type_id, int):
            found_v_type = False
            for k_str, v_info in VESSEL_TYPES.items():
                if v_info['id'] == str(v_type_id):
                    v_type_id = k_str 
                    found_v_type = True
                    break
            if not found_v_type:
                v_type_id = 'otra' 
        elif v_type_id not in VESSEL_TYPES: 
            v_type_id = 'otra'

        marker_details = VESSEL_TYPES.get(v_type_id, DEFAULT_VESSEL_TYPE_INFO)
        
        color = STATUS_COLORS.get("unknown_status") 
        if s_cat_id == "outside_anp":
            color = STATUS_COLORS["outside_anp"]
        else:
            for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                if cat_info['id'] == s_cat_id:
                    color = STATUS_COLORS[cat_info['color_key']]
                    break
        
        if marker_details['desc'] not in legend_elements_types_used_on_this_map:
            legend_elements_types_used_on_this_map[marker_details['desc']] = plt.Line2D([0],[0], marker=marker_details['marker_char'], color='w', label=marker_details['desc'], linestyle='None', markeredgecolor='black', markerfacecolor='dimgray', markersize=7)

        ax.scatter(x_m, y_m, 
                   marker=marker_details['marker_char'], color=color, 
                   s=marker_details['size_factor'] * (0.7 if es_historial_individual else 0.5), 
                   edgecolors='black', linewidths=0.4, zorder=10, alpha=0.75)
        
        ts_obj = data_point['timestamp']
        ts_fmt = ts_obj.strftime('%y-%m-%d %H:%M')
        
        status_desc_display = "Estatus Desconocido"
        found_status_desc = False
        for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
            if cat_info['id'] == s_cat_id: status_desc_display = cat_info['desc']; found_status_desc = True; break
        if not found_status_desc and s_cat_id == "outside_anp": status_desc_display = "Fuera del Polígono ANP"
        
        if es_historial_individual or len(registros_data) < 15: 
            patron_map_text = f"C. {data_point.get('nombre_patron', 'N/A')}" if data_point.get('nombre_patron') and data_point['nombre_patron'].strip().lower() != 'n/a' else "N/A"
            annot_text = (f"{data_point['matricula']}\nPatrón: {patron_map_text}\n{ts_fmt}\n{status_desc_display}")
            ax.annotate(annot_text, (x_m, y_m),
                        xytext=(0, marker_details['size_factor'] * (0.7 if es_historial_individual else 0.5) * 0.05 + 7), 
                        textcoords='offset points', fontsize=4 if not es_historial_individual else 5.5, 
                        ha='center', va='bottom',
                        bbox=dict(boxstyle="round,pad=0.1", fc=color, alpha=0.6, ec='none'))
    
    ax.set_xlabel("X (Web Mercator)"); ax.set_ylabel("Y (Web Mercator)")
    try: cx.add_basemap(ax, crs=crs_mercator.to_string(), source=cx.providers.Esri.WorldImagery, zorder=0, alpha=0.9)
    except Exception as e_ctx: print(f"No se pudo cargar mapa base para resumen: {e_ctx}")

    handles_fig, labels_fig = ax.get_legend_handles_labels(); 
    
    map_legend_handles_display = []
    map_legend_labels_display = []
    seen_labels = set()

    for h,l in zip(handles_fig, labels_fig):
        if l not in seen_labels and l not in [info['desc'] for info in VESSEL_TYPES.values()]:
            map_legend_handles_display.append(h)
            map_legend_labels_display.append(l)
            seen_labels.add(l)

    if map_legend_handles_display:
        leg1 = ax.legend(handles=map_legend_handles_display, labels=map_legend_labels_display, 
                         fontsize='xx-small', loc='upper left', bbox_to_anchor=(1.02, 1), 
                         borderaxespad=0., title="Elementos del Mapa")
        ax.add_artist(leg1)

    if legend_elements_types_used_on_this_map:
        leg_vessel_types = ax.legend(handles=list(legend_elements_types_used_on_this_map.values()), 
                                     fontsize='xx-small', loc='center left', 
                                     bbox_to_anchor=(1.02, 0.65), 
                                     borderaxespad=0., title="Tipos Embarcación")
        ax.add_artist(leg_vessel_types)

    status_legend_handles = []
    all_status_display_ordered = [
        ("blanco", "Paso Inocente"),
        ("verde", "Turístico Autorizado"),
        ("azul_marino", "Investigación Autorizada"),
        ("amarillo", "Inconsistencias Doc. / Nav."),
        ("anaranjado", "Infracción LGPAS (Pesca/Acuacultura)"),
        ("rojo", "Delito Ambiental / Otro"),
        ("outside_anp", "Fuera del Polígono ANP"),
        ("unknown_status", "Estatus Desconocido")
    ]
    for color_key, description in all_status_display_ordered:
        status_legend_handles.append(plt.Line2D([0],[0], marker='s', color='w', label=description, 
                                                 markerfacecolor=STATUS_COLORS[color_key], markersize=7))
    if status_legend_handles:
        leg_status = ax.legend(handles=status_legend_handles, fontsize='xx-small', loc='lower left', 
                           bbox_to_anchor=(1.02, 0.05), borderaxespad=0., title="Semaforo Estatus")
        ax.add_artist(leg_status)
    
    plt.subplots_adjust(left=0.06, right=0.70, bottom=0.05, top=0.92) 
    return fig, ax 


# 7. --- RUTAS DE AUTENTICACIÓN Y APLICACIÓN ---
@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        flash('Ya has iniciado sesión.', 'info')
        return redirect(url_for('index'))

    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password'].strip()

        print(f"DEBUG LOGIN: Intento de login para usuario: '{username}'")
        user = get_user_by_username(username)

        if user:
            print(f"DEBUG LOGIN: Usuario '{username}' encontrado en DB. ID: {user.id}, Aprobado: {user.is_approved}, Rol: {user.role}")

            if not user.is_active: # Usa la propiedad is_active
                flash('Tu cuenta aún no ha sido aprobada por un administrador. Por favor, espera.', 'warning')
                return render_template('login.html')

            if check_password_hash(user.password_hash, password):
                print(f"DEBUG LOGIN: check_password_hash exitoso. Iniciando sesión para '{username}'.")
                login_user(user) 
                flash('Has iniciado sesión exitosamente.', 'success')
                next_page = request.args.get('next') 
                return redirect(next_page or url_for('index'))
            else:
                print(f"DEBUG LOGIN: check_password_hash FALLÓ para usuario '{username}'. Contraseña incorrecta.")
                flash('Nombre de usuario o contraseña incorrectos.', 'error')
        else:
            print(f"DEBUG LOGIN: Usuario '{username}' NO encontrado en DB.")
            flash('Nombre de usuario o contraseña incorrectos.', 'error')
    
    return render_template('login.html')

@app.route('/logout')
@login_required # Cualquier usuario logueado puede cerrar sesión
def logout():
    logout_user() 
    flash('Has cerrado sesión exitosamente.', 'success')
    return redirect(url_for('login')) 

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username'].strip()
        password = request.form['password'].strip()
        confirm_password = request.form['confirm_password'].strip() 

        if password != confirm_password:
            flash('Las contraseñas no coinciden.', 'error')
            return render_template('register.html')
            
        if not username or not password:
            flash('Por favor, ingresa un nombre de usuario y una contraseña.', 'error')
            return render_template('register.html')
        
        if len(password) < 6:
            flash('La contraseña debe tener al menos 6 caracteres.', 'error')
            return render_template('register.html')

        new_user = add_user(username, password)
        if new_user:
            flash(f'Tu cuenta para "{username}" ha sido creada exitosamente. Será activada por un administrador pronto.', 'success')
            return redirect(url_for('login')) 
        else:
            flash('Error al registrar el usuario. El nombre de usuario podría ya existir.', 'error')
    
    return render_template('register.html')


@app.route('/')
@viewer_required # Cualquier usuario aprobado puede ver el índice y registrar (aunque el botón de guardar observación está restringido)
def index():
    vessel_types_for_template = {k: v for k, v in VESSEL_TYPES.items()}
    status_categories_for_template = {k: v for k, v in STATUS_CATEGORIES_INSIDE_ANP.items()}
    status_categories_for_template['outside_anp'] = {"id": "outside_anp", "desc": "Fuera del Polígono ANP"}
    return render_template('index.html', vessel_types=vessel_types_for_template, status_categories=status_categories_for_template)

@app.route('/add_observation', methods=['POST'])
@editor_required # Solo editores y administradores pueden añadir observaciones
def add_observation():
    matricula = request.form['matricula'].strip().upper()
    nombre_embarcacion = request.form['nombre_embarcacion'].strip() or f"Emb. {matricula}"
    nombre_patron = request.form['nombre_patron'].strip() or "N/A"
    
    timestamp_str_input = request.form.get('timestamp', '').strip() 
    
    if not timestamp_str_input: 
        avist_dt_obj = datetime.datetime.now()
    else:
        try:
            avist_dt_obj = datetime.datetime.strptime(timestamp_str_input, '%Y-%m-%dT%H:%M')
        except ValueError:
            flash("Error: Formato de fecha/hora incorrecto. Use AAAA-MM-DD HH:MM o el selector de fecha/hora.", 'error')
            return redirect(url_for('index'))

    print(f"DEBUG: Fecha/hora final para DB (objeto datetime): {avist_dt_obj}")


    coord_format = request.form['coord_format']
    lat_dd, lon_dd = None, None

    try:
        if coord_format == 'gms':
            lat_g = request.form['lat_g']; lat_m = request.form['lat_m']; lat_s = request.form['lat_s']; lat_h = request.form['lat_h'].upper()
            lon_g = request.form['lon_g']; lon_m = request.form['lon_m']; lon_s = request.form['lon_s']; lon_h = request.form['lon_h'].upper()
            lat_dd = gms_to_dd(lat_g, lat_m, lat_s, lat_h)
            lon_dd = gms_to_dd(lon_g, lon_m, lon_s, lon_h)
        elif coord_format == 'dd':
            lat_dd = float(request.form['lat_dd'])
            lon_dd = float(request.form['lon_dd'])
        elif coord_format == 'utm':
            utm_x = float(request.form['utm_x'])
            utm_y = float(request.form['utm_y'])
            lon_dd, lat_dd = transformer_utm_to_geo.transform(utm_x, utm_y)
        elif coord_format == 'gdm':
            lat_g_gdm = request.form['lat_g_gdm']; lat_m_gdm = request.form['lat_m_gdm']; lat_h_gdm = request.form['lat_h_gdm'].upper()
            lon_g_gdm = request.form['lon_g_gdm']; lon_m_gdm = request.form['lon_m_gdm']; lon_h_gdm = request.form['lon_h_gdm'].upper()
            lat_dd = gdm_to_dd(lat_g_gdm, lat_m_gdm, lat_h_gdm)
            lon_dd = gdm_to_dd(lon_g_gdm, lon_m_gdm, lon_h_gdm)
    except ValueError as e:
        flash(f"Error en el formato de coordenadas: {e}", 'error')
        return redirect(url_for('index'))

    curr_pt_geo = Point(lon_dd, lat_dd)
    is_in_anp = anp_maritime_polygon_geo.intersects(curr_pt_geo) if not anp_maritime_polygon_geo.is_empty else False

    status_category_id = "outside_anp"
    if is_in_anp:
        try:
            status_category_key_str = request.form['status_category'] 
            # Se busca el ID de texto directamente, ya que el select en index.html envía el ID de texto
            found = False
            for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                if cat_info['id'] == status_category_key_str:
                    status_category_id = status_category_key_str
                    found = True
                    break
            if not found and status_category_key_str != 'outside_anp': # Asegurar que 'outside_anp' se maneja
                flash("Error: Estatus de categoría inválido o formato incorrecto.", 'error')
                return redirect(url_for('index'))
            elif status_category_key_str == 'outside_anp':
                 status_category_id = 'outside_anp' # Asegura que se guarda correctamente si se selecciona manualmente
            
        except (ValueError, KeyError):
            flash("Error: Estatus de categoría inválido o formato incorrecto.", 'error')
            return redirect(url_for('index'))
    
    notas = request.form.get('notas_adicionales', '')
    vessel_type_id = request.form['vessel_type'] 

    agregar_observacion_db(matricula, nombre_embarcacion, avist_dt_obj, lat_dd, lon_dd, vessel_type_id, status_category_id, notas, nombre_patron)
    flash(f"Observación para matrícula '{matricula}' guardada exitosamente.", 'success') 
    return redirect(url_for('history', matricula=matricula))


@app.route('/history')
@viewer_required # Cualquier usuario aprobado puede ver el historial
def history():
    matricula = request.args.get('matricula', '')
    nombre_embarcacion = request.args.get('nombre_embarcacion', '')
    nombre_patron = request.args.get('nombre_patron', '')
    
    observations_raw = [] 
    message = None 
    
    if matricula:
        observations_raw = buscar_historial_embarcacion(matricula)
        if not observations_raw:
            message = f"No se encontraron observaciones para la matrícula '{matricula}'."
    elif nombre_embarcacion or nombre_patron:
        observations_raw = buscar_por_nombre_o_patron(nombre_embarcacion, nombre_patron)
        if not observations_raw:
            message = "No se encontraron observaciones para el nombre de embarcación o patrón proporcionado."
    
    if not (matricula or nombre_embarcacion or nombre_patron) and not observations_raw:
        message = "Ingrese un criterio de búsqueda (matrícula, nombre de embarcación o patrón)."

    fig, ax = graficar_mapa_general(observations_raw, f"Historial para {matricula or nombre_embarcacion or nombre_patron}", es_historial_individual=True)
    
    img_buffer = io.BytesIO()
    if fig:
        fig.savefig(img_buffer, format='png', bbox_inches='tight', pad_inches=0.1)
        plt.close(fig) 
    img_buffer.seek(0)
    img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
    del fig 

    observations_for_template = []
    for obs in observations_raw:
        temp_obs = obs.copy() 
        if isinstance(temp_obs['timestamp'], datetime.datetime):
            temp_obs['timestamp'] = temp_obs['timestamp'].strftime('%Y-%m-%d %H:%M:%S')
        observations_for_template.append(temp_obs)

    vessel_types_for_template = {k: v for k, v in VESSEL_TYPES.items()}
    status_categories_for_template = {}
    for k_int, v_dict in STATUS_CATEGORIES_INSIDE_ANP.items():
        status_categories_for_template[v_dict['id']] = v_dict 
    status_categories_for_template['outside_anp'] = {"id": "outside_anp", "desc": "Fuera del Polígono ANP"}

    return render_template('history.html', 
                           observations=observations_for_template, 
                           map_image=img_base64, 
                           matricula=matricula, 
                           nombre_embarcacion=nombre_embarcacion, 
                           nombre_patron=nombre_patron, 
                           message=message, 
                           vessel_types=vessel_types_for_template, 
                           status_categories=status_categories_for_template)


# NUEVA RUTA: Editar observación (GET para mostrar formulario)
@app.route('/edit_observation/<int:obs_id>', methods=['GET'])
@editor_required # Solo editores y administradores pueden editar observaciones
def edit_observation(obs_id):
    observation = get_observacion_by_id(obs_id)
    if not observation:
        flash("Observación no encontrada.", 'error')
        return redirect(url_for('history'))
    
    # Formatear el timestamp para el input datetime-local
    if isinstance(observation['timestamp'], datetime.datetime):
        observation['timestamp_formatted'] = observation['timestamp'].strftime('%Y-%m-%dT%H:%M')
    else:
        observation['timestamp_formatted'] = '' # Fallback si no es un datetime object

    vessel_types_for_template = {k: v for k, v in VESSEL_TYPES.items()}
    status_categories_for_template = {k: v for k, v in STATUS_CATEGORIES_INSIDE_ANP.items()}
    status_categories_for_template['outside_anp'] = {"id": "outside_anp", "desc": "Fuera del Polígono ANP"}

    return render_template('edit_observation.html', 
                           observation=observation, 
                           vessel_types=vessel_types_for_template, 
                           status_categories=status_categories_for_template)

# NUEVA RUTA: Actualizar observación (POST para procesar el formulario)
@app.route('/update_observation/<int:obs_id>', methods=['POST'])
@editor_required # Solo editores y administradores pueden actualizar observaciones
def update_observation(obs_id):
    matricula = request.form['matricula'].strip().upper()
    nombre_embarcacion = request.form['nombre_embarcacion'].strip() or f"Emb. {matricula}"
    nombre_patron = request.form['nombre_patron'].strip() or "N/A"
    
    timestamp_str_input = request.form.get('timestamp', '').strip() 
    
    if not timestamp_str_input: 
        flash("Error: La fecha y hora de avistamiento son obligatorias.", 'error')
        return redirect(url_for('edit_observation', obs_id=obs_id))
    
    try:
        avist_dt_obj = datetime.datetime.strptime(timestamp_str_input, '%Y-%m-%dT%H:%M')
    except ValueError:
        flash("Error: Formato de fecha/hora incorrecto. Use AAAA-MM-DD HH:MM o el selector de fecha/hora.", 'error')
        return redirect(url_for('edit_observation', obs_id=obs_id))

    coord_format = request.form['coord_format']
    lat_dd, lon_dd = None, None

    try:
        if coord_format == 'gms':
            lat_g = request.form['lat_g']; lat_m = request.form['lat_m']; lat_s = request.form['lat_s']; lat_h = request.form['lat_h'].upper()
            lon_g = request.form['lon_g']; lon_m = request.form['lon_m']; lon_s = request.form['lon_s']; lon_h = request.form['lon_h'].upper()
            lat_dd = gms_to_dd(lat_g, lat_m, lat_s, lat_h)
            lon_dd = gms_to_dd(lon_g, lon_m, lon_s, lon_h)
        elif coord_format == 'dd':
            lat_dd = float(request.form['lat_dd'])
            lon_dd = float(request.form['lon_dd'])
        elif coord_format == 'utm':
            utm_x = float(request.form['utm_x'])
            utm_y = float(request.form['utm_y'])
            lon_dd, lat_dd = transformer_utm_to_geo.transform(utm_x, utm_y)
        elif coord_format == 'gdm':
            lat_g_gdm = request.form['lat_g_gdm']; lat_m_gdm = request.form['lat_m_gdm']; lat_h_gdm = request.form['lat_h_gdm'].upper()
            lon_g_gdm = request.form['lon_g_gdm']; lon_m_gdm = request.form['lon_m_gdm']; lon_h_gdm = request.form['lon_h_gdm'].upper()
            lat_dd = gdm_to_dd(lat_g_gdm, lat_m_gdm, lat_h_gdm)
            lon_dd = gdm_to_dd(lon_g_gdm, lon_m_gdm, lon_h_gdm)
    except ValueError as e:
        flash(f"Error en el formato de coordenadas: {e}", 'error')
        return redirect(url_for('edit_observation', obs_id=obs_id))

    curr_pt_geo = Point(lon_dd, lat_dd)
    is_in_anp = anp_maritime_polygon_geo.intersects(curr_pt_geo) if not anp_maritime_polygon_geo.is_empty else False

    status_category_id = "outside_anp"
    if is_in_anp:
        try:
            status_category_key_str = request.form['status_category'] 
            found = False
            for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                if cat_info['id'] == status_category_key_str:
                    status_category_id = status_category_key_str
                    found = True
                    break
            if not found and status_category_key_str != 'outside_anp':
                flash("Error: Estatus de categoría inválido o formato incorrecto.", 'error')
                return redirect(url_for('edit_observation', obs_id=obs_id))
            elif status_category_key_str == 'outside_anp':
                 status_category_id = 'outside_anp'
            
        except (ValueError, KeyError):
            flash("Error: Estatus de categoría inválido o formato incorrecto.", 'error')
            return redirect(url_for('edit_observation', obs_id=obs_id))
    
    notas = request.form.get('notas_adicionales', '')
    vessel_type_id = request.form['vessel_type'] 

    if update_observacion_db(obs_id, matricula, nombre_embarcacion, avist_dt_obj, lat_dd, lon_dd, vessel_type_id, status_category_id, notas, nombre_patron):
        flash(f"Observación ID {obs_id} actualizada exitosamente.", 'success') 
        return redirect(url_for('history', matricula=matricula))
    else:
        flash(f"Error al actualizar la observación ID {obs_id}.", 'error')
        return redirect(url_for('edit_observation', obs_id=obs_id))


@app.route('/download_report/<matricula>')
@viewer_required # Cualquier usuario aprobado puede descargar reportes de historial individual
def download_report(matricula):
    observations_raw = buscar_historial_embarcacion(matricula) 
    if not observations_raw:
        flash("No hay datos para generar el reporte.", 'error')
        return redirect(url_for('history', matricula=matricula))
    
    observations_for_report = []
    for obs in observations_raw:
        temp_obs = obs.copy()
        observations_for_report.append(temp_obs)

    fig, ax = graficar_mapa_general(observations_for_report, f"Historial para {matricula}", es_historial_individual=True)
    
    doc_buffer = io.BytesIO()
    
    if fig:
        try:
            generar_reporte_word(fig, observations_for_report, f"Historial de Inspección: {matricula}", filename_or_buffer=doc_buffer)
        except Exception as e:
            print(f"Error generating Word report for download: {e}")
            flash("Error al generar el reporte de Word.", 'error')
            return redirect(url_for('history', matricula=matricula))
    else:
        flash("Error: No se pudo generar la imagen del mapa para el reporte.", 'error')
        return redirect(url_for('history', matricula=matricula))

    doc_buffer.seek(0)
    return send_file(doc_buffer, download_name=f"reporte_historial_{matricula}.docx", as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/download_all_csv')
@viewer_required # Cualquier usuario aprobado puede descargar todos los CSVs
def download_all_csv():
    conn = conectar_db()
    if not conn:
        flash("Error: No se pudo conectar a la base de datos para exportar CSV.", 'error')
        return redirect(url_for('index'))
    
    cursor = conn.cursor()
    try:
        cursor.execute("""
            SELECT id, matricula, nombre_embarcacion, timestamp, latitud_wgs84, longitud_wgs84,
                   tipo_embarcacion_id, estatus_categoria_id, notas_adicionales, nombre_patron
            FROM observaciones_embarcaciones ORDER BY timestamp DESC
        """)
        
        column_names = [desc[0] for desc in cursor.description]
        rows_tuples = cursor.fetchall()
        
        csv_buffer = io.StringIO()
        csv_writer = csv.writer(csv_buffer)
        
        csv_writer.writerow(column_names)
        
        for row in rows_tuples:
            row_list = list(row)
            try:
                timestamp_index = column_names.index('timestamp')
                timestamp_obj = row_list[timestamp_index]
                if isinstance(timestamp_obj, datetime.datetime):
                    row_list[timestamp_index] = timestamp_obj.strftime('%Y-%m-%d %H:%M:%S')
            except ValueError:
                pass 
            csv_writer.writerow(row_list)
        
        csv_buffer.seek(0)
        
        return send_file(io.BytesIO(csv_buffer.getvalue().encode('utf-8')),
                         mimetype='text/csv',
                         download_name='observaciones_anp_todas.csv',
                         as_attachment=True)
    except psycopg2.Error as e:
        print(f"Error al exportar datos a CSV: {e}")
        flash("Error al exportar datos a CSV.", 'error')
        return redirect(url_for('index'))
    finally:
        cursor.close()
        conn.close()


@app.route('/summary_options')
@viewer_required # Cualquier usuario aprobado puede ver opciones de resumen
def summary_options():
    current_year = datetime.datetime.now().year
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Mexico.1252')
        except locale.Error:
            print("ADVERTENCIA: No se pudo configurar el locale español en summary_options.")

    # Pasar todas las categorías de estatus, incluyendo "Fuera del ANP"
    all_status_categories = {v['id']: v['desc'] for v in STATUS_CATEGORIES_INSIDE_ANP.values()}
    all_status_categories['outside_anp'] = "Fuera del Polígono ANP" # Añadir el estatus fuera del ANP

    return render_template('summary_options.html', 
                           current_year=current_year, 
                           calendar=calendar, 
                           datetime=datetime,
                           all_status_categories=all_status_categories)


@app.route('/summary_report', methods=['GET'])
@viewer_required # Cualquier usuario aprobado puede ver reportes de resumen
def summary_report():
    report_type = request.args.get('report_type')
    requested_year = request.args.get('year', type=int)
    requested_month = request.args.get('month', type=int)
    week_num_option = request.args.get('week_num_option', type=int)
    status_category_filter = request.args.get('status_category', '').strip() # NUEVO: Obtener el filtro de estatus

    current_server_year = datetime.datetime.now().year
    current_server_month = datetime.datetime.now().month

    year = requested_year if requested_year is not None and requested_year != 0 else current_server_year
    month = requested_month if requested_month is not None and requested_month != 0 else current_server_month

    print(f"DEBUG_REQUEST (Summary Report): requested_year={requested_year}, requested_month={requested_month}, status_category_filter='{status_category_filter}'")
    print(f"DEBUG_USED (Summary Report): year={year}, month={month}")


    start_date_obj, end_date_obj, map_title_suffix = None, None, ""
    message = None
    
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Mexico.1252')
        except locale.Error:
            print("ADVERTENCIA: No se pudo configurar el locale español en summary_report.")

    try:
        if report_type == "weekly":
            if not week_num_option: 
                flash("Error: El número de semana es obligatorio para el resumen semanal.", 'error')
                return redirect(url_for('summary_options'))

            first_day_of_month = datetime.date(year, month, 1)
            
            # Calcular el primer domingo de la semana del mes
            # weekday() devuelve 0 para lunes, 6 para domingo. (first_day_of_month.weekday() + 1) % 7 da 1 para lunes, 0 para domingo.
            # Esto calcula los días a restar para llegar al domingo anterior o el mismo domingo si first_day_of_month es domingo.
            first_sunday_of_relevant_period_date = first_day_of_month - datetime.timedelta(days=(first_day_of_month.weekday() + 1) % 7)

            start_week_date = first_sunday_of_relevant_period_date + datetime.timedelta(weeks=week_num_option - 1)
            end_week_date = start_week_date + datetime.timedelta(days=6)

            start_date_obj = datetime.datetime.combine(start_week_date, datetime.time.min).replace(microsecond=0)
            end_date_obj = datetime.datetime.combine(end_week_date, datetime.time.max).replace(microsecond=999999)
            
            print(f"DEBUG_FINAL_RANGE (Weekly Report): start_date_obj={start_date_obj}, end_date_obj={end_date_obj}")
            
            display_start_date_title = max(start_week_date, first_day_of_month)
            display_end_date_title = min(end_week_date, datetime.date(year, month, calendar.monthrange(year, month)[1]))
            map_title_suffix = (f"Semana del {display_start_date_title.strftime('%d de %B')} "
                                f"al {display_end_date_title.strftime('%d de %B de %Y')}")

        elif report_type == "monthly":
            _, num_days = calendar.monthrange(year, month)
            
            start_date_obj = datetime.datetime(year, month, 1, 0, 0, 0, 0)
            end_date_obj = datetime.datetime(year, month, num_days, 23, 59, 59, 999999)
            
            print(f"DEBUG_FINAL_RANGE (Monthly Report): start_date_obj={start_date_obj}, end_date_obj={end_date_obj}")

            map_title_suffix = f"{datetime.date(year, month, 1).strftime('%B').capitalize()} {year}"
        
        elif report_type == "annual":
            start_date_obj = datetime.datetime(year, 1, 1, 0, 0, 0, 0)
            end_date_obj = datetime.datetime(year, 12, 31, 23, 59, 59, 999999)
            
            print(f"DEBUG_FINAL_RANGE (Annual Report): start_date_obj={start_date_obj}, end_date_obj={end_date_obj}")

            map_title_suffix = f"Año {year}"

        elif report_type == "total":
            start_date_obj, end_date_obj = None, None 
            map_title_suffix = "Todas las Inspecciones (Neto)"
            print(f"DEBUG_FINAL_RANGE (Total Report): No hay filtro de fechas.")
        else:
            flash("Error: Tipo de reporte no válido.", 'error')
            return redirect(url_for('summary_options'))

    except ValueError as ve: 
        flash(f"Error en la entrada de fecha para resumen: {ve}", 'error')
        print(f"ERROR: ValueError en summary_report: {ve}")
        return redirect(url_for('summary_options'))
    except Exception as e_date: 
        flash(f"Error al procesar fechas para resumen: {e_date}", 'error')
        print(f"ERROR: Excepción inesperada en summary_report: {e_date}")
        return redirect(url_for('summary_options'))

    # Pasar el filtro de estatus a la función de obtención de observaciones
    observations_raw = obtener_observaciones_filtradas(start_date_obj, end_date_obj, status_category_filter if status_category_filter != "" else None)

    # Si se aplicó un filtro de estatus, añadirlo al título
    if status_category_filter:
        status_desc = "Desconocido"
        if status_category_filter == "outside_anp":
            status_desc = "Fuera del Polígono ANP"
        else:
            for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                if cat_info['id'] == status_category_filter:
                    status_desc = cat_info['desc']
                    break
        map_title_suffix += f" (Estatus: {status_desc})"


    if not observations_raw:
        message = f"No se encontraron observaciones para el periodo: {map_title_suffix}."

    fig, ax = graficar_mapa_general(observations_raw, f"Resumen Inspecciones: {map_title_suffix}", es_historial_individual=False)
    
    img_buffer = io.BytesIO()
    if fig:
        fig.savefig(img_buffer, format='png', bbox_inches='tight', pad_inches=0.1)
        plt.close(fig)
    img_buffer.seek(0)
    img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
    del fig 

    observations_for_template = []
    for obs in observations_raw:
        temp_obs = obs.copy() 
        if isinstance(temp_obs['timestamp'], datetime.datetime):
            temp_obs['timestamp'] = temp_obs['timestamp'].strftime('%Y-%m-%d %H:%M:%S')
        observations_for_template.append(temp_obs)

    vessel_types_for_template = {k: v for k, v in VESSEL_TYPES.items()}
    status_categories_for_template = {}
    for k_int, v_dict in STATUS_CATEGORIES_INSIDE_ANP.items():
        status_categories_for_template[v_dict['id']] = v_dict
    status_categories_for_template['outside_anp'] = {"id": "outside_anp", "desc": "Fuera del Polígono ANP"}

    return render_template('summary_report.html', 
                           observations=observations_for_template, 
                           map_image=img_base64, 
                           map_title=f"Resumen Inspecciones: {map_title_suffix}", 
                           message=message,
                           vessel_types=vessel_types_for_template, 
                           status_categories=status_categories_for_template)


@app.route('/download_summary_report/<report_type>')
@viewer_required # Cualquier usuario aprobado puede descargar reportes de resumen DOCX
def download_summary_report(report_type):
    requested_year = request.args.get('year', type=int)
    requested_month = request.args.get('month', type=int)
    week_num_option = request.args.get('week_num_option', type=int)
    status_category_filter = request.args.get('status_category', '').strip() # NUEVO: Obtener el filtro de estatus

    current_server_year = datetime.datetime.now().year
    current_server_month = datetime.datetime.now().month

    year = requested_year if requested_year is not None and requested_year != 0 else current_server_year
    month = requested_month if requested_month is not None and requested_month != 0 else current_server_month

    print(f"DEBUG_REQUEST (Download Summary): requested_year={requested_year}, requested_month={requested_month}, status_category_filter='{status_category_filter}'")
    print(f"DEBUG_USED (Download Summary): year={year}, month={month}")

    start_date_obj, end_date_obj, map_title_suffix = None, None, ""
    
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Mexico.1252')
        except locale.Error:
            print("ADVERTENCIA: No se pudo configurar el locale español para descarga de reporte.")

    try:
        if report_type == "weekly":
            if not week_num_option: 
                flash("Error: El número de semana es obligatorio para reporte semanal.", 'error')
                return redirect(url_for('summary_options'))
            
            first_day_of_month = datetime.date(year, month, 1)
            first_sunday_of_relevant_period_date = first_day_of_month - datetime.timedelta(days=(first_day_of_month.weekday() + 1) % 7)

            start_week_date = first_sunday_of_relevant_period_date + datetime.timedelta(weeks=week_num_option - 1)
            end_week_date = start_week_date + datetime.timedelta(days=6)

            start_date_obj = datetime.datetime.combine(start_week_date, datetime.time.min).replace(microsecond=0)
            end_date_obj = datetime.datetime.combine(end_week_date, datetime.time.max).replace(microsecond=999999)
            
            print(f"DEBUG_FINAL_RANGE (Weekly Download): start_date_obj={start_date_obj}, end_date_obj={end_date_obj}")

            display_start_date_title = max(start_week_date, first_day_of_month)
            display_end_date_title = min(end_week_date, datetime.date(year, month, calendar.monthrange(year, month)[1]))
            map_title_suffix = (f"Semana del {display_start_date_title.strftime('%d de %B')} "
                                f"al {display_end_date_title.strftime('%d de %B de %Y')}")

        elif report_type == "monthly":
            _, num_days = calendar.monthrange(year, month)
            start_date_obj = datetime.datetime(year, month, 1, 0, 0, 0, 0)
            end_date_obj = datetime.datetime(year, month, num_days, 23, 59, 59, 999999)
            
            print(f"DEBUG_FINAL_RANGE (Monthly Download): start_date_obj={start_date_obj}, end_date_obj={end_date_obj}")

            map_title_suffix = f"{datetime.date(year, month, 1).strftime('%B').capitalize()} {year}"
        
        elif report_type == "annual":
            start_date_obj = datetime.datetime(year, 1, 1, 0, 0, 0, 0)
            end_date_obj = datetime.datetime(year, 12, 31, 23, 59, 59, 999999)
            
            print(f"DEBUG_FINAL_RANGE (Annual Download): start_date_obj={start_date_obj}, end_date_obj={end_date_obj}")

            map_title_suffix = f"Año {year}"

        elif report_type == "total":
            start_date_obj, end_date_obj = None, None 
            map_title_suffix = "Todas las Inspecciones (Neto)"
            print(f"DEBUG_FINAL_RANGE (Total Download): No hay filtro de fechas.")
        else:
            flash("Error: Tipo de reporte no válido.", 'error')
            return redirect(url_for('summary_options'))

    except ValueError as ve: 
        flash(f"Error en la entrada de fecha para reporte Word: {ve}", 'error')
        print(f"ERROR: ValueError en download_summary_report: {ve}")
        return redirect(url_for('summary_options'))
    except Exception as e_date: 
        flash(f"Error al procesar fechas para reporte Word: {e_date}", 'error')
        print(f"ERROR: Excepción inesperada en download_summary_report: {e_date}")
        return redirect(url_for('summary_options'))

    # Pasar el filtro de estatus a la función de obtención de observaciones
    observations_raw = obtener_observaciones_filtradas(start_date_obj, end_date_obj, status_category_filter if status_category_filter != "" else None)

    # Si se aplicó un filtro de estatus, añadirlo al título del documento
    if status_category_filter:
        status_desc = "Desconocido"
        if status_category_filter == "outside_anp":
            status_desc = "Fuera del Polígono ANP"
        else:
            for cat_info in STATUS_CATEGORIES_INSIDE_ANP.values():
                if cat_info['id'] == status_category_filter:
                    status_desc = cat_info['desc']
                    break
        map_title_suffix += f" (Estatus: {status_desc})"

    if not observations_raw:
        flash(f"No hay datos para generar el reporte DOCX para el periodo: {map_title_suffix}.", 'error')
        return redirect(url_for('summary_options'))

    observations_for_report = []
    for obs in observations_raw:
        temp_obs = obs.copy()
        observations_for_report.append(temp_obs)

    fig, ax = graficar_mapa_general(observations_for_report, f"Resumen Inspecciones: {map_title_suffix}", es_historial_individual=False)
    
    doc_buffer = io.BytesIO()
    
    if fig:
        try:
            generar_reporte_word(fig, observations_for_report, f"Resumen de Inspección: {map_title_suffix}", filename_or_buffer=doc_buffer)
        except Exception as e:
            print(f"Error generating Word report for download: {e}")
            flash("Error al generar el reporte de Word.", 'error')
            return redirect(url_for('summary_options'))
    else:
        flash("Error: No se pudo generar la imagen del mapa para el reporte DOCX.", 'error')
        return redirect(url_for('summary_options'))

    doc_buffer.seek(0)
    filename = f"resumen_{report_type}"
    if year: filename += f"_{year}"
    if month: filename += f"_{month}"
    if week_num_option: filename += f"_{week_num_option}"
    if status_category_filter: filename += f"_{status_category_filter}" # Añadir estatus al nombre del archivo
    filename += ".docx"

    return send_file(doc_buffer, download_name=filename, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# NUEVA RUTA: Descargar CSV de resumen filtrado
@app.route('/download_summary_csv')
@viewer_required # Cualquier usuario aprobado puede descargar CSVs de resumen
def download_summary_csv():
    report_type = request.args.get('report_type')
    requested_year = request.args.get('year', type=int)
    requested_month = request.args.get('month', type=int)
    week_num_option = request.args.get('week_num_option', type=int)
    status_category_filter = request.args.get('status_category', '').strip()

    current_server_year = datetime.datetime.now().year
    current_server_month = datetime.datetime.now().month

    year = requested_year if requested_year is not None and requested_year != 0 else current_server_year
    month = requested_month if requested_month is not None and requested_month != 0 else current_server_month

    start_date_obj, end_date_obj = None, None
    filename_suffix = ""

    try:
        if report_type == "weekly":
            if not week_num_option: raise ValueError("El número de semana es obligatorio para el resumen semanal.")
            first_day_of_month = datetime.date(year, month, 1)
            first_sunday_of_relevant_period_date = first_day_of_month - datetime.timedelta(days=(first_day_of_month.weekday() + 1) % 7)
            start_week_date = first_sunday_of_relevant_period_date + datetime.timedelta(weeks=week_num_option - 1)
            end_week_date = start_week_date + datetime.timedelta(days=6)
            start_date_obj = datetime.datetime.combine(start_week_date, datetime.time.min).replace(microsecond=0)
            end_date_obj = datetime.datetime.combine(end_week_date, datetime.time.max).replace(microsecond=999999)
            filename_suffix = f"_semanal_{year}_{month}_{week_num_option}"
        elif report_type == "monthly":
            _, num_days = calendar.monthrange(year, month)
            start_date_obj = datetime.datetime(year, month, 1, 0, 0, 0, 0)
            end_date_obj = datetime.datetime(year, month, num_days, 23, 59, 59, 999999)
            filename_suffix = f"_mensual_{year}_{month}"
        elif report_type == "annual":
            start_date_obj = datetime.datetime(year, 1, 1, 0, 0, 0, 0)
            end_date_obj = datetime.datetime(year, 12, 31, 23, 59, 59, 999999)
            filename_suffix = f"_anual_{year}"
        elif report_type == "total":
            filename_suffix = "_total"
        else:
            raise ValueError("Tipo de reporte no válido.")
    except ValueError as e:
        flash(f"Error en la entrada de fecha para CSV de resumen: {e}", 'error')
        return redirect(url_for('summary_options'))
    except Exception as e:
        flash(f"Error al procesar fechas para CSV de resumen: {e}", 'error')
        return redirect(url_for('summary_options'))

    observations_raw = obtener_observaciones_filtradas(start_date_obj, end_date_obj, status_category_filter if status_category_filter != "" else None)

    if not observations_raw:
        flash("No hay datos para generar el CSV de resumen filtrado.", 'error')
        return redirect(url_for('summary_options'))

    # Añadir filtro de estatus al nombre del archivo si aplica
    if status_category_filter:
        filename_suffix += f"_{status_category_filter}"

    csv_buffer = io.StringIO()
    csv_writer = csv.writer(csv_buffer)
    
    # Escribir encabezados
    column_names = ['id', 'matricula', 'nombre_embarcacion', 'timestamp', 'latitud_wgs84', 
                    'longitud_wgs84', 'tipo_embarcacion_id', 'estatus_categoria_id', 'notas_adicionales', 'nombre_patron']
    csv_writer.writerow(column_names)
    
    # Escribir datos
    for obs in observations_raw:
        row = [
            obs.get('id'),
            obs.get('matricula'),
            obs.get('nombre_embarcacion'),
            obs.get('timestamp').strftime('%Y-%m-%d %H:%M:%S') if isinstance(obs.get('timestamp'), datetime.datetime) else '',
            obs.get('latitud_wgs84'),
            obs.get('longitud_wgs84'),
            obs.get('tipo_embarcacion_id'),
            obs.get('estatus_categoria_id'),
            obs.get('notas_adicionales'),
            obs.get('nombre_patron')
        ]
        csv_writer.writerow(row)
    
    csv_buffer.seek(0)
    
    return send_file(io.BytesIO(csv_buffer.getvalue().encode('utf-8')),
                     mimetype='text/csv',
                     download_name=f'resumen_observaciones{filename_suffix}.csv',
                     as_attachment=True)


@app.route('/delete_observation/<int:obs_id>', methods=['POST'])
@editor_required # Solo editores y administradores pueden eliminar observaciones
def delete_observation(obs_id):
    if eliminar_observacion_db(obs_id):
        flash(f"Observación con ID {obs_id} eliminada exitosamente.", 'success')
        matricula_prev = request.form.get('matricula_redirect') 
        if matricula_prev:
            return redirect(url_for('history', matricula=matricula_prev))
        else:
            return redirect(request.referrer or url_for('index'))
    else:
        flash("Error al eliminar la observación.", 'error')
        return "Error al eliminar la observación.", 500

@app.route('/upload_csv_to_db', methods=['GET', 'POST'])
@editor_required # Solo editores y administradores pueden subir CSVs
def upload_csv_to_db():
    if request.method == 'POST':
        if 'csv_file' not in request.files:
            flash('No se seleccionó ningún archivo.', 'error')
            return redirect(request.url)
        
        file = request.files['csv_file']
        if file.filename == '':
            flash('No se seleccionó ningún archivo.', 'error')
            return redirect(request.url)
        
        if file and file.filename.endswith('.csv'):
            try:
                stream = io.StringIO(file.stream.read().decode("UTF8"))
                reader = csv.DictReader(stream)
                
                conn = conectar_db()
                if not conn:
                    flash('Error: No se pudo conectar a la base de datos.', 'error')
                    return render_template('upload_csv.html')

                cursor = conn.cursor()
                total_inserted = 0
                total_skipped = 0

                for row_num, row_data_from_csv in enumerate(reader):
                    try:
                        insert_sql = """
                        INSERT INTO observaciones_embarcaciones (
                            matricula, nombre_embarcacion, timestamp, latitud_wgs84, longitud_wgs84,
                            tipo_embarcacion_id, estatus_categoria_id, notas_adicionales, nombre_patron
                        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                        ON CONFLICT (matricula, timestamp) DO NOTHING; 
                        """
                        
                        matricula_val = row_data_from_csv.get('matricula')
                        nombre_embarcacion_val = row_data_from_csv.get('nombre_embarcacion')
                        timestamp_val_raw = row_data_from_csv.get('timestamp')

                        timestamp_dt_obj = None
                        if timestamp_val_raw:
                            try:
                                timestamp_dt_obj = datetime.datetime.strptime(timestamp_val_raw, '%Y-%m-%d %H:%M:%S')
                            except ValueError:
                                try:
                                    timestamp_dt_obj = datetime.datetime.strptime(timestamp_val_raw, '%Y-%m-%d %H:%M')
                                except ValueError:
                                    try:
                                        timestamp_dt_obj = datetime.datetime.fromisoformat(timestamp_val_raw.replace('Z', '+00:00'))
                                    except ValueError:
                                        print(f"ADVERTENCIA (CSV Import): Timestamp no reconocido en fila {row_num + 1}: '{timestamp_val_raw}'. Esta fila podría fallar la inserción.")
                                        timestamp_dt_obj = None 
                        
                        if timestamp_dt_obj is None:
                            raise ValueError(f"No se pudo parsear el timestamp '{timestamp_val_raw}' a un formato de fecha/hora válido.")

                        latitud_val = None
                        if row_data_from_csv.get('latitud_wgs84'):
                            try:
                                latitud_val = float(row_data_from_csv['latitud_wgs84'])
                            except ValueError:
                                print(f"ADVERTENCIA (CSV Import): Latitud no válida en fila {row_num + 1}: '{row_data_from_csv['latitud_wgs84']}'")
                        
                        longitud_val = None
                        if row_data_from_csv.get('longitud_wgs84'):
                            try:
                                longitud_val = float(row_data_from_csv['longitud_wgs84'])
                            except ValueError:
                                print(f"ADVERTENCIA (CSV Import): Longitud no válida en fila {row_num + 1}: '{row_data_from_csv['longitud_wgs84']}'")

                        tipo_embarcacion_id_val = row_data_from_csv.get('tipo_embarcacion_id')
                        estatus_categoria_id_val = row_data_from_csv.get('estatus_categoria_id')
                        notas_adicionales_val = row_data_from_csv.get('notas_adicionales')
                        nombre_patron_val = row_data_from_csv.get('nombre_patron')

                        cursor.execute(insert_sql, (
                            matricula_val,
                            nombre_embarcacion_val,
                            timestamp_dt_obj, 
                            latitud_val,
                            longitud_val,
                            tipo_embarcacion_id_val,
                            estatus_categoria_id_val,
                            notas_adicionales_val,
                            nombre_patron_val
                        ))
                        if cursor.rowcount > 0: 
                            total_inserted += 1
                        else: 
                            total_skipped += 1

                    except KeyError as e:
                        print(f"ADVERTENCIA (CSV Import): Columna '{e}' no encontrada en CSV para la fila {row_num + 1}. Fila omitida.")
                        total_skipped += 1
                    except (ValueError, TypeError) as e:
                        print(f"ADVERTENCIA (CSV Import): Error de conversión de tipo en fila {row_num + 1} (Matrícula: {row_data_from_csv.get('matricula', 'N/A')}): {e}. Fila omitida.")
                        total_skipped += 1
                    except psycopg2.Error as e: 
                        print(f"ERROR DB (CSV Import) al procesar fila {row_num + 1} (Matrícula: {row_data_from_csv.get('matricula', 'N/A')}): {e}. Fila omitida.")
                        conn.rollback() 
                        total_skipped += 1
                    except Exception as e:
                        print(f"ERROR inesperado (CSV Import) al procesar fila {row_num + 1} (Matrícula: {row_data_from_csv.get('matricula', 'N/A')}): {e}. Fila omitida.")
                        total_skipped += 1

                conn.commit()
                cursor.close()
                conn.close()
                flash(f'CSV importado exitosamente. Se insertaron {total_inserted} registros y se omitieron {total_skipped}.', 'success')
                return redirect(url_for('index')) 
            except Exception as e:
                flash(f'Error al procesar el archivo CSV: {e}', 'error')
                return render_template('upload_csv.html')
        else:
            flash('Tipo de archivo no permitido. Por favor, sube un archivo CSV.', 'error')
            return redirect(request.url)
    return render_template('upload_csv.html')


# Rutas de Administrador
@app.route('/admin_dashboard')
@admin_required # Asegura que solo los administradores puedan acceder
def admin_dashboard():
    return render_template('admin_dashboard.html')

@app.route('/admin_users')
@admin_required # Asegura que solo los administradores puedan acceder
def admin_users():
    users = get_all_users()
    return render_template('admin_users.html', users=users)

@app.route('/update_user/<int:user_id>', methods=['POST'])
@admin_required # Solo administradores pueden actualizar usuarios
def update_user(user_id):
    is_approved = request.form.get('is_approved') == 'True' # 'True' como string desde el formulario
    role = request.form.get('role')

    # Validar el rol recibido
    if role not in ['viewer', 'editor', 'admin']: # Roles posibles
        flash('Rol no válido.', 'error')
        return redirect(url_for('admin_users'))

    # Asegurarse de que el admin no pueda desaprobarse o cambiarse el rol a sí mismo
    # current_user.id == user_id: el admin que hace la acción
    # not is_approved: si intenta desaprobarse
    # not current_user.has_role(role): si intenta degradar su propio rol (ej. de admin a editor/viewer)
    if current_user.id == user_id and (not is_approved or not current_user.has_role(role)): # Simplified check
        flash('Un administrador no puede desaprobarse o degradarse a sí mismo.', 'error')
        return redirect(url_for('admin_users'))

    if update_user_status_and_role(user_id, is_approved, role):
        flash(f'Usuario ID {user_id} actualizado exitosamente.', 'success')
    else:
        flash(f'Error al actualizar usuario ID {user_id}.', 'error')
    return redirect(url_for('admin_users'))

# NUEVA RUTA: Eliminar usuario
@app.route('/delete_user/<int:user_id>', methods=['POST'])
@admin_required # Solo administradores pueden eliminar usuarios
def delete_user(user_id):
    # Evitar que un administrador se elimine a sí mismo
    if current_user.id == user_id:
        flash('Un administrador no puede eliminarse a sí mismo.', 'error')
        return redirect(url_for('admin_users'))

    if delete_user_db(user_id):
        flash(f'Usuario ID {user_id} eliminado exitosamente.', 'success')
    else:
        flash(f'Error al eliminar usuario ID {user_id}.', 'error')
    return redirect(url_for('admin_users'))


# NUEVA RUTA: Panel de Estadísticas y KPIs / Patrones de Anomalías
@app.route('/dashboard_stats')
@viewer_required # Cualquier usuario aprobado puede ver las estadísticas
def dashboard_stats():
    # Obtener datos para gráficos/estadísticas
    observations_by_month_year = get_observation_counts_by_month_year()
    status_distribution = get_status_distribution()
    top_recurrent_vessels = get_top_recurrent_vessels(limit=5) # Top 5 embarcaciones más recurrentes
    repeated_infraction_vessels = get_repeated_infraction_vessels(min_infractions=2) # Embarcaciones con 2 o más infracciones

    # Preprocesar datos para facilitar la visualización en la plantilla
    # Para observaciones por mes/año (ej. para un gráfico de barras/líneas)
    monthly_data = {f"{item['year']}-{str(int(item['month'])).zfill(2)}": item['count'] 
                    for item in observations_by_month_year}
    
    # Para distribución de estatus (ej. para un gráfico de pastel)
    status_labels = [
        STATUS_CATEGORIES_INSIDE_ANP[1]['desc'], STATUS_CATEGORIES_INSIDE_ANP[2]['desc'], 
        STATUS_CATEGORIES_INSIDE_ANP[3]['desc'], STATUS_CATEGORIES_INSIDE_ANP[4]['desc'],
        STATUS_CATEGORIES_INSIDE_ANP[5]['desc'], STATUS_CATEGORIES_INSIDE_ANP[6]['desc'],
        "Fuera del Polígono ANP", "Estatus Desconocido"
    ]
    status_data = [0] * len(status_labels)
    
    # Mapear IDs de estatus a descripciones legibles y sus recuentos
    status_desc_map = {
        v['id']: v['desc'] for v in STATUS_CATEGORIES_INSIDE_ANP.values()
    }
    status_desc_map['outside_anp'] = "Fuera del Polígono ANP"

    for item in status_distribution:
        status_id = item['estatus_categoria_id']
        count = item['count']
        if status_id in status_desc_map:
            # Encuentra el índice de la descripción en status_labels
            try:
                idx = status_labels.index(status_desc_map[status_id])
                status_data[idx] = count
            except ValueError:
                # Si no se encuentra (ej. 'unknown_status' que no está en el mapa directo)
                if status_id == 'unknown_status':
                    try:
                        idx = status_labels.index("Estatus Desconocido")
                        status_data[idx] = count
                    except ValueError:
                        pass # Should not happen if "Estatus Desconocido" is in labels
                pass
        else:
            # Handle unexpected or unknown status IDs from DB
            try:
                idx = status_labels.index("Estatus Desconocido")
                status_data[idx] += count # Add to unknown category
            except ValueError:
                pass


    return render_template('dashboard_stats.html',
                           monthly_data=monthly_data,
                           status_labels=status_labels,
                           status_data=status_data,
                           top_recurrent_vessels=top_recurrent_vessels,
                           repeated_infraction_vessels=repeated_infraction_vessels)

# NUEVA RUTA: Perfil de usuario y cambio de contraseña
@app.route('/user_profile', methods=['GET'])
@login_required # Solo usuarios logueados pueden ver su perfil
def user_profile():
    return render_template('user_profile.html') # Renderiza una nueva plantilla para el perfil

@app.route('/change_password', methods=['GET', 'POST'])
@login_required # Solo usuarios logueados pueden cambiar su contraseña
def change_password():
    if request.method == 'POST':
        old_password = request.form['old_password']
        new_password = request.form['new_password']
        confirm_new_password = request.form['confirm_new_password']

        user = get_user_by_id(current_user.id) # Obtener el usuario actual de la DB para verificar la contraseña
        if not user:
            flash("Error: Usuario no encontrado.", 'error')
            return redirect(url_for('change_password'))

        if not check_password_hash(user.password_hash, old_password):
            flash("La contraseña actual es incorrecta.", 'error')
            return render_template('change_password.html')

        if new_password != confirm_new_password:
            flash("La nueva contraseña y su confirmación no coinciden.", 'error')
            return render_template('change_password.html')
        
        if len(new_password) < 6:
            flash('La nueva contraseña debe tener al menos 6 caracteres.', 'error')
            return render_template('change_password.html')

        new_password_hash = generate_password_hash(new_password)
        if update_user_password(current_user.id, new_password_hash):
            flash("Tu contraseña ha sido actualizada exitosamente.", 'success')
            return redirect(url_for('user_profile')) # Redirige al perfil o a alguna página de éxito
        else:
            flash("Error al actualizar la contraseña.", 'error')
            return render_template('change_password.html')
            
    return render_template('change_password.html')

# NUEVA RUTA API: Sugerencias para autocompletado de matrículas y nombres
@app.route('/api/search_suggestions')
@viewer_required # Cualquier usuario aprobado que pueda buscar, puede obtener sugerencias
def search_suggestions():
    query = request.args.get('q', '').strip().lower()
    if not query:
        return jsonify([])

    conn = conectar_db()
    if not conn:
        return jsonify([])
    cursor = conn.cursor()
    
    suggestions = set()
    try:
        # Buscar matrículas
        cursor.execute("SELECT DISTINCT matricula FROM observaciones_embarcaciones WHERE LOWER(matricula) LIKE %s LIMIT 10", (f'%{query}%',))
        for row in cursor.fetchall():
            suggestions.add(row[0])
        
        # Buscar nombres de embarcación
        cursor.execute("SELECT DISTINCT nombre_embarcacion FROM observaciones_embarcaciones WHERE LOWER(nombre_embarcacion) LIKE %s LIMIT 10", (f'%{query}%',))
        for row in cursor.fetchall():
            if row[0]: # Asegurarse de que no sea None
                suggestions.add(row[0])

        # Buscar nombres de patrón
        cursor.execute("SELECT DISTINCT nombre_patron FROM observaciones_embarcaciones WHERE LOWER(nombre_patron) LIKE %s LIMIT 10", (f'%{query}%',))
        for row in cursor.fetchall():
            if row[0] and row[0].lower() != 'n/a': # Asegurarse de que no sea None o 'N/A'
                suggestions.add(row[0])

    except Exception as e:
        print(f"Error al obtener sugerencias de búsqueda: {e}")
    finally:
        if cursor: cursor.close()
        if conn: conn.close()
    
    return jsonify(sorted(list(suggestions)))


if __name__ == '__main__':
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Mexico.1252')
        except locale.Error:
            print("ADVERTENCIA: No se pudo configurar el locale español.")
    app.run(debug=True)
